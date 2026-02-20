"""
Assessment Generator Module

Streamlit page for generating assessment documents (SAQ, PP, CS, etc.)
from a Facilitator Guide (FG) document.

Workflow:
1. Upload FG → Parse to text (pure Python, no AI)
2. Load pre-generated assessment context JSON (from Claude Code skill)
3. Fill assessment templates → Download

All AI reasoning (FG interpretation, assessment generation) is handled by
Claude Code skills using subscription. No API tokens needed.
"""

import streamlit as st
import os
import io
import zipfile
import json
import tempfile
import re
import hashlib

from copy import deepcopy
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from company.company_manager import get_selected_company, get_company_template

# Try to import pymupdf, fallback to pypdf2
PYMUPDF_AVAILABLE = False
try:
    import pymupdf
    PYMUPDF_AVAILABLE = True
except ImportError:
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        pass

# Initialize session state keys
if 'fg_data' not in st.session_state:
    st.session_state['fg_data'] = None
if 'assessment_generated_files' not in st.session_state:
    st.session_state['assessment_generated_files'] = {}


################################################################################
# Helper function for PDF text extraction
################################################################################
def get_pdf_page_count(pdf_path):
    """Get total page count of a PDF file."""
    if PYMUPDF_AVAILABLE:
        doc = pymupdf.open(pdf_path)
        total_pages = doc.page_count
        doc.close()
        return total_pages
    else:
        from PyPDF2 import PdfReader
        reader = PdfReader(pdf_path)
        return len(reader.pages)


def extract_pdf_text(pdf_path):
    """Extract text from PDF using PyMuPDF or PyPDF2 fallback."""
    if PYMUPDF_AVAILABLE:
        doc = pymupdf.open(pdf_path)
        text_content = []
        for page_num in range(doc.page_count):
            page = doc[page_num]
            text = page.get_text()
            if text.strip():
                text_content.append({"page": page_num + 1, "text": text})
        doc.close()
        return text_content
    else:
        from PyPDF2 import PdfReader
        reader = PdfReader(pdf_path)
        text_content = []
        for page_num, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                text_content.append({"page": page_num + 1, "text": text})
        return text_content


################################################################################
# Parse Facilitator Guide Document (Pure Python, no AI)
################################################################################
def parse_fg(fg_path):
    """Parse Facilitator Guide document using python-docx."""
    # Create cache directory
    cache_dir = ".output/fg_cache"
    os.makedirs(cache_dir, exist_ok=True)

    # Generate cache key from file content hash
    with open(fg_path, 'rb') as f:
        file_hash = hashlib.md5(f.read()).hexdigest()
    cache_path = os.path.join(cache_dir, f"{file_hash}.json")

    # Check cache first
    if os.path.exists(cache_path):
        print(f"Found cached FG parse result (hash: {file_hash})")
        with open(cache_path, 'r', encoding='utf-8') as f:
            return f.read()

    # Parse with python-docx
    print(f"Parsing FG document...")
    doc = Document(fg_path)

    # Extract all text from paragraphs and tables
    all_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            all_text.append(para.text.strip())

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                all_text.append(" | ".join(row_text))

    # Create JSON structure for parsed content
    parsed_content = [{"pages": [{"page": 1, "text": "\n".join(all_text)}]}]
    result_json = json.dumps(parsed_content)

    # Save to cache
    with open(cache_path, 'w', encoding='utf-8') as f:
        f.write(result_json)
    print(f"Cached FG parse result (hash: {file_hash})")

    return result_json


def extract_master_k_a_list(fg_markdown):
    """
    Extracts the master list of K and A statements from the FG markdown using regex.
    Returns a dict with 'knowledge' and 'abilities' lists.
    """
    master_k = []
    master_a = []

    # Parse JSON to extract all text content from pages
    try:
        fg_json = json.loads(fg_markdown)
        all_text = ""
        if isinstance(fg_json, list):
            for doc in fg_json:
                if "pages" in doc:
                    for page in doc["pages"]:
                        if "text" in page:
                            all_text += page["text"] + "\n"
    except json.JSONDecodeError:
        all_text = fg_markdown

    # Try multiple patterns to match different formats
    patterns_k = [
        r'(K\d+):\s*(.+?)(?=\n(?:K\d+:|A\d+:|TSC |##|\n\n|$))',
        r'\*\*?(K\d+):\*\*?\s*(.+?)(?=\n\*\*?(?:K\d+:|A\d+:)|\n\n|$)',
        r'(K\d+)\s*[-–]\s*(.+?)(?=\n(?:K\d+|A\d+)|\n\n|$)',
    ]

    patterns_a = [
        r'(A\d+):\s*(.+?)(?=\n(?:K\d+:|A\d+:|TSC |##|\n\n|$))',
        r'\*\*?(A\d+):\*\*?\s*(.+?)(?=\n\*\*?(?:K\d+:|A\d+:)|\n\n|$)',
        r'(A\d+)\s*[-–]\s*(.+?)(?=\n(?:K\d+|A\d+)|\n\n|$)',
    ]

    for pattern in patterns_k:
        for match in re.finditer(pattern, all_text, re.DOTALL | re.MULTILINE):
            k_id = match.group(1)
            k_text = match.group(2).strip().replace('\n', ' ').replace('*', '')
            if not any(k['id'] == k_id for k in master_k):
                master_k.append({"id": k_id, "text": k_text})

    for pattern in patterns_a:
        for match in re.finditer(pattern, all_text, re.DOTALL | re.MULTILINE):
            a_id = match.group(1)
            a_text = match.group(2).strip().replace('\n', ' ').replace('*', '')
            if not any(a['id'] == a_id for a in master_a):
                master_a.append({"id": a_id, "text": a_text})

    print(f"Master K/A List - Found {len(master_k)} K statements, {len(master_a)} A statements")
    return {"knowledge": master_k, "abilities": master_a}


################################################################################
# Parse Slide Deck Document (Pure Python, no AI)
################################################################################
def parse_slides(slides_path):
    """Parse PDF slides and extract text content."""
    print(f"Parsing slides from: {slides_path}")

    total_pages = get_pdf_page_count(slides_path)
    start_page = min(17, total_pages)
    end_page = max(start_page, total_pages - 6)

    slides_content = []

    if PYMUPDF_AVAILABLE:
        doc = pymupdf.open(slides_path)
        for page_num in range(start_page - 1, end_page):
            if page_num < doc.page_count:
                page = doc[page_num]
                text = page.get_text()
                if text.strip():
                    slides_content.append({"page": page_num + 1, "text": text.strip()})
        doc.close()
    else:
        from PyPDF2 import PdfReader
        reader = PdfReader(slides_path)
        for page_num in range(start_page - 1, end_page):
            if page_num < len(reader.pages):
                text = reader.pages[page_num].extract_text() or ""
                if text.strip():
                    slides_content.append({"page": page_num + 1, "text": text.strip()})

    print(f"Extracted text from {len(slides_content)} pages")
    return {"slides": slides_content, "total_pages": total_pages}


################################################################################
# Utility function to ensure answers are always a list
################################################################################
def _ensure_list(answer):
    if isinstance(answer, list):
        return answer
    elif isinstance(answer, str):
        return [answer]
    return []


################################################################################
# Assessment type name mapping
################################################################################
def _get_assessment_full_name(assessment_type: str) -> str:
    """Map assessment type code to full display name for document titles."""
    mapping = {
        "WA (SAQ)": "Written Assessment (SAQ)",
        "WA-SAQ": "Written Assessment (SAQ)",
        "PP": "Practical Performance (PP)",
        "CS": "Case Study (CS)",
        "OQ": "Oral Questioning (OQ)",
        "OI": "Oral Interview (OI)",
        "DEM": "Demonstration (DEM)",
        "RP": "Role Play (RP)",
        "PRJ": "Project (PRJ)",
        "ASGN": "Assignment (ASGN)",
    }
    return mapping.get(assessment_type, assessment_type)


def _get_assessment_long_name(assessment_type: str) -> str:
    """Map assessment type code to long descriptive name for Section B."""
    mapping = {
        "WA (SAQ)": "Written Assessment \u2013 Short Answer Questions (SAQ)",
        "WA-SAQ": "Written Assessment \u2013 Short Answer Questions (SAQ)",
        "PP": "Practical Performance (PP)",
        "CS": "Case Study (CS)",
        "OQ": "Oral Questioning (OQ)",
        "OI": "Oral Interview (OI)",
        "DEM": "Demonstration (DEM)",
        "RP": "Role Play (RP)",
        "PRJ": "Project (PRJ)",
        "ASGN": "Assignment (ASGN)",
    }
    return mapping.get(assessment_type, assessment_type)


def _format_duration_mins(duration: str) -> str:
    """Convert duration string to minutes format (e.g. '1 hr' -> '60 mins')."""
    if not duration:
        return "60 mins"
    d = duration.strip().lower()
    # Already in mins format
    if 'min' in d and 'hr' not in d and 'hour' not in d:
        m = re.match(r'(\d+)', d)
        return f"{m.group(1)} mins" if m else duration
    # Extract hours
    m = re.match(r'(\d+(?:\.\d+)?)\s*(?:hr|hour)', d)
    if m:
        hours = float(m.group(1))
        return f"{int(hours * 60)} mins"
    return duration


def _set_cell_font(cell, size=11, bold=False, font_name='Calibri'):
    """Set font properties for all paragraphs in a table cell."""
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.size = Pt(size)
            r.font.name = font_name
            r.bold = bold


def _build_ref_string(q: dict) -> str:
    """Build inline reference string like (K3) or (A1, A2) from question data."""
    refs = []
    if q.get('knowledge_id'):
        refs.append(q['knowledge_id'])
    if q.get('ability_id'):
        aids = q['ability_id'] if isinstance(q['ability_id'], list) else [q['ability_id']]
        refs.extend(aids)
    if refs:
        return f"({', '.join(refs)})"
    return ""


################################################################################
# Generate documents (Question and Answer papers) - Template filling only
################################################################################
def _build_answer_doc(context: dict, assessment_type: str, questions: list) -> Document:
    """Build an assessment ANSWER document in WSQ client-ready format.

    Matches reference format: Title ('Answers to ...') + horizontal line +
    questions with suggestive answers inside bordered boxes.
    No Section A/B/C, no trainee info, no instructions, no footer.
    """
    doc = Document()

    # -- Set default font --
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    # -- Set margins (1 inch) --
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    course_title = context.get('course_title', '')
    assessment_code = context.get('assessment_code', assessment_type)

    # =========================================================================
    # Title Block (centered, bold)
    # =========================================================================
    # "Answers to [Course Title]"
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(2)
    run = p_title.add_run(f"Answers to {course_title}")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = 'Arial'

    # Assessment type subtitle
    subtitle_text = _get_assessment_full_name(assessment_code)
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(6)
    run = p_sub.add_run(subtitle_text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = 'Arial'

    # Horizontal line separator after title
    p_line = doc.add_paragraph()
    p_line.paragraph_format.space_before = Pt(0)
    p_line.paragraph_format.space_after = Pt(6)
    pPr = p_line._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="000000"/></w:pBdr>')
    pPr.append(pBdr)

    # =========================================================================
    # Questions with Suggestive Answers in bordered boxes
    # =========================================================================
    for idx, q in enumerate(questions, 1):
        # Scenario (italics) — outside the box
        scenario = q.get('scenario', '')
        if scenario:
            p_scen = doc.add_paragraph()
            p_scen.paragraph_format.space_before = Pt(12)
            p_scen.paragraph_format.space_after = Pt(6)
            run = p_scen.add_run(scenario)
            run.italic = True
            run.font.size = Pt(12)
            run.font.name = 'Arial'

        # Question: "Q1. [text] (K3)" — outside the box, regular weight
        question_text = q.get('question_statement', q.get('question', ''))
        ref_str = _build_ref_string(q)

        p_q = doc.add_paragraph()
        p_q.paragraph_format.space_before = Pt(12) if not scenario else Pt(6)
        p_q.paragraph_format.space_after = Pt(6)

        # "Q1." prefix — regular weight (not bold) to match reference
        run_q = p_q.add_run(f"Q{idx}. {question_text}")
        run_q.font.size = Pt(12)
        run_q.font.name = 'Arial'

        # Inline reference (K#) or (A#)
        if ref_str:
            run_ref = p_q.add_run(f" {ref_str}")
            run_ref.font.size = Pt(12)
            run_ref.font.name = 'Arial'

        # Bordered box containing "Suggestive answers" label + answer text
        answers = _ensure_list(q.get('answer', []))
        ans_table = doc.add_table(rows=1, cols=1)
        ans_table.style = 'Table Grid'
        cell = ans_table.cell(0, 0)
        cell.paragraphs[0].clear()

        # "Suggestive answers (not exhaustive):" — first line inside box
        p_label = cell.paragraphs[0]
        p_label.paragraph_format.space_before = Pt(6)
        p_label.paragraph_format.space_after = Pt(6)
        run = p_label.add_run("Suggestive answers (not exhaustive):")
        run.font.size = Pt(12)
        run.font.name = 'Arial'

        # Answer text inside the same box
        if answers:
            # Check if answers look like bullet items (short items)
            # Use bullets if there are 3+ short items (like the reference K9 example)
            use_bullets = len(answers) >= 3 and all(len(a) < 80 for a in answers)

            if use_bullets:
                for ans in answers:
                    p_ans = cell.add_paragraph()
                    p_ans.paragraph_format.space_before = Pt(2)
                    p_ans.paragraph_format.space_after = Pt(2)
                    # Use list style indent with bullet
                    p_ans.paragraph_format.left_indent = Inches(0.5)
                    run = p_ans.add_run(f"\u25cf {ans}")
                    run.font.size = Pt(12)
                    run.font.name = 'Arial'
            else:
                # Join all answer items into a single paragraph (like reference format)
                combined_answer = " ".join(answers)
                p_ans = cell.add_paragraph()
                p_ans.paragraph_format.space_before = Pt(6)
                p_ans.paragraph_format.space_after = Pt(6)
                run = p_ans.add_run(combined_answer)
                run.font.size = Pt(12)
                run.font.name = 'Arial'

    return doc


def _build_assessment_doc(context: dict, assessment_type: str, questions: list, include_answers: bool) -> Document:
    """Build an assessment Word document in WSQ client-ready format."""

    # Answer documents use a completely different simpler format
    if include_answers:
        return _build_answer_doc(context, assessment_type, questions)

    doc = Document()

    # -- Set default font --
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    # -- Set margins (1 inch = 914400 EMU) --
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    course_title = context.get('course_title', '')
    duration = context.get('duration', '')
    assessment_code = context.get('assessment_code', assessment_type)

    # =========================================================================
    # Title Block (centered)
    # =========================================================================
    # Course title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(2)
    run = p_title.add_run(course_title)
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = 'Arial'

    # Assessment type subtitle
    subtitle_text = _get_assessment_full_name(assessment_code)
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(12)
    run = p_sub.add_run(subtitle_text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = 'Arial'

    # =========================================================================
    # A: Trainee Information
    # =========================================================================
    p_a = doc.add_paragraph()
    p_a.paragraph_format.space_before = Pt(6)
    p_a.paragraph_format.space_after = Pt(6)
    run = p_a.add_run("A: Trainee Information:")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Arial'

    blank_line = "_" * 30
    short_blank = "_" * 15

    p = doc.add_paragraph()
    run = p.add_run(f"Trainee Name (as Per NRIC): {blank_line}")
    run.font.size = Pt(12)
    run.font.name = 'Arial'

    p = doc.add_paragraph()
    run = p.add_run(f"Last three digits and alphabet of NRIC/FIN: {short_blank}")
    run.font.size = Pt(12)
    run.font.name = 'Arial'

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(f"Date: {short_blank}")
    run.font.size = Pt(12)
    run.font.name = 'Arial'

    # =========================================================================
    # B: Assessment Instruction
    # =========================================================================
    p_b = doc.add_paragraph()
    p_b.paragraph_format.space_before = Pt(6)
    p_b.paragraph_format.space_after = Pt(6)
    run = p_b.add_run("B: Assessment Instruction")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Arial'

    assessment_long = _get_assessment_long_name(assessment_code)
    num_questions = len(questions)

    p = doc.add_paragraph()
    run = p.add_run(f"This is the {assessment_long}")
    run.font.size = Pt(12)
    run.font.name = 'Arial'

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"Duration: {_format_duration_mins(duration)}")
    run.font.size = Pt(12)
    run.font.name = 'Arial'

    instructions = [
        f"1. The assessor will pass the questions in hard copy to you. There are {num_questions} questions. You need to answer all the questions.",
        "2. This is an open-book exam that must be completed individually.",
        "3. You need to get all answers correct to be competent.",
    ]
    for instr in instructions:
        p = doc.add_paragraph()
        run = p.add_run(instr)
        run.font.size = Pt(12)
        run.font.name = 'Arial'

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("Submission Procedure:")
    run.font.size = Pt(12)
    run.font.name = 'Arial'

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("1. Please pass the hard copy to the assessor after completion.")
    run.font.size = Pt(12)
    run.font.name = 'Arial'

    # =========================================================================
    # C: Questions and Answers
    # =========================================================================
    p_c = doc.add_paragraph()
    p_c.paragraph_format.space_before = Pt(6)
    p_c.paragraph_format.space_after = Pt(6)
    run = p_c.add_run("C: Questions and Answers")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Arial'

    for idx, q in enumerate(questions, 1):
        # Scenario (italics)
        scenario = q.get('scenario', '')
        if scenario:
            p_scen = doc.add_paragraph()
            p_scen.paragraph_format.space_before = Pt(12)
            p_scen.paragraph_format.space_after = Pt(6)
            run = p_scen.add_run(scenario)
            run.italic = True
            run.font.size = Pt(12)
            run.font.name = 'Arial'

        # Question: "Q1. [text] (K3)"
        question_text = q.get('question_statement', q.get('question', ''))
        ref_str = _build_ref_string(q)

        p_q = doc.add_paragraph()
        p_q.paragraph_format.space_before = Pt(12) if not scenario else Pt(6)
        p_q.paragraph_format.space_after = Pt(6)

        # Bold "Q1." prefix
        run_num = p_q.add_run(f"Q{idx}. ")
        run_num.bold = True
        run_num.font.size = Pt(12)
        run_num.font.name = 'Arial'

        # Question text
        run_text = p_q.add_run(question_text)
        run_text.font.size = Pt(12)
        run_text.font.name = 'Arial'

        # Inline reference (K#) or (A#)
        if ref_str:
            run_ref = p_q.add_run(f" {ref_str}")
            run_ref.font.size = Pt(12)
            run_ref.font.name = 'Arial'

        # Answer label
        p_ans_label = doc.add_paragraph()
        p_ans_label.paragraph_format.space_before = Pt(6)
        p_ans_label.paragraph_format.space_after = Pt(4)
        run = p_ans_label.add_run("Answer:")
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Arial'

        # Empty bordered answer box for writing
        box_table = doc.add_table(rows=1, cols=1)
        box_table.style = 'Table Grid'
        cell = box_table.cell(0, 0)
        cell.text = ""
        for _ in range(6):
            cell.add_paragraph("")
        for cp in cell.paragraphs:
            cp.paragraph_format.space_before = Pt(2)
            cp.paragraph_format.space_after = Pt(2)
            for r in cp.runs:
                r.font.size = Pt(12)
                r.font.name = 'Arial'

    # =========================================================================
    # Footer: For Official Use Only
    # =========================================================================
    doc.add_paragraph()  # spacing

    # Horizontal line
    p_line = doc.add_paragraph()
    p_line.paragraph_format.space_before = Pt(12)
    p_line.paragraph_format.space_after = Pt(6)
    pPr = p_line._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="000000"/></w:pBdr>')
    pPr.append(pBdr)

    # "For Official Use Only" heading
    p_official = doc.add_paragraph()
    p_official.paragraph_format.space_before = Pt(6)
    p_official.paragraph_format.space_after = Pt(6)
    run = p_official.add_run("For Official Use Only")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Arial'

    blank = "_" * 10
    long_blank = "_" * 15

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(f"Grade: {blank}(C / NYC)")
    run.font.size = Pt(12)
    run.font.name = 'Arial'

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(f"Assessor Name: {long_blank}")
    run.font.size = Pt(12)
    run.font.name = 'Arial'
    run = p.add_run(f"    Assessor NRIC: {long_blank}")
    run.font.size = Pt(12)
    run.font.name = 'Arial'

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(f"Date: {long_blank}")
    run.font.size = Pt(12)
    run.font.name = 'Arial'
    run = p.add_run(f"    Signature: {long_blank}")
    run.font.size = Pt(12)
    run.font.name = 'Arial'

    return doc


def generate_documents(context: dict, assessment_type: str, output_dir: str, company: dict = None) -> dict:
    """Generate assessment question paper and answer key as Word documents."""
    os.makedirs(output_dir, exist_ok=True)

    selected_company = company if company is not None else get_selected_company()
    context['company_name'] = selected_company.get('name', 'Tertiary Infotech Academy Pte Ltd')
    context['company_uen'] = selected_company.get('uen', '201200696W')
    context['company_address'] = selected_company.get('address', '')

    questions = context.get("questions", [])

    # Build question paper (no answers)
    q_doc = _build_assessment_doc(context, assessment_type, questions, include_answers=False)
    q_file = tempfile.NamedTemporaryFile(delete=False, suffix=f"_{assessment_type}_Questions.docx")
    q_file.close()
    q_doc.save(q_file.name)

    # Build answer key (with answers)
    a_doc = _build_assessment_doc(context, assessment_type, questions, include_answers=True)
    a_file = tempfile.NamedTemporaryFile(delete=False, suffix=f"_{assessment_type}_Answers.docx")
    a_file.close()
    a_doc.save(a_file.name)

    return {
        "ASSESSMENT_TYPE": assessment_type,
        "QUESTION": q_file.name,
        "ANSWER": a_file.name,
    }


################################################################################
# Streamlit app
################################################################################
DEFAULT_ASSESSMENT_PROMPT = """Generate assessment questions for the course: {course_title}

Assessment types to generate: {assessment_types}

Use the following course data including Learning Units, K/A statements,
topics, and assessment method details to create comprehensive questions.

--- COURSE CONTEXT ---
{course_context}
--- END ---

{master_ka_list}

Generate comprehensive assessment questions following the schema in your instructions.
Return ONLY the JSON object."""


def app():
    st.title("Generate Assessment")

    selected_company = get_selected_company()
    extracted_info = st.session_state.get('extracted_course_info')

    # Prompt Templates (editable, collapsed)
    from utils.prompt_template_editor import render_prompt_templates
    render_prompt_templates("assessment", "Prompt Templates (Assessment)")

    # ----- Dependency check -----
    from utils.agent_runner import submit_agent_job, get_job
    from utils.agent_status import render_page_job_status

    extract_job = get_job("extract_course_info")
    if extract_job and extract_job.get("status") == "running":
        st.info("Course info extraction is still running. Please wait for it to complete.")

    # Generate Assessment
    if st.button("Generate Assessment", type="primary"):
        if not extracted_info:
            st.error("Please extract course info first.")
        elif extract_job and extract_job.get("status") == "running":
            st.warning("Please wait for course info extraction to complete.")
        else:
            # Pre-capture data for background thread (can't access session_state from thread)
            _company = dict(selected_company)
            _extracted_info = dict(extracted_info)

            def _fill_assessment_templates(result):
                """Post-process: fill assessment templates in background thread."""
                if not result or not isinstance(result, dict):
                    return {"fg_data": result, "generated_files": {}}

                assessment_types = result.get('assessment_types', [])
                generated_files = {}
                for assessment in assessment_types:
                    a_type = assessment.get('type', assessment.get('code', 'Unknown'))
                    questions = assessment.get('questions', [])
                    if not questions:
                        continue
                    try:
                        doc_context = {
                            "course_title": result.get('course_title', ''),
                            "duration": assessment.get('duration', ''),
                            "assessment_code": assessment.get('code', a_type),
                            "questions": questions,
                        }
                        files = generate_documents(doc_context, a_type, ".output", company=_company)
                        generated_files[a_type] = files
                    except Exception as e:
                        print(f"Error generating {a_type}: {e}")

                return {"fg_data": result, "generated_files": generated_files}

            from courseware_agents.assessment_generator import generate_assessments
            job = submit_agent_job(
                key="generate_assessment",
                label="Generate Assessment",
                async_fn=generate_assessments,
                kwargs={"course_context": _extracted_info},
                post_process=_fill_assessment_templates,
            )

            if job is None:
                st.warning("Assessment generation is already running.")
            else:
                st.rerun()

    # ----- Agent Status -----
    def _on_assessment_complete(job):
        post = job.get("post_results") or {}
        if post:
            st.session_state['fg_data'] = post.get("fg_data")
            st.session_state['assessment_generated_files'] = post.get("generated_files", {})

    job_status = render_page_job_status(
        "generate_assessment",
        on_complete=_on_assessment_complete,
        running_message="AI Agent generating assessments...",
    )

    # Download
    generated_files = st.session_state.get('assessment_generated_files', {})
    if generated_files and any(
        ((file_paths.get('QUESTION') and os.path.exists(file_paths.get('QUESTION'))) or
        (file_paths.get('ANSWER') and os.path.exists(file_paths.get('ANSWER'))))
        for file_paths in generated_files.values()
    ):
        course_title = "Course Title"
        if st.session_state.get('fg_data'):
            course_title = st.session_state['fg_data'].get("course_title", "Course Title")

        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
            for assessment_type, file_paths in generated_files.items():
                q_path = file_paths.get('QUESTION')
                a_path = file_paths.get('ANSWER')

                if q_path and os.path.exists(q_path):
                    q_file_name = f"{assessment_type} - {course_title}.docx"
                    zipf.write(q_path, arcname=q_file_name)

                if a_path and os.path.exists(a_path):
                    a_file_name = f"Answer to {assessment_type} - {course_title}.docx"
                    zipf.write(a_path, arcname=a_file_name)

        zip_buffer.seek(0)

        st.download_button(
            label="Download All Assessments (ZIP)",
            data=zip_buffer.getvalue(),
            file_name="assessments.zip",
            mime="application/zip"
        )

    # =========================================================================
    # Add Assessment to AP (Annex)
    # =========================================================================
    st.markdown("---")
    st.subheader("Add Assessment to AP")

    from add_assessment_to_ap.annex_assessment_v2 import merge_documents

    plan_file = st.file_uploader(
        "Upload Assessment Plan (.docx)",
        type=["docx"],
        key="plan_upload",
        help="Upload the main Assessment Plan document"
    )

    # Build assessment types dynamically from extracted course info
    assessment_types = []
    if extracted_info:
        methods = extracted_info.get('Assessment_Methods_Details', [])
        for m in methods:
            abbr = m.get('Method_Abbreviation', '')
            name = m.get('Assessment_Method', '')
            if abbr == 'WA-SAQ':
                assessment_types.append("WA (SAQ)")
            elif abbr:
                assessment_types.append(abbr)
            elif name:
                assessment_types.append(name)
    # Fallback if no course info
    if not assessment_types:
        assessment_types = ["WA (SAQ)", "PP", "CS", "Oral Questioning"]
    assessment_files = {}

    for a_type in assessment_types:
        with st.expander(f"{a_type}", expanded=False):
            col1, col2 = st.columns(2)
            with col1:
                q_file = st.file_uploader(
                    "Question Paper", type=["docx"], key=f"q_{a_type}",
                    help=f"Upload the question paper for {a_type}"
                )
            with col2:
                a_file = st.file_uploader(
                    "Answer Paper", type=["docx"], key=f"a_{a_type}",
                    help=f"Upload the answer paper for {a_type}"
                )
            if q_file or a_file:
                assessment_files[a_type] = {"question": q_file, "answer": a_file}

    if st.button("Generate Merged Document", type="primary", key="merge_btn"):
        if not plan_file:
            st.error("Please upload an Assessment Plan document first.")
        elif not assessment_files:
            st.error("Please upload at least one Question or Answer paper.")
        else:
            try:
                with st.spinner("Merging documents..."):
                    merged_doc_bytes = merge_documents(plan_file, assessment_files)
                st.success("Document merged successfully!")
                base_name = os.path.splitext(plan_file.name)[0]
                st.download_button(
                    label="Download Merged Document",
                    data=merged_doc_bytes,
                    file_name=f"{base_name}_with_annex.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="download_merged"
                )
            except Exception as e:
                st.error(f"Error merging documents: {e}")
