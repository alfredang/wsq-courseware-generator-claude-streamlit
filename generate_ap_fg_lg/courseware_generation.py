from generate_ap_fg_lg.utils.agentic_LG import generate_learning_guide
from generate_ap_fg_lg.utils.agentic_AP import generate_assessment_documents
from generate_ap_fg_lg.utils.agentic_LP import generate_lesson_plan
from generate_ap_fg_lg.utils.agentic_FG import generate_facilitators_guide
import os
import io
import zipfile
import tempfile
import json
import re
from datetime import datetime
import streamlit as st
from docx import Document
from docx import Document as DocxDocument
import openpyxl
from pydantic import BaseModel
from typing import List, Optional
from utils.helpers import save_uploaded_file, parse_json_content, copy_to_courseware
import asyncio
from generate_ap_fg_lg.utils.organization_utils import (
    load_organizations,
    save_organizations,
    add_organization,
    update_organization,
    delete_organization,
    Organization
)
from streamlit_modal import Modal

# Initialize session state variables
if 'lg_output' not in st.session_state:
    st.session_state['lg_output'] = None
if 'ap_output' not in st.session_state:
    st.session_state['ap_output'] = None
if 'lp_output' not in st.session_state:
    st.session_state['lp_output'] = None
if 'fg_output' not in st.session_state:
    st.session_state['fg_output'] = None
if 'context' not in st.session_state:
    st.session_state['context'] = None
if 'asr_output' not in st.session_state:
    st.session_state['asr_output'] = None

############################################################
# Pydantic Models
############################################################
class Topic(BaseModel):
    Topic_Title: str
    Bullet_Points: List[str]

class KDescription(BaseModel):
    K_number: str
    Description: str

class ADescription(BaseModel):
    A_number: str
    Description: str

class LearningUnit(BaseModel):
    LU_Title: str
    Topics: List[Topic]
    LO: str
    K_numbering_description: List[KDescription]
    A_numbering_description: List[ADescription]
    Assessment_Methods: List[str]
    Instructional_Methods: List[str]

class EvidenceDetail(BaseModel):
    LO: str
    Evidence: str

class AssessmentMethodDetail(BaseModel):
    Assessment_Method: str
    Method_Abbreviation: str
    Total_Delivery_Hours: str
    Assessor_to_Candidate_Ratio: List[str]
    Evidence: Optional[List[EvidenceDetail]] = None
    Submission: Optional[List[str]] = None
    Marking_Process: Optional[List[str]] = None
    Retention_Period: Optional[str] = None

class CourseData(BaseModel):
    Date: str
    Year: str
    Name_of_Organisation: str
    Course_Title: str
    TSC_Title: str
    TSC_Code: str
    TSC_Category: Optional[str] = None
    TSC_Description: Optional[str] = None
    TSC_Sector: Optional[str] = None
    Proficiency_Level: Optional[str] = None
    Proficiency_Description: Optional[str] = None
    Skills_Framework: Optional[str] = None
    Total_Training_Hours: str
    Total_Assessment_Hours: str
    Total_Course_Duration_Hours: str
    Learning_Units: List[LearningUnit]
    Assessment_Methods_Details: List[AssessmentMethodDetail]

class Session(BaseModel):
    Time: str
    instruction_title: str
    bullet_points: List[str]
    Instructional_Methods: str
    Resources: str

class DayLessonPlan(BaseModel):
    Day: str
    Sessions: List[Session]

class LessonPlan(BaseModel):
    lesson_plan: List[DayLessonPlan]


############################################################
# Course Proposal Document Parsing (Pure Python, no AI)
############################################################

def parse_cp_document(uploaded_file):
    """
    Parses a Course Proposal (CP) document and returns its content as text.

    For Word CP (.docx): Trims to Part 1 through Part 4.
    For Excel CP (.xlsx): Trims to Course Particulars through Declarations.

    Args:
        uploaded_file: The file uploaded via st.file_uploader.

    Returns:
        str: Trimmed text string containing the parsed document content.
    """
    with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[1]) as tmp:
        tmp.write(uploaded_file.read())
        temp_file_path = tmp.name

    try:
        ext = os.path.splitext(temp_file_path)[1].lower()
        text_content = []

        if ext == ".docx":
            doc = DocxDocument(temp_file_path)
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))

        elif ext == ".xlsx":
            wb = openpyxl.load_workbook(temp_file_path, data_only=True)
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                text_content.append(f"## {sheet}")
                for row in ws.iter_rows(values_only=True):
                    row_text = [str(cell) if cell is not None else "" for cell in row]
                    if any(row_text):
                        text_content.append(" | ".join(row_text))
            wb.close()

        markdown_text = "\n".join(text_content)

        # Trim based on file extension
        if ext == ".docx":
            start_pattern = re.compile(r"Part\s*1.*?Particulars\s+of\s+Course", re.IGNORECASE)
            end_pattern = re.compile(r"Part\s*4.*?Facilities\s+and\s+Resources", re.IGNORECASE)
        elif ext == ".xlsx":
            start_pattern = re.compile(r"1\s*-\s*Course\s*Particulars", re.IGNORECASE)
            end_pattern = re.compile(r"4\s*-\s*Declarations", re.IGNORECASE)
        else:
            start_pattern = None
            end_pattern = None

        if start_pattern and end_pattern:
            start_match = start_pattern.search(markdown_text)
            end_match = end_pattern.search(markdown_text)
            if start_match and end_match and end_match.start() > start_match.start():
                markdown_text = markdown_text[start_match.start():end_match.start()].strip()

    finally:
        os.remove(temp_file_path)

    return markdown_text


############################################################
# TSC Field Defaults (Pure Python)
############################################################

SECTOR_MAPPING = {
    "LOG": ("Logistics", "Skills Framework for Logistics"),
    "ICT": ("Infocomm Technology", "Skills Framework for Infocomm Technology"),
    "FIN": ("Financial Services", "Skills Framework for Financial Services"),
    "HR": ("Human Resource", "Skills Framework for Human Resource"),
    "MFG": ("Manufacturing", "Skills Framework for Manufacturing"),
    "RET": ("Retail", "Skills Framework for Retail"),
    "SEC": ("Security", "Skills Framework for Security"),
    "TOU": ("Tourism", "Skills Framework for Tourism"),
    "HEA": ("Healthcare", "Skills Framework for Healthcare"),
    "EDU": ("Education", "Skills Framework for Training and Adult Education"),
    "HAS": ("Hotel and Accommodation Services", "Skills Framework for Hotel and Accommodation Services"),
    "FBS": ("Food Services", "Skills Framework for Food Services"),
    "ATT": ("Attractions", "Skills Framework for Attractions"),
    "TAE": ("Training and Adult Education", "Skills Framework for Training and Adult Education"),
    "SER": ("Services", "Skills Framework for Services"),
    "AIR": ("Air Transport", "Skills Framework for Air Transport"),
    "SEA": ("Sea Transport", "Skills Framework for Sea Transport"),
    "LND": ("Land Transport", "Skills Framework for Land Transport"),
    "ENE": ("Energy and Chemicals", "Skills Framework for Energy and Chemicals"),
    "AER": ("Aerospace", "Skills Framework for Aerospace"),
    "BIO": ("Biopharmaceutical Manufacturing", "Skills Framework for Biopharmaceutical Manufacturing"),
    "MED": ("Media", "Skills Framework for Media"),
    "DES": ("Design", "Skills Framework for Design"),
    "BCE": ("Built Environment", "Skills Framework for Built Environment"),
    "MAR": ("Marine and Offshore", "Skills Framework for Marine and Offshore"),
    "PRE": ("Precision Engineering", "Skills Framework for Precision Engineering"),
    "WSH": ("Workplace Safety and Health", "Skills Framework for Workplace Safety and Health"),
    "PUB": ("Public Service", "Skills Framework for Public Service"),
    "SOC": ("Social Service", "Skills Framework for Social Service"),
    "EAC": ("Early Childhood", "Skills Framework for Early Childhood"),
}


def apply_tsc_defaults(context: dict) -> dict:
    """Apply default values for TSC fields based on TSC Code mapping."""
    def is_empty(val):
        return val is None or val == "" or val == "null" or val == "None"

    tsc_code = context.get("TSC_Code", "")
    if not tsc_code:
        return context

    sector_abbr = tsc_code.split("-")[0] if "-" in tsc_code else ""

    if sector_abbr in SECTOR_MAPPING:
        sector_name, skills_framework_full = SECTOR_MAPPING[sector_abbr]

        if is_empty(context.get("Skills_Framework")) or len(str(context.get("Skills_Framework", ""))) <= 5:
            context["Skills_Framework"] = skills_framework_full
        if is_empty(context.get("TSC_Sector")) or len(str(context.get("TSC_Sector", ""))) <= 5:
            context["TSC_Sector"] = sector_name
        if is_empty(context.get("TSC_Category")) or len(str(context.get("TSC_Category", ""))) <= 5:
            context["TSC_Category"] = sector_name
        context["TSC_Sector_Abbr"] = skills_framework_full

    # Extract proficiency level from TSC Code
    level_match = re.search(r'-(\d+)\.\d+$', tsc_code)
    if level_match and is_empty(context.get("Proficiency_Level")):
        context["Proficiency_Level"] = f"Level {level_match.group(1)}"

    # Generate TSC_Description from Learning Outcomes if not set
    if is_empty(context.get("TSC_Description")):
        learning_units = context.get("Learning_Units", [])
        if learning_units:
            for lu in learning_units:
                lo = lu.get("LO", "")
                if lo:
                    lo_text = lo.split(": ", 1)[-1] if ": " in lo else lo
                    context["TSC_Description"] = lo_text
                    break

    if is_empty(context.get("TSC_Description")):
        tsc_title = context.get("TSC_Title", "")
        course_title = context.get("Course_Title", "")
        if tsc_title:
            context["TSC_Description"] = f"Apply knowledge and skills in {tsc_title.lower()} to meet organizational and industry requirements."
        elif course_title:
            context["TSC_Description"] = f"Apply knowledge and skills in {course_title.lower()} to meet organizational and industry requirements."

    # Copy TSC_Description to Proficiency_Description for template compatibility
    if context.get("TSC_Description") and is_empty(context.get("Proficiency_Description")):
        context["Proficiency_Description"] = context.get("TSC_Description")

    # Final cleanup
    tsc_title = context.get("TSC_Title", "")
    tsc_description = context.get("TSC_Description", tsc_title)
    defaults = {
        "TSC_Category": tsc_title,
        "TSC_Sector": tsc_title,
        "TSC_Description": tsc_title,
        "Skills_Framework": "",
        "Proficiency_Level": "",
        "Proficiency_Description": tsc_description,
        "TSC_Sector_Abbr": context.get("Skills_Framework", ""),
    }
    for field, default_value in defaults.items():
        if is_empty(context.get(field)):
            context[field] = default_value

    return context


############################################################
# Streamlit App
############################################################

def app():
    """Streamlit page for courseware document generation."""

    st.title("Generate AP/FG/LG")

    # Auto-load context from Extract Course Info
    extracted_info = st.session_state.get('extracted_course_info')
    if extracted_info and not st.session_state.get('context'):
        context = apply_tsc_defaults(dict(extracted_info))
        st.session_state['context'] = context

    # Get organization from sidebar
    sidebar_selected_company = st.session_state.get('selected_company', None)
    if sidebar_selected_company:
        selected_org = sidebar_selected_company['name']
    else:
        org_list = load_organizations()
        org_names = [org["name"] for org in org_list] if org_list else []
        selected_org = org_names[0] if org_names else None

    # ----- Course Info Status -----
    from utils.agent_runner import get_job
    extract_job = get_job("extract_course_info")

    if extract_job and extract_job.get("status") == "running":
        st.info("Course info extraction is still running. Please wait for it to complete.")
    elif st.session_state.get('context'):
        ctx = st.session_state['context']
        course_title = ctx.get('Course_Title', 'N/A')
        num_lu = len(ctx.get('Learning_Units', []))
        num_topics = sum(len(lu.get('Topics', [])) for lu in ctx.get('Learning_Units', []))
        num_assess = len(ctx.get('Assessment_Methods_Details', []))
        st.success(
            f"**{course_title}** | "
            f"{num_lu} learning unit(s) | "
            f"{num_topics} topic(s) | "
            f"{num_assess} assessment method(s)"
        )
    else:
        st.warning("No course info loaded. Please extract course info first.")

    # Prompt templates (editable, collapsed)
    from utils.prompt_template_editor import render_prompt_templates
    render_prompt_templates("courseware", "Prompt Templates (AP/FG/LG)")

    # Select and generate documents
    st.subheader("Select Document(s) to Generate")
    generate_lg = st.checkbox("Learning Guide (LG)", value=True)
    generate_ap = st.checkbox("Assessment Plan (AP)", value=True)
    generate_fg = st.checkbox("Facilitator's Guide (FG)", value=True)

    if st.button("Generate Documents"):
        context = st.session_state.get('context')

        if context is None:
            st.error("Please extract course info first.")
            return

        if not selected_org:
            st.error("Please select a company from the sidebar.")
            return

        # Reset previous outputs
        st.session_state['lg_output'] = None
        st.session_state['ap_output'] = None
        st.session_state['asr_output'] = None
        st.session_state['fg_output'] = None

        # Add metadata
        current_datetime = datetime.now()
        context["Date"] = current_datetime.strftime("%d %b %Y")
        context["Year"] = current_datetime.year
        context["TGS_Ref_No"] = context.get("TGS_Ref_No", "")

        org_list = load_organizations()
        selected_org_data = next((org for org in org_list if org["name"] == selected_org), None)
        if selected_org_data:
            context["UEN"] = selected_org_data["uen"]

        st.session_state['context'] = context

        # Generate Learning Guide
        if generate_lg:
            try:
                with st.spinner('Generating Learning Guide...'):
                    lg_output = generate_learning_guide(context, selected_org)
                if lg_output:
                    st.success("Learning Guide generated successfully!")
                    st.session_state['lg_output'] = lg_output
            except Exception as e:
                st.error(f"Error generating Learning Guide: {e}")

        # Generate Assessment Plan
        if generate_ap:
            try:
                with st.spinner('Generating Assessment Plan and Assessment Summary Record...'):
                    ap_output, asr_output = generate_assessment_documents(context, selected_org)
                if ap_output:
                    st.success("Assessment Plan generated successfully!")
                    st.session_state['ap_output'] = ap_output
                if asr_output:
                    st.success("Assessment Summary Record generated successfully!")
                    st.session_state['asr_output'] = asr_output
            except Exception as e:
                st.error(f"Error generating Assessment Documents: {e}")

        # Generate Facilitator's Guide
        if generate_fg:
            try:
                with st.spinner("Generating Facilitator's Guide..."):
                    fg_output = generate_facilitators_guide(context, selected_org)
                if fg_output:
                    st.success("Facilitator's Guide generated successfully!")
                    st.session_state['fg_output'] = fg_output
            except Exception as e:
                st.error(f"Error generating Facilitator's Guide: {e}")

    # Download section
    if any([
        st.session_state.get('lg_output'),
        st.session_state.get('ap_output'),
        st.session_state.get('asr_output'),
        st.session_state.get('fg_output')
    ]):
        st.subheader("Download All Generated Documents as ZIP")

        def sanitize_filename(name):
            if not name:
                return "Document"
            name = str(name).strip()
            invalid_chars = ['/', '\\', ':', '*', '?', '"', '<', '>', '|', '\n', '\r', '\t', '&', '%', '$', '#', '@', '!', '^', '~', '`', "'", ';', ',', '(', ')', '[', ']', '{', '}']
            for char in invalid_chars:
                name = name.replace(char, '_')
            name = re.sub(r'[\s_]+', '_', name)
            name = name.strip('_')
            name = ''.join(c if ord(c) < 128 else '_' for c in name)
            name = name[:50] if len(name) > 50 else name
            return name if name else "Document"

        ctx = st.session_state.get('context', {}) or {}
        course_title = sanitize_filename(ctx.get('Course_Title', 'Course'))
        tgs_ref_no = sanitize_filename(ctx.get('TGS_Ref_No', ''))

        # Copy generated files to Courseware subfolders
        _file_map = {
            "lg_output": ("LG", "Learner Guide"),
            "ap_output": ("Assessment_Plan", "Assessment Plan"),
            "asr_output": ("Assessment_Summary_Record", "Assessment Plan"),
            "fg_output": ("FG", "Facilitator Guide"),
        }
        for key, (prefix, subfolder) in _file_map.items():
            fpath = st.session_state.get(key)
            if fpath and os.path.exists(fpath):
                if tgs_ref_no:
                    fname = f"{prefix}_{tgs_ref_no}_{course_title}_v1.docx"
                else:
                    fname = f"{prefix}_{course_title}_v1.docx"
                copy_to_courseware(fpath, subfolder, fname, ctx)

        zip_buffer = io.BytesIO()
        try:
            with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
                def add_file(file_path, prefix):
                    if file_path and os.path.exists(file_path):
                        if tgs_ref_no:
                            file_name = f"{prefix}_{tgs_ref_no}_{course_title}_v1.docx"
                        else:
                            file_name = f"{prefix}_{course_title}_v1.docx"
                        with open(file_path, 'rb') as f:
                            zipf.writestr(file_name, f.read())

                add_file(st.session_state.get('lg_output'), "LG")
                add_file(st.session_state.get('ap_output'), "Assessment_Plan")
                add_file(st.session_state.get('asr_output'), "Assessment_Summary_Record")
                add_file(st.session_state.get('fg_output'), "FG")

            zip_buffer.seek(0)
            zip_filename = f"courseware_{course_title}.zip" if course_title else "courseware_documents.zip"

            st.download_button(
                label="Download All Documents (ZIP)",
                data=zip_buffer.getvalue(),
                file_name=zip_filename,
                mime="application/zip"
            )
        except Exception as e:
            st.error(f"Error creating ZIP file: {e}")
