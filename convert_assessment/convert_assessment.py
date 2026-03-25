"""
Document Conversion Module

Streamlit page for batch converting existing documents to the
standardised WSQ client-ready format.

Tabs:
1. Convert Assessment — question papers & answer keys
2. Convert Courseware — AP, FG, LG, ASR
3. Convert Lesson Plan — LP
"""

import streamlit as st
import os
import io
import tempfile
import zipfile

from docx import Document
from courseware_agents.base import run_agent_json
from generate_assessment.assessment_generation import (
    _build_assessment_doc,
    _get_assessment_full_name,
)


# ──────────────────────────────────────────────
# Shared helpers
# ──────────────────────────────────────────────

def _extract_docx_text(docx_file) -> str:
    """Extract all text from a DOCX file (uploaded file or path)."""
    doc = Document(docx_file)
    all_text = []

    for para in doc.paragraphs:
        if para.text.strip():
            all_text.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                all_text.append(" | ".join(row_text))

    return "\n".join(all_text)


def _convert_and_download(uploaded_files, extract_fn, build_fn, session_key, zip_name):
    """Generic convert-and-download flow."""
    if not uploaded_files:
        return

    st.markdown(f"**{len(uploaded_files)} file(s) uploaded:**")
    for f in uploaded_files:
        st.markdown(f"- {f.name}")

    if st.button("Convert All to WSQ Format", type="primary", key=f"btn_{session_key}"):
        converted_files = {}
        progress_bar = st.progress(0)
        status_text = st.empty()
        total = len(uploaded_files)

        for i, uploaded_file in enumerate(uploaded_files):
            status_text.text(f"Converting {i+1}/{total}: {uploaded_file.name}...")
            progress_bar.progress(i / total)

            try:
                uploaded_file.seek(0)
                doc_text = _extract_docx_text(uploaded_file)
                if not doc_text.strip():
                    st.error(f"Could not extract text from {uploaded_file.name}")
                    continue

                import asyncio
                extracted = asyncio.run(extract_fn(doc_text))
                if not extracted or not isinstance(extracted, dict):
                    st.error(f"AI could not parse {uploaded_file.name}")
                    continue

                result = build_fn(extracted, uploaded_file.name)
                if result:
                    converted_files.update(result)

            except Exception as e:
                st.error(f"Error converting {uploaded_file.name}: {e}")

        progress_bar.progress(1.0)
        status_text.text(f"Done! {len(converted_files)}/{total} files converted.")

        if converted_files:
            st.session_state[session_key] = converted_files
            st.rerun()

    # Download section
    converted = st.session_state.get(session_key, {})
    if converted:
        st.markdown("---")
        st.subheader(f"Download Converted Files ({len(converted)} files)")

        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
            for name, data in converted.items():
                zipf.writestr(name, data)
        zip_buffer.seek(0)

        st.download_button(
            label=f"Download All ({len(converted)} files) as ZIP",
            data=zip_buffer.getvalue(),
            file_name=zip_name,
            mime="application/zip",
            key=f"dl_{session_key}",
        )


# ──────────────────────────────────────────────
# Tab 1: Convert Assessment
# ──────────────────────────────────────────────

ASSESSMENT_EXTRACTION_PROMPT = """You are an expert at reading assessment documents and extracting structured data.

You will receive the full text content of an assessment document (question paper or answer key).
Extract ALL information into the JSON schema below.

CRITICAL RULES:
- Respond with ONLY a valid JSON object. No preamble or commentary.
- Start your response with { and end with }.
- Do NOT use tools — all data is provided in the prompt.
- Extract EVERY question found in the document. Do not skip any.
- For knowledge_id, look for patterns like (K1), (K2), K1, K2 near the question.
- For ability_id, look for patterns like (A1), (A2), A1, A2 near the question.
- For learning_outcome_id, look for patterns like LO1, LO2 near the question.
- If the document contains answers, extract them. If not, leave answer as empty list.
- Detect the assessment type from the document title or content (e.g., SAQ, PP, CS, OQ).
- Detect whether this is a question paper (no answers) or answer key (has answers).

JSON Schema:
{
    "course_title": "string (the course name from the document title)",
    "assessment_type": "string (e.g., WA (SAQ), PP, CS, OQ)",
    "assessment_code": "string (e.g., WA-SAQ, PP, CS, OQ)",
    "duration": "string (e.g., 60 mins, 1 hr) - extract from document or use empty string",
    "has_answers": true/false,
    "questions": [
        {
            "scenario": "string (the scenario text, or empty string if none)",
            "question_statement": "string (the question text)",
            "knowledge_id": "string (e.g., K1) or null",
            "ability_id": ["string (e.g., A1)"] or null,
            "learning_outcome_id": "string (e.g., LO1) or null",
            "answer": ["string (answer bullet point 1)", "string (answer bullet point 2)"]
        }
    ]
}
"""


async def _extract_assessment_data(doc_text: str) -> dict:
    prompt = f"""Extract the structured assessment data from this document text.

--- DOCUMENT TEXT ---
{doc_text}
--- END ---

Return ONLY the JSON object following the schema in your instructions."""

    return await run_agent_json(
        prompt=prompt,
        system_prompt=ASSESSMENT_EXTRACTION_PROMPT,
        tools=[],
        max_turns=3,
    )


def _build_assessment_result(extracted: dict, filename: str) -> dict:
    course_title = extracted.get('course_title', '')
    assessment_code = extracted.get('assessment_code', extracted.get('assessment_type', 'SAQ'))
    duration = extracted.get('duration', '60 mins')
    questions = extracted.get('questions', [])
    has_answers = extracted.get('has_answers', False)

    if not questions:
        st.warning(f"No questions found in {filename}")
        return {}

    context = {
        'course_title': course_title,
        'duration': duration,
        'assessment_code': assessment_code,
    }

    assessment_type = extracted.get('assessment_type', assessment_code)
    new_doc = _build_assessment_doc(context, assessment_type, questions, include_answers=has_answers)

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp.close()
    new_doc.save(tmp.name)
    with open(tmp.name, 'rb') as f:
        doc_bytes = f.read()
    os.unlink(tmp.name)

    base_name = os.path.splitext(filename)[0]
    out_name = f"{base_name} (WSQ Format).docx"

    doc_type = "Answer Key" if has_answers else "Question Paper"
    st.success(f"{filename} — {len(questions)} questions ({doc_type})")

    return {out_name: doc_bytes}


def _tab_convert_assessment():
    st.markdown("Upload existing assessment documents to batch convert them to the standardised WSQ client-ready format.")

    uploaded_files = st.file_uploader(
        "Upload Assessment Documents (.docx)",
        type=["docx"],
        accept_multiple_files=True,
        key="convert_assessment_upload",
        help="Upload all your assessment DOCX files at once (question papers and answer keys)"
    )

    if not uploaded_files:
        st.info("Upload your assessment DOCX files — you can select multiple files at once. They will all be converted and downloaded as a ZIP.")
        return

    _convert_and_download(
        uploaded_files,
        _extract_assessment_data,
        _build_assessment_result,
        "converted_assessment_files",
        "converted_assessments.zip",
    )


# ──────────────────────────────────────────────
# Tab 2: Convert Courseware (AP, FG, LG, ASR)
# ──────────────────────────────────────────────

COURSEWARE_EXTRACTION_PROMPT = """You are an expert at reading WSQ courseware documents and extracting structured data.

You will receive the full text content of a courseware document (Assessment Plan, Facilitator Guide, Learner Guide, or Assessment Summary Record).
Extract ALL information into the JSON schema below.

CRITICAL RULES:
- Respond with ONLY a valid JSON object. No preamble or commentary.
- Start your response with { and end with }.
- Do NOT use tools — all data is provided in the prompt.
- Detect the document type from content: AP (Assessment Plan), FG (Facilitator Guide), LG (Learner Guide), ASR (Assessment Summary Record).
- Extract ALL learning units, topics, K statements, A statements, learning outcomes.
- Extract company info, course details, training/assessment hours.

JSON Schema:
{
    "doc_type": "string (AP, FG, LG, or ASR)",
    "Course_Title": "string",
    "TGS_Ref_No": "string (e.g., TGS-2026062163)",
    "Name_of_Organisation": "string (company name)",
    "UEN": "string",
    "TSC_Code": "string (e.g., ICT-INT-0047-1.1)",
    "TSC_Title": "string",
    "TSC_Sector": "string",
    "TSC_Category": "string",
    "TSC_Sector_Abbr": "string",
    "Proficiency_Level": "string",
    "Proficiency_Description": "string",
    "Total_Training_Hours": "string (e.g., 30)",
    "Total_Assessment_Hours": "string (e.g., 2)",
    "Total_Course_Duration_Hours": "string (e.g., 32)",
    "Learning_Units": [
        {
            "LU_Title": "string (e.g., LU1. Topic Name)",
            "LO": "string (learning outcome description)",
            "Assessment_Methods": ["string (e.g., WA-SAQ, PP)"],
            "Topics": [
                {
                    "Topic_Title": "string",
                    "Bullet_Points": ["string"]
                }
            ],
            "K_numbering_description": [
                {"K_number": "K1", "Description": "string"}
            ],
            "A_numbering_description": [
                {"A_number": "A1", "Description": "string"}
            ],
            "Instructional_Methods": ["string"]
        }
    ],
    "Assessment_Methods_Details": [
        {
            "Assessment_Method": "string (e.g., Written Assessment - Short Answer Questions)",
            "Method_Abbreviation": "string (e.g., WA-SAQ)",
            "Total_Delivery_Hours": "string",
            "Assessor_to_Candidate_Ratio": ["string"]
        }
    ]
}
"""


async def _extract_courseware_data(doc_text: str) -> dict:
    prompt = f"""Extract the structured courseware data from this document text.

--- DOCUMENT TEXT ---
{doc_text}
--- END ---

Return ONLY the JSON object following the schema in your instructions."""

    return await run_agent_json(
        prompt=prompt,
        system_prompt=COURSEWARE_EXTRACTION_PROMPT,
        tools=[],
        max_turns=3,
    )


def _build_courseware_result(extracted: dict, filename: str) -> dict:
    from docxtpl import DocxTemplate

    doc_type = extracted.get('doc_type', '').upper()
    if not doc_type:
        # Try to detect from filename
        fname_lower = filename.lower()
        if 'ap' in fname_lower or 'assessment plan' in fname_lower:
            doc_type = 'AP'
        elif 'fg' in fname_lower or 'facilitator' in fname_lower:
            doc_type = 'FG'
        elif 'lg' in fname_lower or 'learner' in fname_lower:
            doc_type = 'LG'
        elif 'asr' in fname_lower or 'summary' in fname_lower:
            doc_type = 'ASR'
        else:
            st.warning(f"Could not detect document type for {filename}")
            return {}

    template_map = {
        'AP': 'courseware_agents/courseware/templates/AP_TGS-Ref-No_Course-Title_v1.docx',
        'FG': 'courseware_agents/courseware/templates/FG_TGS-Ref-No_Course-Title_v1.docx',
        'LG': 'courseware_agents/courseware/templates/LG_TGS-Ref-No_Course-Title_v1.docx',
        'ASR': 'courseware_agents/courseware/templates/ASR_TGS-Ref-No_Course-Title_v1.docx',
    }

    template_path = template_map.get(doc_type)
    if not template_path or not os.path.exists(template_path):
        st.error(f"Template not found for {doc_type}: {template_path}")
        return {}

    try:
        # Add Date field
        from datetime import datetime
        extracted['Date'] = datetime.now().strftime('%d %B %Y')

        # Add company logo placeholder
        extracted['company_logo'] = ''

        doc = DocxTemplate(template_path)
        doc.render(extracted)

        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        tmp.close()
        doc.save(tmp.name)
        with open(tmp.name, 'rb') as f:
            doc_bytes = f.read()
        os.unlink(tmp.name)

        base_name = os.path.splitext(filename)[0]
        out_name = f"{base_name} (WSQ Format).docx"

        lu_count = len(extracted.get('Learning_Units', []))
        st.success(f"{filename} — Converted as {doc_type} ({lu_count} learning units)")

        return {out_name: doc_bytes}

    except Exception as e:
        st.error(f"Error building {doc_type} for {filename}: {e}")
        return {}


def _tab_convert_courseware():
    st.markdown("Upload existing courseware documents (AP, FG, LG, ASR) to convert them to the standardised WSQ template format.")

    uploaded_files = st.file_uploader(
        "Upload Courseware Documents (.docx)",
        type=["docx"],
        accept_multiple_files=True,
        key="convert_courseware_upload",
        help="Upload AP, FG, LG, or ASR documents — doc type is auto-detected"
    )

    if not uploaded_files:
        st.info("Upload your courseware DOCX files (AP, FG, LG, ASR). The document type will be auto-detected from the content.")
        return

    _convert_and_download(
        uploaded_files,
        _extract_courseware_data,
        _build_courseware_result,
        "converted_courseware_files",
        "converted_courseware.zip",
    )


# ──────────────────────────────────────────────
# Tab 3: Convert Lesson Plan
# ──────────────────────────────────────────────

LP_EXTRACTION_PROMPT = """You are an expert at reading WSQ Lesson Plan documents and extracting structured data.

You will receive the full text content of a Lesson Plan document.
Extract ALL information into the JSON schema below.

CRITICAL RULES:
- Respond with ONLY a valid JSON object. No preamble or commentary.
- Start your response with { and end with }.
- Do NOT use tools — all data is provided in the prompt.
- Extract the full timetable/schedule with all sessions for each day.
- Extract topics, instructional methods, resources for each session.
- A session is a time block (e.g., 9:00 AM - 10:30 AM) with its content.

JSON Schema:
{
    "Course_Title": "string",
    "TGS_Ref_No": "string (e.g., TGS-2026062163)",
    "Name_of_Organisation": "string (company name)",
    "UEN": "string",
    "Total_Training_Hours": "string",
    "Total_Assessment_Hours": "string",
    "lesson_plan": [
        {
            "Day": "string (e.g., Day 1)",
            "Sessions": [
                {
                    "Time": "string (e.g., 9:00 AM - 10:30 AM)",
                    "instruction_title": "string (topic or activity name)",
                    "bullet_points": ["string (sub-topic or detail)"],
                    "Instructional_Methods": "string (e.g., Lecture, Group Discussion)",
                    "Resources": "string (e.g., Slides, Whiteboard)"
                }
            ]
        }
    ]
}
"""


async def _extract_lp_data(doc_text: str) -> dict:
    prompt = f"""Extract the structured lesson plan data from this document text.

--- DOCUMENT TEXT ---
{doc_text}
--- END ---

Return ONLY the JSON object following the schema in your instructions."""

    return await run_agent_json(
        prompt=prompt,
        system_prompt=LP_EXTRACTION_PROMPT,
        tools=[],
        max_turns=3,
    )


def _build_lp_result(extracted: dict, filename: str) -> dict:
    from docxtpl import DocxTemplate

    template_path = 'courseware_agents/lesson_plan/templates/LP_TGS-Ref-No_Course-Title_v1.docx'
    if not os.path.exists(template_path):
        st.error(f"LP template not found: {template_path}")
        return {}

    try:
        from datetime import datetime
        extracted['Date'] = datetime.now().strftime('%d %B %Y')
        extracted['company_logo'] = ''

        doc = DocxTemplate(template_path)
        doc.render(extracted)

        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        tmp.close()
        doc.save(tmp.name)
        with open(tmp.name, 'rb') as f:
            doc_bytes = f.read()
        os.unlink(tmp.name)

        base_name = os.path.splitext(filename)[0]
        out_name = f"{base_name} (WSQ Format).docx"

        days = len(extracted.get('lesson_plan', []))
        st.success(f"{filename} — Converted LP ({days} day(s))")

        return {out_name: doc_bytes}

    except Exception as e:
        st.error(f"Error building LP for {filename}: {e}")
        return {}


def _tab_convert_lesson_plan():
    st.markdown("Upload existing Lesson Plan documents to convert them to the standardised WSQ template format.")

    uploaded_files = st.file_uploader(
        "Upload Lesson Plan Documents (.docx)",
        type=["docx"],
        accept_multiple_files=True,
        key="convert_lp_upload",
        help="Upload LP documents to convert"
    )

    if not uploaded_files:
        st.info("Upload your Lesson Plan DOCX files. They will be converted to the standardised WSQ LP template.")
        return

    _convert_and_download(
        uploaded_files,
        _extract_lp_data,
        _build_lp_result,
        "converted_lp_files",
        "converted_lesson_plans.zip",
    )


# ──────────────────────────────────────────────
# Main app
# ──────────────────────────────────────────────

def app():
    st.title("Convert Documents")

    tab1, tab2, tab3 = st.tabs([
        "Convert Assessment",
        "Convert Courseware",
        "Convert Lesson Plan",
    ])

    with tab1:
        _tab_convert_assessment()

    with tab2:
        _tab_convert_courseware()

    with tab3:
        _tab_convert_lesson_plan()
