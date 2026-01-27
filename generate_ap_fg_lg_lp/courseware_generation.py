"""
File: courseware_generation.py

===============================================================================
Courseware Document Generator
===============================================================================
Description:
    This module serves as the main entry point for the Courseware Document Generator
    application. It is designed to parse Course Proposal (CP) documents, extract and
    interpret the course data, and generate multiple courseware documents such as:
      - Learning Guide (LG)
      - Assessment Plan (AP)
      - Lesson Plan (LP)
      - Facilitator's Guide (FG)
      - Timetable (as needed)
      
    The application utilizes both AI-based processing (via OpenAI and autogen agents)
    and conventional document parsing and web scraping methods to ensure that the CP data
    is accurately transformed into a structured format for document generation.

Main Functionalities:
    1. Data Models:
        - Defines several Pydantic models (e.g., Topic, LearningUnit, CourseData, etc.)
          to validate and structure the course proposal and generated document data.
          
    2. Document Parsing:
        - Function: parse_cp_document(uploaded_file)
          Parses a CP document (Word or Excel) into a trimmed Markdown string based on
          regex patterns to capture only the relevant sections of the document.
          
    3. Web Scraping:
        - Function: web_scrape(course_title, name_of_org)
          Automates a headless browser session using Selenium to retrieve TGS Ref No (and UEN)
          from the MySkillsFuture portal based on the provided course title and organization.
          
    4. Data Interpretation:
        - Function: interpret_cp(raw_data, model_client)
          Leverages an AI assistant (via the OpenAIChatCompletionClient) to extract and structure
          the course proposal data into a comprehensive JSON dictionary as defined by the CourseData model.
          
    5. Streamlit Application:
        - Function: app()
          Implements the user interface using Streamlit. This interface guides users through:
            - Uploading a Course Proposal document.
            - Managing organization details (CRUD operations via a modal).
            - Optionally uploading an updated Skills Framework dataset.
            - Selecting which courseware documents to generate.
            - Executing the parsing, data extraction, document generation processes,
              and finally providing a ZIP file download of all generated documents.
              
Dependencies:
    - Custom Courseware Utilities:
        • Courseware.utils.agentic_LG         : For generating the Learning Guide.
        • Courseware.utils.agentic_AP         : For generating Assessment Documents.
        • Courseware.utils.timetable_generator : For generating the course timetable.
        • Courseware.utils.agentic_LP         : For generating the Lesson Plan.
        • Courseware.utils.agentic_FG         : For generating the Facilitator's Guide.
        • Courseware.utils.model_configs       : For model configuration and selection.
        • Courseware.utils.organization_utils  : For managing organization data (CRUD).
    - External Libraries:
        • os, io, zipfile, tempfile, json, time, asyncio, datetime
        • streamlit                        : For building the web UI.
        • BeautifulSoup                    : For web scraping tasks.
        • docx                             : For generating and modifying Word documents.
        • pydantic                         : For data validation and structured models.
        • autogen_agentchat & autogen_core   : For AI-assisted text generation and processing.
        • urllib.parse                     : For URL manipulation.
    
Usage:
    - Configure API keys and endpoints in st.secrets (e.g., OPENAI_API_KEY, OPENROUTER_API_KEY,
      BROWSER_TOKEN, BROWSER_WEBDRIVER_ENDPOINT, etc.).
    - Run this module using Streamlit, e.g., `streamlit run <this_file.py>`, to launch the web interface.
    - Follow the on-screen instructions to upload your CP document, manage organization data, select
      the desired courseware documents, and generate/download the outputs.

Author: 
    Derrick Lim
Date:
    4 March 2025

Notes:
    - This module uses asynchronous functions and external AI services for data extraction.
    - The Selenium web scraping component is configured to run headlessly with optimized options
      suitable for both local and containerized environments.
    - Organization management is performed using a JSON-based system via utility functions provided
      in the Courseware.utils.organization_utils module.
    - Ensure all dependencies are installed and properly configured before running the application.

===============================================================================
"""


from generate_ap_fg_lg_lp.utils.agentic_LG import generate_learning_guide
from generate_ap_fg_lg_lp.utils.agentic_AP import generate_assessment_documents
from generate_ap_fg_lg_lp.utils.timetable_generator import generate_timetable
from generate_ap_fg_lg_lp.utils.agentic_LP import generate_lesson_plan
from generate_ap_fg_lg_lp.utils.agentic_FG import generate_facilitators_guide
from settings.model_configs import get_model_config
import os
import io
import zipfile
import tempfile
import json 
import time
import asyncio
from datetime import datetime
import streamlit as st
import urllib.parse
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table
from bs4 import BeautifulSoup
from pydantic import BaseModel
from typing import List, Optional
from openai import OpenAI
from utils.helpers import save_uploaded_file, parse_json_content
# Import organisation CRUD utilities and model
from generate_ap_fg_lg_lp.utils.organization_utils import (
    load_organizations,
    save_organizations,
    add_organization,
    update_organization,
    delete_organization,
    Organization
)
from streamlit_modal import Modal

# Initialize session state variables
if 'lg_output' not in st.session_state:
    st.session_state['lg_output'] = None
if 'ap_output' not in st.session_state:
    st.session_state['ap_output'] = None
if 'lp_output' not in st.session_state:
    st.session_state['lp_output'] = None
if 'fg_output' not in st.session_state:
    st.session_state['fg_output'] = None
if 'context' not in st.session_state:
    st.session_state['context'] = None
if 'asr_output' not in st.session_state:
    st.session_state['asr_output'] = None
# Note: selected_model is set in app.py sidebar based on user selection and database defaults
# Do not set a hardcoded default here - let app.py handle model selection

############################################################
# 1. Pydantic Models
############################################################
class Topic(BaseModel):
    Topic_Title: str
    Bullet_Points: List[str]

class KDescription(BaseModel):
    K_number: str
    Description: str

class ADescription(BaseModel):
    A_number: str
    Description: str

class LearningUnit(BaseModel):
    LU_Title: str
    Topics: List[Topic]
    LO: str
    K_numbering_description: List[KDescription]
    A_numbering_description: List[ADescription]
    Assessment_Methods: List[str]
    Instructional_Methods: List[str]

class EvidenceDetail(BaseModel):
    LO: str
    Evidence: str

class AssessmentMethodDetail(BaseModel):
    Assessment_Method: str
    Method_Abbreviation: str
    Total_Delivery_Hours: str
    Assessor_to_Candidate_Ratio: List[str]
    Evidence: Optional[List[EvidenceDetail]] = None
    Submission: Optional[List[str]] = None
    Marking_Process: Optional[List[str]] = None
    Retention_Period: Optional[str] = None

class CourseData(BaseModel):
    Date: str 
    Year: str
    Name_of_Organisation: str
    Course_Title: str
    TSC_Title: str
    TSC_Code: str
    Total_Training_Hours: str 
    Total_Assessment_Hours: str 
    Total_Course_Duration_Hours: str 
    Learning_Units: List[LearningUnit]
    Assessment_Methods_Details: List[AssessmentMethodDetail]

class Session(BaseModel):
    Time: str
    instruction_title: str
    bullet_points: List[str]
    Instructional_Methods: str
    Resources: str

class DayLessonPlan(BaseModel):
    Day: str
    Sessions: List[Session]

class LessonPlan(BaseModel):
    lesson_plan: List[DayLessonPlan]

############################################################
# 2. Course Proposal Document Parsing
############################################################
from docx import Document as DocxDocument
import openpyxl
import os
import re

def parse_cp_document(uploaded_file):
    """
    Parses a Course Proposal (CP) document (UploadedFile) and returns its content as text,
    trimmed based on the document type using regex patterns.

    For Word CP (.docx):
      - Excludes everything before a line matching "Part 1" and "Particulars of Course"
      - Excludes everything after a line matching "Part 4" and "Facilities and Resources"

    For Excel CP (.xlsx):
      - Excludes everything before a line matching "1 - Course Particulars"
      - Excludes everything after a line matching "4 - Declarations"

    Args:
        uploaded_file (UploadedFile): The file uploaded via st.file_uploader.

    Returns:
        str: A trimmed text string containing the parsed document content.
    """
    # Write the uploaded file to a temporary file.
    with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[1]) as tmp:
        tmp.write(uploaded_file.read())
        temp_file_path = tmp.name

    try:
        ext = os.path.splitext(temp_file_path)[1].lower()
        text_content = []

        if ext == ".docx":
            # Parse Word document using python-docx
            doc = DocxDocument(temp_file_path)
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))

        elif ext == ".xlsx":
            # Parse Excel document using openpyxl
            wb = openpyxl.load_workbook(temp_file_path, data_only=True)
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                text_content.append(f"## {sheet}")
                for row in ws.iter_rows(values_only=True):
                    row_text = [str(cell) if cell is not None else "" for cell in row]
                    if any(row_text):
                        text_content.append(" | ".join(row_text))
            wb.close()

        markdown_text = "\n".join(text_content)

        # Set up regex patterns based on file extension
        if ext == ".docx":
            start_pattern = re.compile(r"Part\s*1.*?Particulars\s+of\s+Course", re.IGNORECASE)
            end_pattern = re.compile(r"Part\s*4.*?Facilities\s+and\s+Resources", re.IGNORECASE)
        elif ext == ".xlsx":
            start_pattern = re.compile(r"1\s*-\s*Course\s*Particulars", re.IGNORECASE)
            end_pattern = re.compile(r"4\s*-\s*Declarations", re.IGNORECASE)
        else:
            start_pattern = None
            end_pattern = None

        # If both patterns exist, search for the matches and trim the text
        if start_pattern and end_pattern:
            start_match = start_pattern.search(markdown_text)
            end_match = end_pattern.search(markdown_text)
            if start_match and end_match and end_match.start() > start_match.start():
                markdown_text = markdown_text[start_match.start():end_match.start()].strip()

    finally:
        # Clean up the temporary file
        os.remove(temp_file_path)

    return markdown_text

############################################################
# 2. Interpret Course Proposal Data
############################################################
def create_openai_client(model_choice: str = "GPT-4o-Mini"):
    """
    Create an OpenAI client configured with the specified model choice.

    Args:
        model_choice: Model choice string (e.g., "DeepSeek-Chat", "GPT-4o-Mini")

    Returns:
        tuple: (OpenAI client instance, model configuration dict)
    """
    autogen_config = get_model_config(model_choice)
    config_dict = autogen_config.get("config", {})

    base_url = config_dict.get("base_url", "https://openrouter.ai/api/v1")
    api_key = config_dict.get("api_key", "")
    model = config_dict.get("model", "gpt-4o-mini")
    temperature = config_dict.get("temperature", 0.2)

    # Fallback: If no API key in config, get it dynamically based on api_provider
    if not api_key:
        from settings.api_manager import load_api_keys
        api_provider = autogen_config.get("api_provider", "OPENROUTER")
        api_keys = load_api_keys()
        api_key = api_keys.get(f"{api_provider}_API_KEY", "")

    client = OpenAI(
        base_url=base_url,
        api_key=api_key
    )

    model_config = {
        "model": model,
        "temperature": temperature,
        "base_url": base_url
    }

    return client, model_config


async def interpret_cp(raw_data: dict, model_choice: str = "GPT-4o-Mini") -> dict:
    """
    Interprets and extracts structured data from a raw Course Proposal (CP) document using OpenAI SDK.

    This function processes raw CP data using an AI model to extract
    structured information such as course details, learning units, topics,
    assessment methods, and instructional methods.

    Args:
        raw_data (dict):
            The unstructured data extracted from the CP document.
        model_choice (str):
            The model choice string for selecting the AI model.

    Returns:
        dict:
            A structured dictionary containing course details.

    Raises:
        Exception:
            If the AI-generated response does not contain the expected fields.
    """
    client, config = create_openai_client(model_choice)

    system_message = f"""
        You are an AI assistant that helps extract specific information from a JSON object containing a Course Proposal Form (CP). Your task is to interpret the JSON data, regardless of its structure, and extract the required information accurately.

        ---

        **Task:** Extract the following information from the provided JSON data:

        ### Part 1: Particulars of Course

        - Name of Organisation
        - Course Title
        - TSC Title
        - TSC Code
        - Total Training Hours/ Total Instructional Duration (calculated as the sum of Classroom Facilitation, Workplace Learning: On-the-Job (OJT), Practicum, Practical, E-learning: Synchronous and Asynchronous), formatted with units (e.g., "30 hrs", "1 hr")
        - Total Assessment Hours/ Total Assessment Duration, formatted with units (e.g., "2 hrs")
        - Total Course Duration Hours, formatted with units (e.g., "42 hrs")

        ### Part 3: Curriculum Design

        From the Learning Units and Topics Table:

        For each Learning Unit (LU):
        - Learning Unit Title (include the "LUx: " prefix)
        - Topics Covered Under Each LU:
        - For each Topic:
            - **Topic_Title** (include the "Topic x: " prefix and the associated K and A statements in parentheses)
            - **Bullet_Points** (a list of bullet points under the topic; remove any leading bullet symbols such as "-" so that only the content remains)
        - Learning Outcomes (LOs) (include the "LOx: " prefix for each LO)
        - Numbering and Description for the "K" (Knowledge) Statements (as a list of dictionaries with keys "K_number" and "Description")
        - Numbering and Description for the "A" (Ability) Statements (as a list of dictionaries with keys "A_number" and "Description")
        - **Assessment_Methods** (a list of assessment method abbreviations; e.g., ["WA-SAQ", "CS"]). Note: If the CP contains the term "Written Exam", output it as "Written Assessment - Short Answer Questions". If it contains "Practical Exam", output it as "Practical Performance".
        - **Duration Calculation:** When extracting the duration for each assessment method:
            1. If the extracted duration is not exactly 0.5 or a whole number (e.g., 0.5, 1, 2, etc.), interpret it as minutes.
            2. If duplicate entries for the same assessment method occur within the same LU, sum their durations to obtain a total duration.
            3. For CPs in Excel format, under 3 - Summary sheet, the duration appears in the format "(Assessor-to-Candidate Ratio, duration)"—for example, "Written Exam (1:20, 20)" means 20 minutes, and "Others: Case Study (1:20, 25)" appearing twice should result in a total of 50 minutes for Case Study.
        - **Instructional_Methods** (a list of instructional method abbreviations or names)

        ### Part E: Details of Assessment Methods Proposed

        For each Assessment Method in the CP, extract:
        - **Assessment_Method** (always use the full term, e.g., "Written Assessment - Short Answer Questions", "Practical Performance", "Case Study", "Oral Questioning", "Role Play")
        - **Method_Abbreviation** (if provided in parentheses or generated according to the rules)
        - **Total_Delivery_Hours** (formatted with units, e.g., "1 hr")
        - **Assessor_to_Candidate_Ratio** (a list of minimum and maximum ratios, e.g., ["1:3 (Min)", "1:5 (Max)"])

        **Additionally, if the CP explicitly provides the following fields, extract them. Otherwise, do not include them in the final output:**
        - **Type_of_Evidence**
        - For PP and CS assessment methods, the evidence may be provided as a dictionary where keys are LO identifiers (e.g., "LO1", "LO2", "LO3") and values are the corresponding evidence text. In that case, convert the dictionary into a list of dictionaries with keys `"LO"` and `"Evidence"`.
        - If the evidence is already provided as a list (for example, a list of strings or a list of dictionaries), keep it as is.
        - **Manner_of_Submission** (as a list, e.g., ["Submission 1", "Submission 2"])
        - **Marking_Process** (as a list, e.g., ["Process 1", "Process 2"])
        - **Retention_Period**: **Extract the complete retention description exactly as provided in the CP.**
        - **No_of_Role_Play_Scripts** (only if the assessment method is Role Play and this information is provided)

        ---

        **Instructions:**

        - Carefully parse the JSON data and locate the sections corresponding to each part.
        - Even if the JSON structure changes, use your understanding to find and extract the required information.
        - Ensure that the `Topic_Title` includes the "Topic x: " prefix and the associated K and A statements in parentheses exactly as they appear.
        - For Learning Outcomes (LOs), always include the "LOx: " prefix (where x is the number).
        - Present the extracted information in a structured JSON format where keys correspond exactly to the placeholders required for the Word document template.
        - Ensure all extracted information is normalized by:
            - Replacing en dashes (–) and em dashes (—) with hyphens (-)
            - Converting curly quotes (" ") to straight quotes (")
            - Replacing other non-ASCII characters with their closest ASCII equivalents.
        - **Time fields** must include units (e.g., "40 hrs", "1 hr", "2 hrs").
        - For `Assessment_Methods`, always use the abbreviations (e.g., WA-SAQ, PP, CS, OQ, RP) as per the following rules:
            1. Use the abbreviation provided in parentheses if available.
            2. Otherwise, generate an abbreviation by taking the first letters of the main words (ignoring articles/prepositions) and join with hyphens.
            3. For methods containing "Written Assessment", always prefix with "WA-".
            4. If duplicate or multiple variations exist, use the standard abbreviation.
        - **Important:** Verify that the sum of `Total_Delivery_Hours` for all assessment methods equals the `Total_Assessment_Hours`. If individual delivery hours for assessment methods are not specified, divide the `Total_Assessment_Hours` equally among them.
        - For bullet points in each topic, ensure that the number of bullet points exactly matches those in the CP. Re-extract if discrepancies occur.
        - **If the same K or A statement (same numbering and description) appears multiple times within the same LU, keep only one instance. If the same K or A statement appears in different LUs, keep it as it is.**
        - Do not include any extraneous information or duplicate entries.

        Generate structured output matching this schema:
        {json.dumps(CourseData.model_json_schema(), indent=2)}
        """

    agent_task = f"""
    Please extract and structure the following data: {raw_data}.
    **Return the extracted information as a complete JSON dictionary containing the specified fields. Do not truncate or omit any data. Include all fields and their full content. Do not use '...' or any placeholders to replace data.**
    Simply return the JSON dictionary object directly.
    """

    try:
        completion = client.chat.completions.create(
            model=config["model"],
            temperature=config["temperature"],
            messages=[
                {"role": "system", "content": system_message},
                {"role": "user", "content": agent_task}
            ],
            response_format={"type": "json_object"},
            max_tokens=16384
        )

        raw_content = completion.choices[0].message.content

        if not raw_content:
            print("ERROR: No response from LLM during CP interpretation")
            return "No content found in the agent's last message."

        # Debug: Log response length and preview
        print(f"CP Interpretation - Response length: {len(raw_content)} chars")
        print(f"CP Interpretation - First 500 chars: {raw_content[:500]}")
        print(f"CP Interpretation - Last 500 chars: {raw_content[-500:]}")

        context = parse_json_content(raw_content)
        if context is None:
            print("ERROR: parse_json_content returned None - invalid JSON")
            print(f"Raw response (truncated): {raw_content[:1000]}...")
            raise Exception(f"Failed to parse JSON from model response. Raw content: {raw_content[:500]}...")

        # Debug: Check if K and A statements were extracted
        if "Learning_Units" in context:
            for lu in context["Learning_Units"]:
                k_count = len(lu.get("K_numbering_description", []))
                a_count = len(lu.get("A_numbering_description", []))
                print(f"LU: {lu.get('LU_Title', 'Unknown')} - K statements: {k_count}, A statements: {a_count}")

        return context

    except Exception as e:
        print(f"ERROR: Exception during CP interpretation: {e}")
        raise Exception(f"Error parsing structured output: {e}")

# Streamlit App
def app():
    """
    Streamlit web application for generating courseware documents.

    This function serves as the entry point for the user interface,
    allowing users to upload a Course Proposal document, select 
    their organization, and generate various courseware documents.

    The app guides users through:
    - Uploading a Course Proposal (CP) document.
    - Selecting an organization from a predefined list.
    - Uploading an optional updated Skills Framework (SFw) dataset.
    - Selecting documents to generate (Learning Guide, Lesson Plan, etc.).
    - Processing and downloading the generated documents.

    Raises:
        ValueError: 
            If required input fields are missing.
        Exception: 
            If any step in the document generation process fails.
    """

    st.title("Generate AP/FG/LG/LP")

    # ================================================================
    # Step 1: Upload Course Proposal (CP) Document
    # ================================================================
    st.subheader("Step 1: Upload Course Proposal (CP) Document")
    cp_file = st.file_uploader("Upload Course Proposal (CP) Document", type=["docx", "xlsx"])

    # ================================================================
    # Step 2: Select Name of Organisation
    # ================================================================
    # Create a modal instance with a unique key and title
    crud_modal = Modal(key="crud_modal", title="Manage Organisations")

    st.subheader("Step 2: Enter Relevant Details")
    tgs_course_code = st.text_input("Enter TGS Course Code", key="tgs_course_code", placeholder="e.g., TGS-2023039181")

    # Load organisations from JSON using the utility function - cached
    @st.cache_data
    def get_cached_organizations():
        return load_organizations()

    org_list = get_cached_organizations()
    org_names = [org["name"] for org in org_list] if org_list else []

    # Get the company selected from sidebar (automatically use it)
    sidebar_selected_company = st.session_state.get('selected_company', None)
    if sidebar_selected_company:
        selected_org = sidebar_selected_company['name']
    else:
        selected_org = org_names[0] if org_names else None

    # Hidden manage button (keep modal functionality for future use)
    # with col2:
    #     if st.button("Manage", key="manage_button", use_container_width=True):
    #         crud_modal.open()

    # ---------------------------
    # Modal: CRUD Interface
    # ---------------------------
    if crud_modal.is_open():
        with crud_modal.container():
            
            # ---- Add New Organisation Form ----
            st.write("#### Add New Organisation")
            with st.form("new_org_form"):
                new_name = st.text_input("Organisation Name", key="new_org_name")
                new_uen = st.text_input("UEN", key="new_org_uen")
                # Use file uploader for the logo instead of a text input
                new_logo_file = st.file_uploader("Upload Logo (optional)", type=["png", "jpg", "jpeg"], key="new_org_logo_file")
                new_submitted = st.form_submit_button("Add Organisation")
                if new_submitted:
                    logo_path = None
                    if new_logo_file is not None:
                        # Construct a safe filename based on the organisation name and file extension
                        _, ext = os.path.splitext(new_logo_file.name)
                        safe_filename = new_name.lower().replace(" ", "_") + ext
                        save_path = os.path.join("Courseware", "utils", "logo", safe_filename)
                        with open(save_path, "wb") as f:
                            f.write(new_logo_file.getvalue())
                        logo_path = save_path
                    new_org = Organization(name=new_name, uen=new_uen, logo=logo_path)
                    add_organization(new_org)
                    st.success(f"Organisation '{new_name}' added.")
                    st.rerun()
            
            # ---- Display Existing Organisations with Edit/Delete Buttons ----
            st.write("#### Existing Organisations")
            org_list = load_organizations()  # Refresh the list

            # Table header
            col_sno, col_name, col_uen, col_logo, col_edit, col_delete = st.columns([1, 3, 2, 2, 1, 2])
            col_sno.markdown("**SNo**")
            col_name.markdown("**Name**")
            col_uen.markdown("**UEN**")
            col_logo.markdown("**Logo**")
            col_edit.markdown("**Edit**")
            col_delete.markdown("**Delete**")

            # Table rows
            for display_idx, org in enumerate(org_list, start=1):
                # The actual index in the list is display_idx - 1
                real_index = display_idx - 1

                row_sno, row_name, row_uen, row_logo, row_edit, row_delete = st.columns([1, 3, 2, 2, 1, 2])
                row_sno.write(display_idx)
                row_name.write(org["name"])
                row_uen.write(org["uen"])
                
                if org["logo"] and os.path.exists(org["logo"]):
                    row_logo.image(org["logo"], width=70)
                else:
                    row_logo.write("No Logo")

                # Edit/Delete Buttons
                if row_edit.button("Edit", key=f"edit_{display_idx}", type="secondary"):
                    st.session_state["org_edit_index"] = real_index
                    st.rerun()
                if row_delete.button("Delete", key=f"delete_{display_idx}", type="primary"):
                    if org["logo"] and os.path.exists(org["logo"]):
                        os.remove(org["logo"])
                    delete_organization(real_index)
                    st.success(f"Organisation '{org['name']}' deleted.")
                    st.rerun()

            # ---- Edit Organisation Form (if a row is selected for editing) ----
            if "org_edit_index" in st.session_state:
                edit_index = st.session_state["org_edit_index"]
                org_to_edit = load_organizations()[edit_index]
                st.write(f"#### Edit Organisation: {org_to_edit['name']}")
                with st.form("edit_org_form"):
                    edited_name = st.text_input("Organisation Name", value=org_to_edit["name"], key="edited_name")
                    edited_uen = st.text_input("UEN", value=org_to_edit["uen"], key="edited_uen")
                    # File uploader for updating the logo image
                    edited_logo_file = st.file_uploader("Upload Logo (optional)", type=["png", "jpg", "jpeg"], key="edited_logo_file")
                    edit_submitted = st.form_submit_button("Update Organisation")
                    if edit_submitted:
                        logo_path = org_to_edit.get("logo", None)
                        if edited_logo_file is not None:
                            _, ext = os.path.splitext(edited_logo_file.name)
                            safe_filename = edited_name.lower().replace(" ", "_") + ext
                            save_path = os.path.join("Courseware", "utils", "logo", safe_filename)
                            with open(save_path, "wb") as f:
                                f.write(edited_logo_file.getvalue())
                            logo_path = save_path
                        updated_org = Organization(name=edited_name, uen=edited_uen, logo=logo_path)
                        update_organization(edit_index, updated_org)
                        st.success(f"Organisation '{edited_name}' updated.")
                        del st.session_state["org_edit_index"]
                        st.rerun()

    # ================================================================
    # Step 3 (Optional): Upload Updated SFW Dataset
    # ================================================================
    st.subheader("Step 3 (Optional): Upload Updated Skills Framework (SFw) Dataset")
    sfw_file = st.file_uploader("Upload Updated SFw Dataset (Excel File)", type=["xlsx"])
    if sfw_file:
        sfw_data_dir = save_uploaded_file(sfw_file, 'input/dataset')
        st.success(f"Updated SFw dataset saved to {sfw_data_dir}")
    else:
        sfw_data_dir = "generate_ap_fg_lg_lp/input/dataset/Sfw_dataset-2022-03-30 copy.xlsx"

    # ================================================================
    # Step 4: Select Document(s) to Generate using Checkboxes
    # ================================================================
    st.subheader("Step 4: Select Document(s) to Generate")
    generate_lg = st.checkbox("Learning Guide (LG)", value=True)
    generate_ap = st.checkbox("Assessment Plan (AP)", value=True)
    generate_lp = st.checkbox("Lesson Plan (LP)", value=True)
    generate_fg = st.checkbox("Facilitator's Guide (FG)", value=True)

    # ================================================================
    # Step 5: Generate Documents
    # ================================================================
    if st.button("Generate Documents"):
        if cp_file is not None and selected_org:
            # Reset previous output document paths
            st.session_state['lg_output'] = None
            st.session_state['ap_output'] = None
            st.session_state['asr_output'] = None
            st.session_state['lp_output'] = None
            st.session_state['fg_output'] = None
            # Use the selected model choice for all OpenAI SDK calls
            if 'selected_model' not in st.session_state or not st.session_state['selected_model']:
                st.error("❌ No model selected. Please select a model from the sidebar.")
                return

            model_choice = st.session_state['selected_model']
            selected_config = get_model_config(model_choice)

            # Get API key from config or load dynamically based on api_provider
            api_key = selected_config["config"].get("api_key")

            # Fallback: If no API key in config, get it dynamically based on api_provider
            if not api_key:
                from settings.api_manager import load_api_keys
                api_provider = selected_config.get("api_provider", "OPENROUTER")
                api_keys = load_api_keys()
                api_key = api_keys.get(f"{api_provider}_API_KEY", "")

            if not api_key:
                api_provider = selected_config.get("api_provider", "OPENROUTER")
                st.error(f"❌ API key for {model_choice} ({api_provider}) is not provided.")
                st.info(f"💡 **Solution**: Go to Settings → API & Models to add {api_provider}_API_KEY")
                return
            model_name = selected_config["config"]["model"]
            base_url = selected_config["config"].get("base_url", None)

            # Step 1: Parse the CP document
            try:
                with st.spinner('Parsing the Course Proposal...'):
                    raw_data = parse_cp_document(cp_file)
            except Exception as e:
                st.error(f"Error parsing the Course Proposal: {e}")
                return
            
            try:
                with st.spinner('Extracting Information from Course Proposal...'):
                    context = asyncio.run(interpret_cp(raw_data=raw_data, model_choice=model_choice))

            except Exception as e:
                st.error(f"Error extracting Course Proposal: {e}")
                return

            # After obtaining the context
            if context:
                # Step 2: Add the current date to the raw_data
                current_datetime = datetime.now()
                current_date = current_datetime.strftime("%d %b %Y")
                year = current_datetime.year
                context["Date"] = current_date
                context["Year"] = year
                # Find the selected organisation UEN in the organisation's record
                selected_org_data = next((org for org in org_list if org["name"] == selected_org), None)
                if selected_org_data:
                    context["UEN"] = selected_org_data["uen"]

                tgs_course_code = st.session_state.get("tgs_course_code", "")
                context["TGS_Ref_No"] = tgs_course_code

                st.session_state['context'] = context  # Store context in session state

                # Generate Learning Guide
                if generate_lg:
                    try:
                        with st.spinner('Generating Learning Guide...'):
                            lg_output = generate_learning_guide(context, selected_org, model_choice)
                        if lg_output:
                            st.success(f"Learning Guide generated: {lg_output}")
                            st.session_state['lg_output'] = lg_output  # Store output path in session state
                    except Exception as e:
                        st.error(f"Error generating Learning Guide: {e}")

                # Generate Assessment Plan
                if generate_ap:
                    try:
                        with st.spinner('Generating Assessment Plan and Assessment Summary Record...'):
                            ap_output, asr_output = generate_assessment_documents(context, selected_org, None, model_name, api_key, base_url)
                        
                        if ap_output:
                            st.success(f"Assessment Plan generated: {ap_output}")
                            st.session_state['ap_output'] = ap_output  # Store output path in session state

                        if asr_output:
                            st.success(f"Assessment Summary Record generated: {asr_output}")
                            st.session_state['asr_output'] = asr_output  # Store output path in session state

                    except Exception as e:
                        st.error(f"Error generating Assessment Documents: {e}")

                # Check if any documents require the timetable
                needs_timetable = (generate_lp or generate_fg)

                # Generate the timetable if needed and not already generated
                if needs_timetable and 'lesson_plan' not in context:
                    try:
                        with st.spinner("Generating Timetable..."):
                            hours = int(''.join(filter(str.isdigit, context["Total_Course_Duration_Hours"])))
                            num_of_days = hours / 8
                            timetable_data = asyncio.run(generate_timetable(context, num_of_days, model_choice))
                            context['lesson_plan'] = timetable_data['lesson_plan']
                        st.session_state['context'] = context  # Update context in session state
                    except Exception as e:
                        st.error(f"Error generating timetable: {e}")
                        return  # Exit if timetable generation fails
                    
                # Now generate Lesson Plan
                if generate_lp:
                    try:
                        with st.spinner("Generating Lesson Plan..."):
                            lp_output = generate_lesson_plan(context, selected_org)
                        if lp_output:
                            st.success(f"Lesson Plan generated: {lp_output}")
                            st.session_state['lp_output'] = lp_output  # Store output path in session state
     
                    except Exception as e:
                        st.error(f"Error generating Lesson Plan: {e}")

                # Generate Facilitator's Guide
                if generate_fg:
                    try:
                        with st.spinner("Generating Facilitator's Guide..."):
                            fg_output = generate_facilitators_guide(context, selected_org)
                        if fg_output:
                            st.success(f"Facilitator's Guide generated: {fg_output}")
                            st.session_state['fg_output'] = fg_output  # Store output path in session state

                    except Exception as e:
                        st.error(f"Error generating Facilitator's Guide: {e}")
            else:
                st.error("Context is empty. Cannot proceed with document generation.")
        else:
            st.error("Please upload a CP document and select a Name of Organisation.")

    # Check if any courseware document was generated
    if any([
        st.session_state.get('lg_output'),
        st.session_state.get('ap_output'),
        st.session_state.get('asr_output'),
        st.session_state.get('lp_output'),
        st.session_state.get('fg_output')
    ]):
        st.subheader("Download All Generated Documents as ZIP")

        # Create an in-memory ZIP file
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
            
            # Helper function to add a file to the zip archive
            def add_file(file_path, prefix):
                if file_path and os.path.exists(file_path):
                    # Determine file name based on TGS_Ref_No (if available) or fallback to course title
                    if 'TGS_Ref_No' in st.session_state['context'] and st.session_state['context']['TGS_Ref_No']:
                        file_name = f"{prefix}_{st.session_state['context']['TGS_Ref_No']}_{st.session_state['context']['Course_Title']}_v1.docx"
                    else:
                        file_name = f"{prefix}_{st.session_state['context']['Course_Title']}_v1.docx"
                    zipf.write(file_path, arcname=file_name)
            
            # Add each generated document if it exists
            add_file(st.session_state.get('lg_output'), "LG")
            add_file(st.session_state.get('ap_output'), "Assessment_Plan")
            add_file(st.session_state.get('asr_output'), "Assessment_Summary_Record")
            add_file(st.session_state.get('lp_output'), "LP")
            add_file(st.session_state.get('fg_output'), "FG")
        
        # Reset the buffer's position to the beginning
        zip_buffer.seek(0)
        
        # Create a download button for the ZIP archive
        st.download_button(
            label="Download All Documents (ZIP)",
            data=zip_buffer.getvalue(),
            file_name="courseware_documents.zip",
            mime="application/zip"
        )
