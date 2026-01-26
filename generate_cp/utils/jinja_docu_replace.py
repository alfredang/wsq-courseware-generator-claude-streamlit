import json
import sys
import re
from docx import Document
from docxtpl import DocxTemplate

def replace_placeholders_with_docxtpl(json_path, doc_path, new_doc_name):
    # Load the JSON data
    with open(json_path, 'r') as file:
        json_data = json.load(file)

        # Debug print to see the structure
    print("DEBUG: JSON data structure keys:", json_data.keys())
    
    # Check for None values in the data that might be iterated
    for key, value in json_data.items():
        if value is None:
            print(f"WARNING: Key '{key}' has None value")
        elif isinstance(value, dict):
            for subkey, subvalue in value.items():
                if subvalue is None:
                    print(f"WARNING: Nested key '{key}.{subkey}' has None value")


    # Preprocess JSON keys to make them valid Python variable names
    def preprocess_json_keys(json_data):
        new_data = {}
        for key, value in json_data.items():
            # Remove special characters and adjust keys
            new_key = re.sub(r'[^0-9a-zA-Z_]', '_', key)
            new_key = new_key.strip('_')
            # Recursively preprocess if value is a dict
            if isinstance(value, dict):
                value = preprocess_json_keys(value)
            new_data[new_key] = value
        return new_data

    context = preprocess_json_keys(json_data)

    # List of placeholders to process
    placeholders_to_process = ['Placeholder_1'] + [f'Topics_{i}' for i in range(6)] + ['AssessmentJustification','Sequencing']  # Adjust range as needed

    # Process specified placeholders
    for placeholder in placeholders_to_process:
        if placeholder in context:
            value = context[placeholder]
            if isinstance(value, list) and len(value) > 0:
                context[placeholder] = process_placeholder(value)
            elif isinstance(value, str):
                context[placeholder] = process_placeholder(value)
    # CHECK THE CONTEXT BEFORE RENDERING
    print("Context being passed to the template:")
    print(json.dumps(context, indent=4))  # Prints the context in a readable format to the console

    # Alternatively, write the context to a file for inspection
    with open('generate_cp/json_output/context_output.json', 'w') as outfile:
        json.dump(context, outfile, indent=4)
    print("Context written to 'context_output.json'")

    # Load the template document
    tpl = DocxTemplate(doc_path)

    # # Render the document with the context
    # tpl.render(context, autoescape=True)

    data = json_data
    
    try:
        print("DEBUG: About to render template...")
        tpl.render(context, autoescape=True)
        print("DEBUG: Template rendered successfully")
    except Exception as e:
        print(f"ERROR during template rendering: {str(e)}")
        print(f"Error type: {type(e)}")
        # Try to identify which variable caused the issue
        if "is not iterable" in str(e):
            # Extract the problematic template section if possible
            print("Possible template issue with a None value being iterated")
            # Print relevant parts of the context that might be None but expected to be lists
            for key in ['Learning Outcomes', 'TSC and Topics', 'Assessment Methods']:
                if key in data:
                    print(f"DEBUG: '{key}' structure: {type(data[key])}")

    # Save the new document after rendering placeholders
    temp_doc_path = "generate_cp/output_docs/temp_with_placeholders_replaced.docx"
    tpl.save(temp_doc_path)

    # Load the modified document to clean tables
    doc = Document(temp_doc_path)

    # Clean the tables in the document
    clean_tables(doc)

    # Save the final cleaned document
    doc.save(new_doc_name)
    print(f"Updated document saved as: {new_doc_name}")

def clean_tables(doc):
    """Cleans the tables in the document by removing rows with empty first cells."""
    # Access all the tables in the document
    tables = [doc.tables[i] for i in range(len(doc.tables))]

    # Iterate over each table and clean it
    for table in tables:
        print_and_modify_table(table)

def print_and_modify_table(table, indent_level=0):
    indent = "\t" * indent_level

    # Create a list of rows to remove to avoid modifying the list while iterating
    rows_to_remove = []

    for row in table.rows:
        first_cell_text = row.cells[0].text.strip()
        if first_cell_text == "":
            # Mark the row for removal if the first cell is empty
            rows_to_remove.append(row)
        else:
            # Print row content if not empty
            row_content = []
            for cell in row.cells:
                if cell.tables:
                    # If there are nested tables, handle them recursively
                    for nested_table in cell.tables:
                        print(f"{indent}Nested Table Content:")
                        print_and_modify_table(nested_table, indent_level + 1)
                else:
                    row_content.append(cell.text)
            if row_content:
                print(f"{indent}" + "\t".join(row_content))

    # Remove rows that have an empty first cell
    for row in rows_to_remove:
        table._tbl.remove(row._tr)

def process_placeholder(value):
    import re
    items = []
    
    # Define the phrases to be bolded
    bold_phrases = ["Performance Gaps:", "Attributes Gained:", "Post-Training Benefits to Learners:"]
    
    # If value is a list, process each entry in the list
    if isinstance(value, list):
        for idx, entry in enumerate(value):
            lines = entry.split('\n')
            current_paragraph = []
            bullet_points = []
            
            for line in lines:
                line = line.strip()
                
                if not line:
                    # Empty line encountered; finalize current paragraph or bullets
                    if current_paragraph:
                        items.append({'type': 'paragraph', 'content': ' '.join(current_paragraph)})
                        current_paragraph = []
                    if bullet_points:
                        items.append({'type': 'bullets', 'content': bullet_points})
                        bullet_points = []
                    # Do not add empty paragraph for spacing
                    
                elif line.startswith('•'):
                    # Bullet point
                    if current_paragraph:
                        items.append({'type': 'paragraph', 'content': ' '.join(current_paragraph)})
                        current_paragraph = []
                    bullet_points.append(line.lstrip('•').strip())
                    
                elif re.match(r'^LU\d+:\s', line) or line in bold_phrases:
                    # LU title or specific bold phrases
                    if current_paragraph:
                        items.append({'type': 'paragraph', 'content': ' '.join(current_paragraph)})
                        current_paragraph = []
                    if bullet_points:
                        items.append({'type': 'bullets', 'content': bullet_points})
                        bullet_points = []
                    # Add the line as a bold paragraph
                    items.append({'type': 'bold_paragraph', 'content': line})
                    
                else:
                    # Regular line
                    if bullet_points:
                        items.append({'type': 'bullets', 'content': bullet_points})
                        bullet_points = []
                    current_paragraph.append(line)
                    
            # Handle any remaining content in the entry
            if current_paragraph:
                items.append({'type': 'paragraph', 'content': ' '.join(current_paragraph)})
                current_paragraph = []
            if bullet_points:
                items.append({'type': 'bullets', 'content': bullet_points})
                bullet_points = []
                
    else:
        # Handle single string value (e.g., Conclusion)
        lines = value.split('\n')
        current_paragraph = []
        bullet_points = []
        
        for line in lines:
            line = line.strip()
            
            if not line:
                # Empty line encountered; finalize current paragraph or bullets
                if current_paragraph:
                    items.append({'type': 'paragraph', 'content': ' '.join(current_paragraph)})
                    current_paragraph = []
                if bullet_points:
                    items.append({'type': 'bullets', 'content': bullet_points})
                    bullet_points = []
                # Do not add empty paragraph for spacing
                
            elif line.startswith('•'):
                if current_paragraph:
                    items.append({'type': 'paragraph', 'content': ' '.join(current_paragraph)})
                    current_paragraph = []
                bullet_points.append(line.lstrip('•').strip())
                
            elif line in bold_phrases:
                # Specific bold phrases
                if current_paragraph:
                    items.append({'type': 'paragraph', 'content': ' '.join(current_paragraph)})
                    current_paragraph = []
                if bullet_points:
                    items.append({'type': 'bullets', 'content': bullet_points})
                    bullet_points = []
                # Add the line as a bold paragraph
                items.append({'type': 'bold_paragraph', 'content': line})
                
            else:
                if bullet_points:
                    items.append({'type': 'bullets', 'content': bullet_points})
                    bullet_points = []
                current_paragraph.append(line)
                
        if current_paragraph:
            items.append({'type': 'paragraph', 'content': ' '.join(current_paragraph)})
        if bullet_points:
            items.append({'type': 'bullets', 'content': bullet_points})
    
    return items

# Example of how to use this function
if __name__ == "__main__":
    # Ensure correct number of arguments
    if len(sys.argv) != 4:
        print("Usage: python script.py <json_file> <template_docx> <new_docx>")
        sys.exit(1)

    # Parameters
    json_file = sys.argv[1]
    word_file = sys.argv[2]
    new_word_file = sys.argv[3]

    # Call the function
    replace_placeholders_with_docxtpl(json_file, word_file, new_word_file)
