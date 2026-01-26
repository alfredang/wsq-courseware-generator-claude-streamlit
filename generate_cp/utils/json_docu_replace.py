import json
from docx import Document
import sys
from datetime import datetime

def replace_placeholders_in_doc(json_path, doc_path, new_doc_name, response_set):
    """Replace placeholders in a Word document with values from a JSON file."""
    
    def read_json(json_path):
        """Load the JSON file."""
        with open(json_path, 'r') as file:
            return json.load(file)

    def detect_and_replace_placeholders_in_paragraph(paragraph, replacements):
        """Replace placeholders in a paragraph's text while preserving formatting."""
        full_text = ''.join(run.text for run in paragraph.runs)
        
        for key, value in replacements.items():
            if isinstance(value, list):
                value = ' '.join(str(v) for v in value)  # Convert all elements to strings and join
            else:
                value = str(value)  # Ensure the value is a string
            full_text = full_text.replace(key, value)

        # Clear the existing runs and create new runs for the replaced text
        for i in range(len(paragraph.runs)):
            paragraph.runs[i].text = ''  # Clear the text in each run

        if len(paragraph.runs) > 0:
            paragraph.runs[0].text = full_text  # Assign the full text to the first run

    def process_table(table, replacements):
        """Recursively process tables, replacing placeholders."""
        for row in table.rows:
            for cell in row.cells:
                # Process paragraphs in the cell
                for paragraph in cell.paragraphs:
                    detect_and_replace_placeholders_in_paragraph(paragraph, replacements)
                # Check for nested tables within the cell and process them
                for nested_table in cell.tables:
                    process_table(nested_table, replacements)

    def extract_and_replace_text_from_docx(doc_path, replacements, new_doc_name):
        """Extract and replace text in the Word document based on the provided replacements."""
        doc = Document(doc_path)

        # Process paragraphs in the main document body
        for para in doc.paragraphs:
            detect_and_replace_placeholders_in_paragraph(para, replacements)

        # Process tables, including nested ones
        for table in doc.tables:
            process_table(table, replacements)

        # Save the updated document
        doc.save(new_doc_name)

        return new_doc_name

    # Function to get today's date
    def get_today_date():
        """Return today's date in the format 'Month Day, Year' (e.g., 'April 5, 2023')."""
        return datetime.now().strftime("%d %B %Y")

    # Load the JSON with values
    json_data = read_json(json_path)

    # Prepare replacements dictionary for placeholders
    tsc_combined = f"{json_data['course_info']['TSC Title']} {json_data['course_info']['TSC Code']}"


    # Prepare replacements dictionary for placeholders
    replacements = {
        "#Course": json_data["course_info"]["Course Title"],
        "#TSC": tsc_combined,
        "#Title": json_data["course_info"]["Industry"],
        "#Q1": response_set["What are the performance gaps in the industry?"],
        "#Q2": response_set["Why do you think this WSQ course will address the training needs for the industry?"],
        "#Date": get_today_date()  # Add today's date using the get_today_date function
    }

    # Replace the placeholders in the document using the JSON data directly
    updated_doc_path = extract_and_replace_text_from_docx(doc_path, replacements, new_doc_name)
    return updated_doc_path

if __name__ == "__main__":
    # Get input and output file paths from command-line arguments
    if len(sys.argv) != 4:
        print("Usage: python json_docu_replace.py <json_file> <word_template> <output_word_file>")
        sys.exit(1)

    json_file = sys.argv[1] #temp json files
    word_template = sys.argv[2] #iteration of the word templates
    output_word_file = sys.argv[3] # iteration of output files

    # Load the JSON file
    with open(json_file, 'r') as f:
        json_data = json.load(f)

    # Now ensure that both course_info and analyst_responses are present
    course_info = json_data.get("course_info", {})
    analyst_response = json_data["analyst_responses"][0]  # Use the first (and only) response in the list

    if not course_info:
        print("Error: 'course_info' is missing from the JSON.")
        sys.exit(1)

    # Call the function to replace placeholders and generate the new document
    updated_document = replace_placeholders_in_doc(json_file, word_template, output_word_file, analyst_response)
    print(f"Updated document saved as: {updated_document}") 
