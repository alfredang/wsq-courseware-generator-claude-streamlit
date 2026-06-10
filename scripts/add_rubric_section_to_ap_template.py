"""
One-off builder: add a "Marking Rubric" section to the AP DOCX template.

Appends a section that renders one Competent / Not Yet Competent (C/NYC)
checklist table per assessment method, driven by docxtpl Jinja loops over
`Assessment_Methods_Details` and `mtd.Rubric_Criteria`.

Idempotent: if the section marker already exists, the script exits without
re-adding it. Run from the project root:

    python scripts/add_rubric_section_to_ap_template.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATE = ".claude/skills/generate_assessment_plan/templates/AP_TGS-Ref-No_Course-Title_v1.docx"

HEADER_BG = "4472C4"      # steel blue
HEADER_FG = RGBColor(0xFF, 0xFF, 0xFF)
MARKER = "mtd.Rubric_Criteria"  # unique to the rubric section


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def style_run(run, *, bold=False, color=None, size=10):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell(cell, text, *, bold=False, color=None, bg=None, align=None, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    style_run(run, bold=bold, color=color, size=size)
    if bg:
        set_cell_bg(cell, bg)


def add_plain_paragraph(doc, text, *, bold=False, size=10, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    return p


def main():
    doc = Document(TEMPLATE)

    # Idempotency guard: bail if the rubric loop is already present (it lives
    # inside a table cell, so scan both paragraphs and table cells).
    blobs = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                blobs.append(c.text)
    if any(MARKER in b for b in blobs):
        print("Rubric section already present — nothing to do.")
        return

    # --- Section start ---------------------------------------------------
    # Page break before the new section.
    pb = doc.add_paragraph()
    pb.add_run().add_break(WD_BREAK.PAGE)

    h = doc.add_paragraph("Marking Rubric", style="Heading 1")

    add_plain_paragraph(
        doc,
        "The following rubric is used to mark each assessment. For every "
        "performance criterion, the assessor records a judgement of Competent (C) "
        "or Not Yet Competent (NYC), with remarks as evidence. A candidate must be "
        "assessed Competent on all criteria to achieve an overall Competent outcome.",
        italic=True,
    )

    # --- docxtpl loop over assessment methods ---------------------------
    add_plain_paragraph(doc, "{% for mtd in Assessment_Methods_Details %}")

    # Per-method sub-heading.
    sub = doc.add_paragraph(style="Heading 2")
    r = sub.add_run("{{ mtd.Assessment_Method }} ({{ mtd.Method_Abbreviation }})")
    r.font.name = "Calibri"

    # Rubric table: header row + one row per criterion. Each assessment method
    # has a fixed set of CRITERIA_PER_METHOD criteria (see rubrics.py), so the
    # rows are emitted explicitly with indexed access — this avoids docxtpl's
    # {%tr%} row-repeat (which mis-pairs when nested inside the method loop).
    CRITERIA_PER_METHOD = 4
    table = doc.add_table(rows=1 + CRITERIA_PER_METHOD, cols=4)
    table.style = "Table Grid"
    table.autofit = False

    headers = ["Performance Criteria", "Competent (C)", "Not Yet Competent (NYC)", "Remarks"]
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell(cell, text, bold=True, color=HEADER_FG, bg=HEADER_BG,
                 align=WD_ALIGN_PARAGRAPH.CENTER)

    for i in range(CRITERIA_PER_METHOD):
        body = table.rows[1 + i].cells
        set_cell(body[0], "{{ mtd.Rubric_Criteria[%d] }}" % i)
        set_cell(body[1], "☐", align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(body[2], "☐", align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(body[3], "")

    # Column widths.
    widths = [Inches(3.3), Inches(1.0), Inches(1.6), Inches(1.6)]
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = w

    # Per-method overall outcome line.
    add_plain_paragraph(
        doc,
        "Overall Outcome:   ☐ Competent      ☐ Not Yet Competent",
        bold=True,
    )
    add_plain_paragraph(doc, "")  # spacer

    add_plain_paragraph(doc, "{% endfor %}")

    doc.save(TEMPLATE)
    print("Marking Rubric section added to AP template.")


if __name__ == "__main__":
    main()
