"""
File: annex_assessment_v2.py

===============================================================================
Assessment into AP Annex Module (Local File Upload Version)
===============================================================================
Description:
    This module integrates assessment Q&A documents into the annex section of an
    Assessment Plan (AP) document using locally uploaded files. Unlike the Google
    Drive version, this module allows users to directly upload files through the
    Streamlit interface without requiring Google Drive authentication.

Main Functionalities:
    1. File Upload Interface:
         • Allows users to upload Assessment Plan (.docx)
         • Allows users to upload Question papers (.docx)
         • Allows users to upload Answer papers (.docx)

    2. Document Merging:
         • Merges Q&A documents into the annex section of the assessment plan
         • Inserts centered headers with annex labels
         • Maintains document formatting

    3. Download:
         • Provides download button for the merged document

Dependencies:
    - Standard Libraries: os, io, tempfile
    - External Libraries:
         • streamlit              – For building the web application interface
         • docx, docxcompose      – For document manipulation and merging

Usage:
    - Run the module with Streamlit:
          streamlit run annex_assessment_v2.py
    - Follow the on-screen instructions to:
          1. Upload the Assessment Plan document
          2. Upload Question and Answer papers for different assessment types
          3. Generate the merged document
          4. Download the final document

Author:
    Wong Xin Ping
Date:
    17 October 2025
===============================================================================
"""

import os
import io
import tempfile
import streamlit as st
from docx import Document
from docxcompose.composer import Composer
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

###############################################################################
# HELPER FUNCTIONS
###############################################################################

def insert_centered_header(doc, text, annex_label):
    """
    Inserts a centered header with a specified annex label and text into a Word document.

    Args:
        doc (Document): The python-docx Document object to modify.
        text (str): The header text to insert.
        annex_label (str): The annex label (e.g., "Annex A").

    Returns:
        None
    """
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    table = doc.add_table(rows=1, cols=1)
    table.allow_autofit = True
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    table_cell = table.cell(0, 0)
    table_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    row = table.rows[0]._tr
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), '12000')
    rowPr = row.get_or_add_trPr()
    rowPr.append(trHeight)

    paragraph = table_cell.paragraphs[0]
    run = paragraph.add_run(f"{annex_label}:\n{text}")
    run.bold = True
    run.font.size = Pt(24)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def get_annex_label(index: int) -> str:
    """
    Returns an annex label based on the provided index.

    Args:
        index (int): The zero-based index.

    Returns:
        str: The annex label (e.g., "Annex A", "Annex B").
    """
    letter = chr(ord("A") + index)
    return f"Annex {letter}"


def merge_documents(plan_file, assessment_files):
    """
    Merges question and answer documents into the annex section of the assessment plan.

    Args:
        plan_file: The uploaded Assessment Plan file object
        assessment_files: Dictionary of assessment files organized by type
            Format: {
                "WA (SAQ)": {"question": file_obj, "answer": file_obj},
                "PP": {"question": file_obj, "answer": file_obj},
                ...
            }

    Returns:
        bytes: The merged document as bytes
    """
    # Load the base document
    base_doc = Document(plan_file)
    composer = Composer(base_doc)
    annex_index = 0

    # Process each assessment type
    for assessment_type, files in assessment_files.items():
        question_file = files.get("question")
        answer_file = files.get("answer")

        # Add question paper if provided
        if question_file:
            annex_label = get_annex_label(annex_index)
            annex_index += 1

            temp_doc = Document()
            insert_centered_header(temp_doc, f"QUESTION PAPER OF {assessment_type} ASSESSMENT", annex_label)
            composer.append(temp_doc)

            question_doc = Document(question_file)
            composer.append(question_doc)

        # Add answer paper if provided
        if answer_file:
            annex_label = get_annex_label(annex_index)
            annex_index += 1

            temp_doc = Document()
            insert_centered_header(temp_doc, f"SUGGESTED ANSWER TO {assessment_type} ASSESSMENT QUESTIONS", annex_label)
            composer.append(temp_doc)

            answer_doc = Document(answer_file)
            composer.append(answer_doc)

    # Save to bytes
    output = io.BytesIO()
    composer.save(output)
    output.seek(0)
    return output.getvalue()


###############################################################################
# STREAMLIT APP
###############################################################################

def app():
    """
    Streamlit application to merge assessment Q&A documents into the annex of
    an Assessment Plan using locally uploaded files.
    """
    st.title("Add Assessment to AP")

    st.markdown("""
    ### Instructions:
    1. Upload your **Assessment Plan** document (.docx)
    2. Upload **Question** and **Answer** papers for each assessment type
    3. Click **Generate Merged Document** to combine them
    4. Download the final merged document

    **Note:** This version works without Google Drive - just upload files directly!
    """)

    # Upload Assessment Plan
    st.subheader("Step 1: Upload Assessment Plan")
    plan_file = st.file_uploader(
        "Upload Assessment Plan (.docx)",
        type=["docx"],
        key="plan_upload",
        help="Upload the main Assessment Plan document"
    )

    # Assessment types configuration
    assessment_types = ["WA (SAQ)", "PP", "CS", "Oral Questioning"]

    st.subheader("Step 2: Upload Assessment Documents")
    st.write("Upload Question and/or Answer papers for each assessment type:")

    # Dictionary to store uploaded files
    assessment_files = {}

    # Create upload widgets for each assessment type
    for assessment_type in assessment_types:
        with st.expander(f"📝 {assessment_type}", expanded=False):
            col1, col2 = st.columns(2)

            with col1:
                question_file = st.file_uploader(
                    f"Question Paper",
                    type=["docx"],
                    key=f"q_{assessment_type}",
                    help=f"Upload the question paper for {assessment_type}"
                )

            with col2:
                answer_file = st.file_uploader(
                    f"Answer Paper",
                    type=["docx"],
                    key=f"a_{assessment_type}",
                    help=f"Upload the answer paper for {assessment_type}"
                )

            # Store files if uploaded
            if question_file or answer_file:
                assessment_files[assessment_type] = {
                    "question": question_file,
                    "answer": answer_file
                }

    # Generate merged document
    st.subheader("Step 3: Generate Merged Document")

    if st.button("🚀 Generate Merged Document", type="primary"):
        if not plan_file:
            st.error("❌ Please upload an Assessment Plan document first.")
            return

        if not assessment_files:
            st.error("❌ Please upload at least one Question or Answer paper.")
            return

        try:
            with st.spinner("Merging documents..."):
                merged_doc_bytes = merge_documents(plan_file, assessment_files)

            st.success("✅ Document merged successfully!")

            # Generate download filename
            plan_filename = plan_file.name
            base_name = os.path.splitext(plan_filename)[0]
            download_filename = f"{base_name}_with_annex.docx"

            # Download button
            st.download_button(
                label="📥 Download Merged Document",
                data=merged_doc_bytes,
                file_name=download_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

            # Show summary
            st.info(f"**Summary:** Added {sum(1 for files in assessment_files.values() for f in [files.get('question'), files.get('answer')] if f)} documents to the annex.")

        except Exception as e:
            st.error(f"❌ Error merging documents: {e}")
            st.exception(e)

    # Help section
    with st.expander("ℹ️ Tips & Troubleshooting"):
        st.markdown("""
        **Tips:**
        - You don't need to upload both Question and Answer for each type - upload what you have
        - The annex labels (Annex A, B, C, etc.) are automatically generated
        - Documents are merged in the order: WA (SAQ), PP, CS, Oral Questioning

        **Common Issues:**
        - **"Invalid file format"**: Make sure all files are .docx format (not .doc or Google Docs)
        - **"Document not loading"**: Try re-uploading the file
        - **Formatting issues**: Ensure all documents are properly formatted Word documents
        """)
