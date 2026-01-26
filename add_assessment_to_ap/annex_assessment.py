"""
File: annex_assessment.py

===============================================================================
Assessment into AP Annex Module
===============================================================================
Description:
    This module integrates assessment Q&A documents into the annex section of an 
    Assessment Plan (AP) document. It processes course folder contents in Google Drive 
    to classify and retrieve relevant files (such as assessment plans, question papers, 
    and answer papers) using OpenAI. The module then merges the Q&A documents into the 
    annex section of the AP, updates version information on the cover page and within the 
    filename, and tracks changes via a version control record. Overall, it automates the 
    process of adding Assessment documents into the AP Annex.

Main Functionalities:
    0. Data Model & OpenAI Classification:
         • Defines the FileClassification model (using Pydantic) to structure file metadata 
           and classification details.
         • Implements classify_files_with_openai() to classify files (assessment plan, question, 
           and answer papers) using the GPT-4o-mini model.
    1. Helper Functions:
         • Provides utilities for authenticating with Google, downloading files from Google Drive, 
           parsing version strings, selecting the latest file versions, building method data for 
           Q&A documents, and deleting temporary files.
    2. Processing Course Folder:
         • Functions process_course_folder() and process_course_folder_direct() retrieve and classify 
           files from specific subfolders (e.g., "Assessment Plan" and "Assessment") and compile structured 
           data for further processing.
    3. Merging Documents into Annex & Version Updates:
         • insert_centered_header(): Inserts a centered header with an annex label into a Word document.
         • insert_answers_under_heading(): Merges Q&A documents into the annex section of the assessment plan.
         • update_cover_page_version(), update_version_number(), get_annex_label(): Update version numbers 
           and generate annex labels.
         • update_version_control_record(), bump_filename_version(), upload_updated_doc(): Handle version 
           control updates and file renaming/uploading on Google Drive.
    4. Tracking Processes:
         • track_edited_assessment_plan(): Records editing events by appending new entries to an Excel file.
    5. Main Function (app):
         • Provides a Streamlit web application interface that:
             - Prompts the user to enter a course TGS code.
             - Searches for and processes the corresponding course folder in Google Drive.
             - Downloads the assessment plan and Q&A files.
             - Merges the Q&A documents into the annex of the assessment plan.
             - Updates version information and uploads the revised document back to Google Drive.
             - Cleans up temporary downloaded files after processing.

Dependencies:
    - Standard Libraries: os, io, re, json, pandas, datetime
    - External Libraries:
         • streamlit              – For building the web application interface.
         • googleapiclient        – For interacting with Google Drive and Docs APIs.
         • google.oauth2          – For service account authentication.
         • docx, docxcompose      – For document manipulation and merging.
         • pydantic               – For data validation and modeling.
         • openai                 – For file classification via GPT-4o-mini.
         • Additional modules from typing and various docx enums/shared/oxml for styling.

Usage:
    - Ensure that all required API keys and credentials (e.g., OPENAI_API_KEY, GOOGLE_API_CREDS, 
      BROWSER_TOKEN, BROWSER_WEBDRIVER_ENDPOINT) are properly configured in st.secrets.
    - Run the module with Streamlit:
          streamlit run annex_assessment.py
    - Follow the on-screen instructions to:
          1. Input the course TGS code.
          2. Search for and process the corresponding course folder.
          3. Download, merge, and update assessment Q&A documents into the annex of the AP.
          4. Upload the final updated document back to Google Drive.
          5. Track changes via an Excel log.

Author:
    Derrick Lim
Date:
    3 March 2025
===============================================================================
"""

import os
import io
import re
import json
import pandas as pd
from datetime import datetime
import streamlit as st
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload, MediaFileUpload
from google.oauth2 import service_account
from docx import Document
from docxcompose.composer import Composer
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pydantic import BaseModel, ValidationError
from typing import Optional, List
from openai import OpenAI

###############################################################################
# 0. DATA MODEL & OPENAI CLASSIFICATION
###############################################################################

class FileClassification(BaseModel):
    file_id: str
    file_name: str
    is_assessment_plan: bool = False
    assessment_type: Optional[str] = None
    is_question_paper: bool = False
    is_answer_paper: bool = False
    version: Optional[str] = None


def classify_files_with_openai(file_metadata: List[dict]) -> List[FileClassification]:
    """
    Uses OpenAI to classify files into assessment plan, question paper, or answer paper.

    This function sends a prompt to the OpenAI API (using the GPT-4o-mini model) with the metadata of files
    from a WSQ course folder and expects a JSON array as a response. It then parses the JSON, filters for files 
    that are relevant (i.e. assessment plans, question papers, or answer papers), and returns a list of 
    FileClassification objects.

    Args:
        file_metadata (List[dict]): A list of dictionaries containing file metadata (e.g., file id and file name).

    Returns:
        List[FileClassification]: A list of classified files.
    """
    from settings.api_manager import load_api_keys
    api_keys = load_api_keys()
    openai_api_key = api_keys.get("OPENAI_API_KEY", "")
    client = OpenAI(api_key=openai_api_key)

    # Prepare file metadata for OpenAI
    file_info = "\n".join([f"{file['id']} - {file['name']}" for file in file_metadata])
    system_message = (
        "You are an AI assistant tasked with analyzing file names related to WSQ assessments. "
        "For each file, identify:\n"
        "1. Whether it is an assessment plan, question paper, or answer paper.\n"
        "   - Recognize variations like 'AP_' prefixes for assessment plans or '(Draft)' in filenames.\n"
        "2. For question or answer papers, identify the assessment type (e.g., WA (SAQ), PP, CS, Oral Questioning).\n"
        "   - Include incomplete or draft versions, noting them explicitly.\n"
        "3. Extract the version (e.g., v2.1, v1.5) if available.\n"
        "4. If a file cannot be classified, mark it as irrelevant.\n"
        "\n"
        "Return only a valid JSON array with no additional text or explanation, using this format:\n"
        "[{\n"
        "  \"file_id\": \"\",\n"
        "  \"file_name\": \"\",\n"
        "  \"is_assessment_plan\": false,\n"
        "  \"assessment_type\": null,\n"
        "  \"is_question_paper\": false,\n"
        "  \"is_answer_paper\": false,\n"
        "  \"version\": \"\"\n"
        "}]\n"
        "\n"
        "If no valid classifications are found, return an empty JSON array: []"
    )

    user_message = (
        "Below is a list of files from a WSQ course folder. "
        "Classify each file as an assessment plan, question paper, or answer paper. "
        "Recognize drafts (e.g., '(Draft)') and variations in naming. "
        "For unclassifiable files, mark them as irrelevant. "
        "Return the results in valid JSON format only.\n\n"
        "Files:\n"
        f"{file_info}\n\n"
        "Provide the classification in JSON format only."
    )

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": system_message},
            {"role": "user", "content": user_message},
        ],
        max_tokens=1000,
        temperature=0.3,
    )

    content = response.choices[0].message.content.strip()
    try:
        # Attempt to extract JSON content
        json_start = content.find("[")
        json_end = content.rfind("]") + 1
        if json_start != -1 and json_end != -1:
            json_content = content[json_start:json_end]
        else:
            raise ValueError("No JSON array found in response")

        json_content = json_content.replace("'", '"')  # ensure valid JSON
        raw_list = json.loads(json_content)

        # Filter only relevant files
        filtered_list = [
            item for item in raw_list
            if item["is_assessment_plan"] or item["is_question_paper"] or item["is_answer_paper"]
        ]

        return [FileClassification(**item) for item in filtered_list]

    except (json.JSONDecodeError, ValidationError, ValueError) as e:
        print("Error parsing OpenAI response:", e)
        print("OpenAI response content was:", content)
        return []


###############################################################################
# 1. HELPER FUNCTIONS
###############################################################################

def authenticate():
    """
    Authenticates with Google using credentials from Streamlit secrets.

    Returns:
        google.oauth2.service_account.Credentials: A credentials object for accessing Google services.
        Returns None if authentication fails.
    """
    try:
        creds = service_account.Credentials.from_service_account_info(
            st.secrets["GOOGLE_API_CREDS"]
        )
        return creds
    except Exception as e:
        print(f"Error during authentication: {e}")
        return None


def download_file(file_id, file_name, drive_service, download_dir="./downloads"):
    """
    Downloads a file (Google Doc or Word .docx) from Google Drive.

    The function exports Google Docs as .docx files and downloads files using the Google Drive API.

    Args:
        file_id (str): The unique ID of the file to download.
        file_name (str): The name of the file.
        drive_service: The Google Drive API service instance.
        download_dir (str, optional): The local directory to store downloaded files (default is "./downloads").

    Returns:
        str or None: The local file path to the downloaded file, or None if the file type is unsupported.
    """
    if not os.path.exists(download_dir):
        os.makedirs(download_dir, exist_ok=True)

    file_info = drive_service.files().get(fileId=file_id, fields="mimeType").execute()
    mime_type = file_info.get("mimeType")

    if mime_type == "application/vnd.google-apps.document":
        export_mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        request = drive_service.files().export_media(fileId=file_id, mimeType=export_mime_type)
        base_name, _ = os.path.splitext(file_name)
        file_name = base_name + ".docx"
    elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        request = drive_service.files().get_media(fileId=file_id)
    else:
        print(f"Skipping file (not .docx or Google Doc): {file_name}")
        return None

    file_path = os.path.join(download_dir, file_name)
    with io.BytesIO() as fh:
        downloader = MediaIoBaseDownload(fh, request)
        done = False
        while not done:
            _, done = downloader.next_chunk()
        fh.seek(0)
        with open(file_path, "wb") as f:
            f.write(fh.read())

    return file_path


def parse_version(version_str: Optional[str]) -> tuple:
    """
    Extracts a (major, minor) tuple from a version string like 'v2.1' or 'v1'.

    If the version string is invalid or missing, returns (0, 0).

    Args:
        version_str (Optional[str]): The version string to parse.

    Returns:
        tuple: A tuple (major, minor) as integers.
    """
    if version_str:
        match = re.match(r"v(\d+)(\.\d+)?", version_str.lower())
        if match:
            major = int(match.group(1))
            minor = int(match.group(2).lstrip(".")) if match.group(2) else 0
            return (major, minor)
    return (0, 0)


def select_latest_version(file_classifications: List[FileClassification]) -> Optional[FileClassification]:
    """
    Selects and returns the FileClassification with the highest version number.

    Args:
        file_classifications (List[FileClassification]): A list of FileClassification objects.

    Returns:
        Optional[FileClassification]: The file with the highest version, or None if the list is empty.
    """

    sorted_files = sorted(
        file_classifications,
        key=lambda f: parse_version(f.version),
        reverse=True
    )
    return sorted_files[0] if sorted_files else None


def select_latest_assessment_plan(file_classifications: List[FileClassification]) -> Optional[FileClassification]:
    """
    Selects the latest version of an assessment plan from a list of FileClassification objects.

    Args:
        file_classifications (List[FileClassification]): A list of FileClassification objects.

    Returns:
        Optional[FileClassification]: The latest assessment plan, or None if none are found.
    """
    plans = [f for f in file_classifications if f.is_assessment_plan]
    return select_latest_version(plans)


def build_method_data(file_classifications: List[FileClassification], abbreviations: List[str]) -> dict:
    """
    Builds a dictionary mapping assessment method abbreviations to the latest question and answer files.

    For each abbreviation, this function filters the file classifications for matching question and answer papers,
    selects the latest versions, and stores them in a nested dictionary.

    Args:
        file_classifications (List[FileClassification]): A list of FileClassification objects.
        abbreviations (List[str]): A list of assessment method abbreviations (e.g., ["WA (SAQ)", "PP"]).

    Returns:
        dict: A dictionary of the form {abbr: {"question": {...}, "answer": {...}}}.
    """
    method_data = {}

    for abbr in abbreviations:
        question_files = [
            f for f in file_classifications
            if f.assessment_type == abbr and f.is_question_paper
        ]
        answer_files = [
            f for f in file_classifications
            if f.assessment_type == abbr and f.is_answer_paper
        ]

        latest_question = select_latest_version(question_files)
        latest_answer = select_latest_version(answer_files)

        if latest_question or latest_answer:
            method_data[abbr] = {
                "question": {
                    "id": latest_question.file_id,
                    "name": latest_question.file_name
                } if latest_question else None,
                "answer": {
                    "id": latest_answer.file_id,
                    "name": latest_answer.file_name
                } if latest_answer else None,
            }

    return method_data


def delete_irrelevant_files(download_dir="./downloads", keep_filename=None):
    """
    Deletes all files in the specified download directory except for a given filename.

    If keep_filename is None, deletes all files.

    Args:
        download_dir (str, optional): The directory containing downloaded files (default is "./downloads").
        keep_filename (Optional[str]): The filename to keep.
    """
    for file_name in os.listdir(download_dir):
        file_path = os.path.join(download_dir, file_name)
        if keep_filename is not None and file_name == keep_filename:
            continue
        try:
            os.remove(file_path)
            print(f"Deleted local file: {file_path}")
        except PermissionError:
            print(f"Could not delete (in use): {file_path}")


###############################################################################
# 2. PROCESS COURSE FOLDER (WITH OPENAI CLASSIFICATION)
###############################################################################

def process_course_folder(course_folder_id, drive_service, abbreviations):
    """
    Processes the course folder by classifying files in its 'Assessment Plan' and 'Assessment' subfolders.

    Steps:
      1. Retrieves all subfolders within the given course folder.
      2. Looks for the 'Assessment Plan' folder and downloads its files.
      3. Classifies the files using OpenAI and selects the latest assessment plan.
      4. Looks for the 'Assessment' folder, classifies its files, and builds method data for Q&A documents.

    Args:
        course_folder_id (str): The ID of the course folder.
        drive_service: The Google Drive API service instance.
        abbreviations (List[str]): A list of assessment method abbreviations.

    Returns:
        dict or None: A dictionary containing:
            - "assessment_plan": { "id": <file_id>, "name": <file_name> }
            - "method_data": A dictionary of method data.
        Returns None if no valid assessment plan is identified.
    """
    # Retrieve subfolders
    subfolders = drive_service.files().list(
        q=f"'{course_folder_id}' in parents and mimeType='application/vnd.google-apps.folder'"
    ).execute().get('files', [])

    # Find "Assessment Plan" folder (case-insensitive check)
    assessment_plan_folder = next(
        (f for f in subfolders if f['name'].strip().lower() == 'assessment plan'),
        None
    )

    plan_files = []
    if assessment_plan_folder:
        plan_files = drive_service.files().list(
            q=(
                f"'{assessment_plan_folder['id']}' in parents and "
                "(mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document' or "
                "mimeType='application/vnd.google-apps.document')"
            )
        ).execute().get('files', [])
    else:
        print(f"No 'Assessment Plan' folder found in {course_folder_id}.")

    if plan_files:
        plan_classifications = classify_files_with_openai(plan_files)
        assessment_plan = select_latest_assessment_plan(plan_classifications)
    else:
        print(f"No files found in Assessment Plan folder for {course_folder_id}.")
        assessment_plan = None

    # If no plan found, check all files in the course folder for a possible misclassified plan
    if not assessment_plan:
        print(f"No valid assessment plan found for {course_folder_id}. Checking further...")
        all_course_files = drive_service.files().list(
            q=(
                f"'{course_folder_id}' in parents and "
                "(mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document' or "
                "mimeType='application/vnd.google-apps.document')"
            )
        ).execute().get('files', [])
        all_classifications = classify_files_with_openai(all_course_files)
        assessment_plan = select_latest_assessment_plan(all_classifications)

    if not assessment_plan:
        print(f"No valid assessment plan could be identified for folder ID {course_folder_id}.")
        return None

    # Find "Assessment" folder (case-insensitive check)
    assessment_folder = next(
        (f for f in subfolders if f['name'].strip().lower() == 'assessment'),
        None
    )

    if assessment_folder:
        assessment_files = drive_service.files().list(
            q=(
                f"'{assessment_folder['id']}' in parents and "
                "(mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document' or "
                "mimeType='application/vnd.google-apps.document')"
            )
        ).execute().get('files', [])
    else:
        print(f"No 'Assessment' folder found for {course_folder_id}.")
        assessment_files = []

    if assessment_files:
        assessment_classifications = classify_files_with_openai(assessment_files)
        method_data = build_method_data(assessment_classifications, abbreviations)
    else:
        method_data = {}

    return {
        "assessment_plan": {
            "id": assessment_plan.file_id,
            "name": assessment_plan.file_name
        },
        "method_data": method_data,
    }


###############################################################################
# 3. MERGING DOCUMENTS INTO ANNEX & VERSION UPDATES
###############################################################################

def insert_centered_header(doc, text, annex_label):
    """
    Inserts a centered header with a specified annex label and text into a Word document.

    The header is inserted using a single-cell table that is centered horizontally and vertically.
    A page break is added after the table.

    Args:
        doc (Document): The python-docx Document object to modify.
        text (str): The header text to insert.
        annex_label (str): The annex label (e.g., "Annex A").

    Returns:
        None
    """

    # Insert the centered header using a single-cell table...
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    table = doc.add_table(rows=1, cols=1)
    table.allow_autofit = True
    table.alignment = WD_TABLE_ALIGNMENT.CENTER  # or CENTER, RIGHT, etc.

    # Vertical alignment
    table_cell = table.cell(0, 0)
    table_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # Row height (10 inches as an example)
    row = table.rows[0]._tr
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), '12000')  # 12000 twips ~ 8.3 inches; adjust as needed
    rowPr = row.get_or_add_trPr()
    rowPr.append(trHeight)

    # Add the centered text to the table
    paragraph = table_cell.paragraphs[0]
    run = paragraph.add_run(f"{annex_label}:\n{text}")
    run.bold = True
    run.font.size = Pt(24)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add a page break after the table
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def insert_answers_under_heading(plan_path, heading_map, method_data):
    """
    Inserts question and answer documents into the annex section of the assessment plan.

    This function:
      - Reads the base plan document.
      - Appends headers and corresponding Q&A documents using a mapping of headings to assessment abbreviations.
      - Saves the updated document.

    Args:
        plan_path (str): The file path to the assessment plan document.
        heading_map (dict): A dictionary mapping heading text to assessment method abbreviations.
        method_data (dict): The dictionary produced by build_method_data containing Q&A file info.

    Returns:
        tuple: A tuple (updated_doc_path, changes_made) where:
            - updated_doc_path (str) is the path to the modified document.
            - changes_made (bool) indicates if any Q&A documents were appended.
    """
    base_doc = Document(plan_path)
    composer = Composer(base_doc)
    changes_made = False
    annex_index = 0

    for heading_text, abbr in heading_map.items():
        if abbr in method_data:
            files = method_data[abbr]
            q_file = files.get('question')
            a_file = files.get('answer')

            # If there's a question doc
            if q_file and 'local_path' in q_file:

                annex_label = get_annex_label(annex_index)  # e.g. "Annex A"
                annex_index += 1
                
                temp_doc = Document()
                insert_centered_header(temp_doc, f"QUESTION PAPER OF {abbr} ASSESSMENT", annex_label)
                composer.append(temp_doc)

                question_doc = Document(q_file['local_path'])
                composer.append(question_doc)
                changes_made = True

            # If there's an answer doc
            if a_file and 'local_path' in a_file:

                annex_label = get_annex_label(annex_index)  # e.g. "Annex A"
                annex_index += 1

                temp_doc = Document()
                insert_centered_header(temp_doc, f"SUGGESTED ANSWER TO {abbr} ASSESSMENT QUESTIONS", annex_label)
                composer.append(temp_doc)

                answer_doc = Document(a_file['local_path'])
                composer.append(answer_doc)
                changes_made = True

    if changes_made:
        updated_path = plan_path.replace(".docx", "_Answers_Only.docx")
        composer.save(updated_path)

        # Auto-fit tables in the merged doc
        updated_doc = Document(updated_path)
        updated_doc.save(updated_path)

        return updated_path, True
    else:
        print("No Q&A appended to annex.")
        return plan_path, False


def update_cover_page_version(doc_path):
    """
    Increments the version number on the cover page of a document.

    The function searches for a paragraph starting with "Version", parses the current version,
    increments it by one major version (e.g., from "Version 1.0" to "Version 2.0"), updates the text,
    and saves the document with a new filename suffixed with "_Updated".

    Args:
        doc_path (str): The file path to the document to update.

    Returns:
        str: The file path to the updated document.

    Raises:
        ValueError: If no "Version" text is found or the version format is invalid.
    """
    doc = Document(doc_path)
    updated = False

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("Version"):
            # Attempt to parse old version
            try:
                current_version = float(text.split()[1])
                new_version = f"Version {int(current_version) + 1}.0"
            except (IndexError, ValueError):
                raise ValueError("Invalid version format on cover page.")

            paragraph.clear()
            run = paragraph.add_run(new_version)
            run.font.name = "Arial"
            run.font.size = Pt(14)
            run.font.bold = True
            r = run._element
            r.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

            print(f"Updated cover page to: {new_version}")
            updated = True
            break

    if not updated:
        raise ValueError("No 'Version' text found on the cover page.")

    updated_doc_path = doc_path.replace(".docx", "_Updated.docx")
    doc.save(updated_doc_path)
    print(f"Cover page version updated => {updated_doc_path}")
    return updated_doc_path


def update_version_number(last_version_str):
    """
    Calculates the next major version from a version string.

    Converts a version string like '1.0' or '2.1' to the next major version (e.g., '2.0' or '3.0').

    Args:
        last_version_str (str): The current version string.

    Returns:
        str: The next major version string.
    """
    try:
        last_version = float(last_version_str)
        next_version = int(last_version) + 1
        return f"{next_version}.0"
    except ValueError:
        return "2.0"

def get_annex_label(index: int) -> str:
    """
    Returns an annex label based on the provided index.

    For example, index=0 returns "Annex A", index=1 returns "Annex B", etc.

    Args:
        index (int): The zero-based index.

    Returns:
        str: The annex label.
    """
    letter = chr(ord("A") + index)
    return f"Annex {letter}"

def update_version_control_record(doc_path, changes, developer="Tertiary Infotech"):
    """
    Appends a new row to the version control table in the document to track changes.

    The table is expected to have the columns: Version, Effective Date, Changes, Developer.

    Args:
        doc_path (str): The file path to the document.
        changes (str): A description of the changes made.
        developer (str, optional): The name of the developer making the update (default is "Tertiary Infotech").

    Returns:
        None
    """
    doc = Document(doc_path)

    if not doc.tables:
        print("No tables found. Can't update version control record.")
        return

    version_table = doc.tables[0]  # assume first table is version control
    # The last version in the first column
    last_version_str = version_table.rows[-1].cells[0].text.strip()
    next_version_str = update_version_number(last_version_str)

    eff_date = datetime.now().strftime("%d %b %Y")
    row_cells = version_table.add_row().cells
    row_cells[0].text = next_version_str
    row_cells[1].text = eff_date
    row_cells[2].text = changes
    row_cells[3].text = developer

    # Center/align each column
    for idx, cell in enumerate(row_cells):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            if idx == 2:  # "Changes" column left-aligned
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(doc_path)
    print(f"Version control record updated => {doc_path}")


def bump_filename_version(doc_path):
    """
    Increments the numeric version in the filename.

    For example, renames 'xxx_v2.docx' to 'xxx_v3.docx'. For fractional versions (e.g., v2.1),
    increments the float by +1.0 (e.g., v2.1 becomes v3.1). If no version is found, returns the original path.

    Args:
        doc_path (str): The original file path.

    Returns:
        str: The new file path with an updated version number.
    """
    base, ext = os.path.splitext(doc_path)
    pattern = r'(?:-|_)v(\d+(\.\d+)*)'

    match = re.search(pattern, base, re.IGNORECASE)
    if not match:
        return doc_path  # no 'vXX' in filename => do nothing

    old_version_str = match.group(1)
    try:
        old_ver_float = float(old_version_str)
        new_ver_float = old_ver_float + 1.0
        # Keep the minor if it existed
        new_version_str = f"{int(new_ver_float)}" if old_ver_float.is_integer() else f"{new_ver_float}"
    except ValueError:
        new_version_str = "2"  # fallback

    def replacement(m):
        prefix = m.group(0)[:-len(old_version_str)]
        return prefix + new_version_str

    new_base = re.sub(pattern, replacement, base, flags=re.IGNORECASE)
    new_doc_path = f"{new_base}{ext}"

    if os.path.exists(doc_path):
        os.rename(doc_path, new_doc_path)
        print(f"Renamed {doc_path} => {new_doc_path}")

    return new_doc_path


def upload_updated_doc(drive_service, file_id, local_doc_path, original_filename):
    """
    Uploads an updated document to Google Drive, cleaning up the filename by removing specific suffixes.

    Args:
        drive_service: The Google Drive API service instance.
        file_id (str): The ID of the file to update.
        local_doc_path (str): The local path of the updated document.
        original_filename (str): The original filename.

    Returns:
        dict: The updated file metadata as returned by the Drive API.
    """
    base_name, ext = os.path.splitext(original_filename)
    # Remove suffixes for cleanliness
    base_name = re.sub(r'_Answers_Only', '', base_name, flags=re.IGNORECASE)
    base_name = re.sub(r'_Updated', '', base_name, flags=re.IGNORECASE)
    new_filename = base_name + ext

    media_body = MediaFileUpload(
        local_doc_path,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        resumable=True
    )

    updated_file = drive_service.files().update(
        fileId=file_id,
        media_body=media_body,
        body={"name": new_filename}
    ).execute()

    print(f"Updated file in Google Drive: {updated_file.get('name')} (ID: {file_id})")
    return updated_file

###############################################################################
# 4. TRACKING PROCESSES FUNCTION
###############################################################################

def track_edited_assessment_plan(course_title, excel_file="edited_assessment_plans.xlsx"):
    """
    Records the editing of an assessment plan by appending a new entry to an Excel file.

    If the Excel file does not exist, it is created with the appropriate headers.

    Args:
        course_title (str): The title of the course whose assessment plan was edited.
        excel_file (str, optional): The path to the Excel file (default is "edited_assessment_plans.xlsx").

    Returns:
        None
    """
    # Check if the Excel file exists
    if not os.path.exists(excel_file):
        # Create a new DataFrame with headers if the file doesn't exist
        df = pd.DataFrame(columns=["Course Title"])
    else:
        # Load the existing Excel file
        df = pd.read_excel(excel_file)

    # Create a new DataFrame for the new entry
    new_entry = pd.DataFrame([[course_title]], columns=["Course Title"])

    # Concatenate the new entry with the existing DataFrame
    df = pd.concat([df, new_entry], ignore_index=True)

    # Save the updated DataFrame back to the Excel file
    df.to_excel(excel_file, index=False)
    print(f"Updated tracking file: {excel_file} with new entry: {course_title}")


###############################################################################
# 5. MAIN FUNCTION
###############################################################################

def process_course_folder_direct(course_folder_id, drive_service, abbreviations):
    """
    Processes the 'Assessment Plan' and 'Assessment' subfolders in a course folder.

    The function:
      - Retrieves subfolders within the given course folder.
      - Looks for an 'Assessment Plan' folder and selects the latest valid plan.
      - Retrieves an 'Assessment' folder (if available) and builds a dictionary of Q&A files for each assessment type.
      - Returns a dictionary with keys "assessment_plan" and "method_data".

    Args:
        course_folder_id (str): The ID of the course folder.
        drive_service: The Google Drive API service instance.
        abbreviations (List[str]): A list of assessment method abbreviations to filter files.

    Returns:
        dict or None: A dictionary with processed data, or None if no valid assessment plan is found.
    """
    # Define the target folder names
    target_folders = {"assessment plan": None, "assessment": None}

    # Retrieve all subfolders in the course folder
    subfolders = drive_service.files().list(
        q=f"'{course_folder_id}' in parents and mimeType='application/vnd.google-apps.folder'",
        fields="files(id, name)"
    ).execute().get("files", [])

    # Match and map the folder names
    for subfolder in subfolders:
        folder_name = subfolder["name"].strip().lower()  # Normalize folder name
        if folder_name in target_folders:
            target_folders[folder_name] = subfolder["id"]

    # Retrieve 'Assessment Plan' files
    assessment_plan = None
    if target_folders["assessment plan"]:
        plan_files = drive_service.files().list(
            q=(
                f"'{target_folders['assessment plan']}' in parents and "
                "(mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document' or "
                "mimeType='application/vnd.google-apps.document')"
            ),
            fields="files(id, name)"
        ).execute().get("files", [])

        if plan_files:
            plan_classifications = classify_files_with_openai(plan_files)
            assessment_plan = select_latest_assessment_plan(plan_classifications)
        else:
            print(f"No files found in 'Assessment Plan' folder for course folder ID {course_folder_id}.")

    # Retrieve 'Assessment' files
    method_data = {}
    if target_folders["assessment"]:
        assessment_files = drive_service.files().list(
            q=(
                f"'{target_folders['assessment']}' in parents and "
                "(mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document' or "
                "mimeType='application/vnd.google-apps.document')"
            ),
            fields="files(id, name)"
        ).execute().get("files", [])

        if assessment_files:
            assessment_classifications = classify_files_with_openai(assessment_files)
            method_data = build_method_data(assessment_classifications, abbreviations)

    # Return the processed data
    if not assessment_plan:
        print(f"No valid assessment plan identified for course folder ID {course_folder_id}.")
        return None

    return {
        "assessment_plan": {
            "id": assessment_plan.file_id,
            "name": assessment_plan.file_name
        },
        "method_data": method_data,
    }

def app():
    """
    Streamlit application to process a course folder and integrate assessment Q&A into the annex of the assessment plan document.

    The app performs the following steps:
      1. Authenticates with Google Drive.
      2. Prompts the user to enter a course TGS code.
      3. Searches for the course folder within "1 WSQ Documents" by querying for folders whose name contains the TGS code.
      4. Processes the course folder to classify and select files from the 'Assessment Plan' and 'Assessment' subfolders.
      5. Downloads the assessment plan and Q&A files.
      6. Merges the Q&A documents into the annex of the assessment plan.
      7. Updates the filename/version and uploads the updated document back to Google Drive.

    Returns:
        None
    """
    st.title("📄 Add Assessment to AP")
    st.subheader("Instructions:")
    st.markdown("""
    **What you must input:**  
    Please enter the **TGS code** for the course folder.  
    *Example: `TGS-2023039343`*  

    #### Prerequisite for the folder:  
    The course folder must be located under the **"1 WSQ Documents"** folder in Google Drive, and its name must contain the TGS code you provide.
    """)
    st.markdown("""
    **📂 File Organization and Naming Instructions**

    - **`Assessment Plan`** folder:  
      Place the **assessment plan file** here.  
      Example: `Assessment Plan_TGS-[Course Code] - [Course Title]_vX.docx`

    - **`Assessment`** folder:  
      Place **question and answer files** here.  
      Examples:  
      - `WA (SAQ) - [Course Title] - vX.docx`  
      - `Answer to WA (SAQ) - [Course Title] - vX.docx`
    """)

    # Abbreviations and heading map
    abbreviations = ["WA (SAQ)", "PP", "CS", "RP", "Oral Questioning"]
    heading_map = {
        "Assessment Questions and Answers for WA(SAQ)": "WA (SAQ)",
        "Assessment Questions and Practical Performance": "PP",
        "Assessment Questions and Case Study": "CS",
        "Assessment Questions and Oral Questioning (OQ)": "Oral Questioning",
    }

    # Authenticate with Google Drive
    with st.spinner("Authenticating with Google Drive..."):
        creds = authenticate()
    if not creds:
        st.error("Authentication failed. Please check your credentials.")
        return

    drive_service = build("drive", "v3", credentials=creds)

    # Helper function to extract TGS code from text
    import re
    def extract_tgs_code(text):
        match = re.search(r'TGS-\d+', text, re.IGNORECASE)
        return match.group(0).upper() if match else None

    # Input: Course TGS code
    course_tgs_code = st.text_input("Enter the Course TGS code:", "", placeholder="e.g., TGS-2023039181").strip().upper()

    if st.button("Process Document"):
        if not course_tgs_code:
            st.error("Please provide a course TGS code to proceed.")
            return

        try:
            # Retrieve the top-level folder "1 WSQ Documents"
            with st.spinner("Looking for the top-level folder..."):
                top_folder_name = "1 WSQ Documents"
                wsq_folder_list = drive_service.files().list(
                    q=f"name='{top_folder_name}' and mimeType='application/vnd.google-apps.folder'",
                    fields="files(id, name)"
                ).execute().get("files", [])

                if not wsq_folder_list:
                    st.error(f"Top-level folder '{top_folder_name}' not found.")
                    return

                wsq_documents_folder_id = wsq_folder_list[0]["id"]

            # Search for the course folder using a query that checks if the folder name contains the TGS code
            with st.spinner("Searching for the course folder..."):
                query = (
                    f"'{wsq_documents_folder_id}' in parents and "
                    "mimeType='application/vnd.google-apps.folder' and "
                    f"name contains '{course_tgs_code}'"
                )
                course_folders = drive_service.files().list(
                    q=query,
                    fields="files(id, name)"
                ).execute().get("files", [])

                # Further filter by extracting the TGS code from folder names
                matching_course_folder = next(
                    (folder for folder in course_folders if extract_tgs_code(folder["name"]) == course_tgs_code),
                    None
                )

                if not matching_course_folder:
                    st.error(f"Course folder containing TGS Code '{course_tgs_code}' not found.")
                    return

                course_folder_id = matching_course_folder["id"]
                st.success(f"Found course folder with TGS Code: {course_tgs_code}")

            # Process the course folder
            with st.spinner("Processing the course folder..."):
                result = process_course_folder_direct(course_folder_id, drive_service, abbreviations)
                if not result:
                    st.error("No valid assessment plan or files found in the folder.")
                    return

                assessment_plan = result["assessment_plan"]
                method_data = result["method_data"]

                st.write(f"Assessment Plan: {assessment_plan['name']}")
                st.write(f"Method Data: {json.dumps(method_data, indent=2)}")

                # Download the assessment plan
                with st.spinner("Downloading assessment plan..."):
                    plan_path = download_file(assessment_plan["id"], assessment_plan["name"], drive_service)
                if not plan_path:
                    st.error(f"Failed to download assessment plan: {assessment_plan['name']}")
                    return

                # Download and append Q&A documents
                with st.spinner("Downloading and appending Q&A documents..."):
                    for abbr, doc_dict in method_data.items():
                        for doc_type in ["question", "answer"]:
                            doc_info = doc_dict.get(doc_type)
                            if doc_info:
                                local_path = download_file(doc_info["id"], doc_info["name"], drive_service)
                                if local_path:
                                    doc_info["local_path"] = local_path

                # Merge Q&A into the annex
                with st.spinner("Merging Q&A into the annex..."):
                    merged_doc_path, changes_made = insert_answers_under_heading(plan_path, heading_map, method_data)
                if changes_made:
                    with st.spinner("Renaming and uploading the file..."):
                        final_doc_path = bump_filename_version(merged_doc_path)
                        upload_updated_doc(
                            drive_service=drive_service,
                            file_id=assessment_plan["id"],
                            local_doc_path=final_doc_path,
                            original_filename=assessment_plan["name"]
                        )
                    st.success(f"Updated assessment plan uploaded: {assessment_plan['name']}")
                else:
                    st.info("No changes made to the assessment plan.")
                
                # Delete downloaded files now that processing is complete
                delete_irrelevant_files(download_dir="./downloads")
        except Exception as e:
            st.error(f"An error occurred: {e}")
