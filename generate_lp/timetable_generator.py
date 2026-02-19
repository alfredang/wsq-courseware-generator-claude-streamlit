"""
Timetable Generator Module

Pure Python lesson plan schedule builder using the barrier algorithm,
and DOCX 4-column table generator.

No AI/agent calls - instant schedule generation.

Algorithm rules (from SKILL.md):
- 9:00 AM to 6:00 PM daily
- Lunch: fixed 45 mins, 12:30 PM - 1:15 PM
- Assessment: fixed 4:00 PM - 6:00 PM on last day only
- Topic duration = instructional_hours * 60 / num_topics (equal allocation)
- Topics CAN split across lunch/day-end barriers
- Minimum session: 15 minutes (avoid tiny splits; use Break instead)
- Fill remaining gaps with Breaks to fit exactly 9AM-6PM
"""

import os
import re
import tempfile
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docxtpl import DocxTemplate
from PIL import Image


# =============================================================================
# Constants
# =============================================================================

DAY_START = 9 * 60            # 9:00 AM in minutes from midnight
DAY_END = 18 * 60             # 6:00 PM
LUNCH_START = 12 * 60 + 30    # 12:30 PM
LUNCH_DURATION = 45            # 45 minutes
LUNCH_END = LUNCH_START + LUNCH_DURATION  # 1:15 PM
ASSESS_START = 16 * 60         # 4:00 PM (last day only)
MIN_SESSION = 15               # Minimum session length in minutes

HEADING_COLOR = RGBColor(0x44, 0x72, 0xC4)  # Steel blue


# =============================================================================
# Utility Functions
# =============================================================================

def _parse_hours(value) -> float:
    """Parse hours from various formats: '14 hrs', '2 hours', '16', 14.5, etc."""
    if isinstance(value, (int, float)):
        return float(value)
    if not value:
        return 0.0
    s = str(value).strip()
    match = re.search(r'(\d+\.?\d*)', s)
    return float(match.group(1)) if match else 0.0


def _round5(minutes: int) -> int:
    """Round minutes to the nearest 5-minute step."""
    return 5 * round(minutes / 5)


def _fmt_time(minutes_from_midnight: int) -> str:
    """Format minutes from midnight as 'H:MM AM/PM'."""
    h = minutes_from_midnight // 60
    m = minutes_from_midnight % 60
    period = "AM" if h < 12 else "PM"
    display_h = h % 12 or 12
    return f"{display_h}:{m:02d} {period}"


def _make_slot(start: int, end: int, description: str, methods: str, **extra) -> dict:
    """Create a schedule slot dictionary with optional extra fields (lu_num, lu_title)."""
    duration = end - start
    slot = {
        "timing": f"{_fmt_time(start)} - {_fmt_time(end)}",
        "duration": f"{duration} mins",
        "description": description,
        "methods": methods,
    }
    slot.update(extra)
    return slot


def _collect_lu_blocks(context: dict) -> list:
    """Collect Learning Unit blocks with aggregated topics for LU-level scheduling.

    Each LU becomes one schedulable block. Topic numbering matches the CP:
    each LU restarts at T1. Returns list of LU block dicts.
    """
    lu_blocks = []
    for lu_idx, lu in enumerate(context.get("Learning_Units", []), start=1):
        lu_title = lu.get("LU_Title", f"Learning Unit {lu_idx}")
        # Strip existing "LU1:" or "LU 1:" prefix to avoid "LU1: LU1: ..."
        lu_title = re.sub(r'^LU\s*\d+\s*[:.\-]\s*', '', lu_title).strip() or lu_title

        methods = lu.get("Instructional_Methods", [])
        normalized = []
        for m in methods:
            if m == "Classroom":
                normalized.append("Lecture")
            elif m == "Practical":
                normalized.append("Practice")
            elif m == "Discussion":
                normalized.append("Group Discussion")
            else:
                normalized.append(m)

        topic_labels = []
        for t_idx, topic in enumerate(lu.get("Topics", []), start=1):
            title = topic.get("Topic_Title", f"Topic {t_idx}")
            if re.match(r'^T\d+[\s:.]', title):
                topic_labels.append(title)
            else:
                topic_labels.append(f"T{t_idx}: {title}")

        lu_blocks.append({
            "lu_num": lu_idx,
            "lu_title": lu_title,
            "topic_labels": topic_labels,
            "methods": normalized or ["Lecture"],
            "num_topics": len(topic_labels),
        })
    return lu_blocks


def extract_unique_instructional_methods(course_context):
    """
    Extracts and processes unique instructional method combinations from the provided course context.

    Args:
        course_context: Dictionary containing course details with Learning Units.

    Returns:
        Set of unique instructional method combinations as strings.
    """
    unique_methods = set()

    valid_im_pairs = {
        ("Lecture", "Didactic Questioning"),
        ("Lecture", "Peer Sharing"),
        ("Lecture", "Group Discussion"),
        ("Demonstration", "Practice"),
        ("Demonstration", "Group Discussion"),
        ("Case Study",),
        ("Role Play",),
    }

    for lu in course_context.get("Learning_Units", []):
        extracted_methods = lu.get("Instructional_Methods", [])

        corrected_methods = []
        for method in extracted_methods:
            if method == "Classroom":
                corrected_methods.append("Lecture")
            elif method == "Practical":
                corrected_methods.append("Practice")
            elif method == "Discussion":
                corrected_methods.append("Group Discussion")
            else:
                corrected_methods.append(method)

        method_pairs = set()
        for pair in valid_im_pairs:
            if all(method in corrected_methods for method in pair):
                method_pairs.add(", ".join(pair))

        if not method_pairs and corrected_methods:
            if len(corrected_methods) == 1:
                method_pairs.add(corrected_methods[0])
            elif len(corrected_methods) == 2:
                method_pairs.add(", ".join(corrected_methods))
            else:
                method_pairs.add(", ".join(corrected_methods[:2]))
                if len(corrected_methods) > 2:
                    method_pairs.add(", ".join(corrected_methods[-2:]))

        unique_methods.update(method_pairs)

    return unique_methods


# =============================================================================
# Schedule Builder (Barrier Algorithm)
# =============================================================================

def build_lesson_plan_schedule(context: dict) -> dict:
    """Build a lesson plan schedule using the barrier algorithm at the LU level.

    Each Learning Unit is scheduled as a single block. If an LU spans a barrier
    (lunch, day-end), it splits into multiple slots marked with is_contd.

    Returns:
        Dict with keys: num_days, instructional_hours, assessment_hours,
        per_topic_mins, days (dict of day_num -> list of slot dicts).
    """
    total_hours = _parse_hours(context.get("Total_Course_Duration_Hours", "16"))
    instr_hours = _parse_hours(context.get("Total_Training_Hours", ""))
    if not instr_hours:
        instr_hours = total_hours
    assess_hours = _parse_hours(context.get("Total_Assessment_Hours", "0"))

    num_days = max(1, round(total_hours / 8)) if total_hours >= 8 else 1

    lu_blocks = _collect_lu_blocks(context)
    total_topics = sum(b["num_topics"] for b in lu_blocks)
    if total_topics == 0:
        return {
            "num_days": num_days,
            "instructional_hours": instr_hours,
            "assessment_hours": assess_hours,
            "per_topic_mins": 0,
            "days": {},
        }

    per_topic = (instr_hours * 60) / total_topics
    assess_mins = int(assess_hours * 60)

    # Calculate duration for each LU block
    for block in lu_blocks:
        block["duration_mins"] = block["num_topics"] * per_topic

    days = {}
    lu_idx = 0
    lu_remaining = lu_blocks[0]["duration_mins"]
    is_contd = False

    for day in range(1, num_days + 1):
        slots = []
        current = DAY_START
        is_last_day = (day == num_days)
        lunch_done = False
        instr_end = ASSESS_START if (is_last_day and assess_mins > 0) else DAY_END

        while lu_idx < len(lu_blocks) and current < instr_end:
            if not lunch_done and current >= LUNCH_START:
                lunch_end = current + LUNCH_DURATION
                slots.append(_make_slot(current, lunch_end, "Lunch Break", "-"))
                current = lunch_end
                lunch_done = True
                continue

            next_barrier = LUNCH_START if (not lunch_done) else instr_end
            available = next_barrier - current

            if available <= 0:
                break

            block = lu_blocks[lu_idx]
            topics_text = "\n".join(block["topic_labels"])
            t_methods = ", ".join(block["methods"])
            lu_extra = {
                "lu_num": block["lu_num"],
                "lu_title": block["lu_title"],
                "is_contd": is_contd,
            }

            if lu_remaining <= available:
                end = _round5(current + int(round(lu_remaining)))
                slots.append(_make_slot(current, end, topics_text, t_methods, **lu_extra))
                current = end
                lu_idx += 1
                if lu_idx < len(lu_blocks):
                    lu_remaining = lu_blocks[lu_idx]["duration_mins"]
                    is_contd = False

            elif available >= MIN_SESSION:
                slots.append(_make_slot(current, next_barrier, topics_text, t_methods, **lu_extra))
                lu_remaining -= available
                is_contd = True
                current = next_barrier

            else:
                if not lunch_done:
                    lunch_end = current + LUNCH_DURATION
                    slots.append(_make_slot(current, lunch_end, "Lunch Break", "-"))
                    current = lunch_end
                    lunch_done = True
                else:
                    slots.append(_make_slot(current, next_barrier, "Break", "-"))
                    current = next_barrier

        # Ensure lunch is placed even if LUs ended early
        if not lunch_done:
            if current < LUNCH_START:
                slots.append(_make_slot(current, LUNCH_START, "Break", "-"))
                current = LUNCH_START
            lunch_end = current + LUNCH_DURATION
            slots.append(_make_slot(current, lunch_end, "Lunch Break", "-"))
            current = lunch_end

        # Assessment on last day
        if is_last_day and assess_mins > 0:
            if current < ASSESS_START:
                slots.append(_make_slot(current, ASSESS_START, "Break", "-"))
                current = ASSESS_START

            am_details = context.get("Assessment_Methods_Details", [])
            # Assessment always fills from ASSESS_START to DAY_END (4-6 PM)
            am_end = DAY_END
            total_am_mins = am_end - current
            if am_details and len(am_details) > 1:
                # Group all assessments into one row, split time equally
                per_am = total_am_mins // len(am_details)
                lines = []
                for am in am_details:
                    name = am.get("Assessment_Method", "Assessment")
                    lines.append(f"Assessment: {name} ({per_am} mins)")
                label = "\n".join(lines)
                slots.append(_make_slot(current, am_end, label, "Assessment"))
                current = am_end
            elif am_details:
                name = am_details[0].get("Assessment_Method", "Assessment")
                label = f"Assessment: {name} ({total_am_mins} mins)"
                slots.append(_make_slot(current, am_end, label, "Assessment"))
                current = am_end
            else:
                slots.append(_make_slot(current, am_end, "Assessment", "Assessment"))
                current = am_end

        # Fill remaining time to end of day
        if current < DAY_END:
            slots.append(_make_slot(current, DAY_END, "Break", "-"))

        days[day] = slots

    return {
        "num_days": num_days,
        "instructional_hours": instr_hours,
        "assessment_hours": assess_hours,
        "per_topic_mins": round(per_topic, 1),
        "days": days,
    }


# =============================================================================
# Template-based Cover Page & Version Control
# =============================================================================

LP_TEMPLATE_PATH = ".claude/skills/generate_lesson_plan/templates/LP_template_v2.docx"


def _render_lp_template(context: dict, company: dict) -> Document:
    """Render the LP docxtpl template with cover page and version control.

    Returns a python-docx Document ready for appending schedule tables.
    """
    from generate_ap_fg_lg.utils.helper import process_logo_image

    tpl = DocxTemplate(LP_TEMPLATE_PATH)

    org_name = company.get("name", "") if company else ""

    # Build template context (matches AP/FG/LG variable names)
    current_date = datetime.now()
    tpl_ctx = {
        "Course_Title": context.get("Course_Title", "Course"),
        "TGS_Ref_No": context.get("TGS_Ref_No", ""),
        "Name_of_Organisation": org_name,
        "UEN": company.get("uen", "") if company else "",
        "Date": current_date.strftime("%d %b %Y"),
        "Year": str(current_date.year),
    }

    # Process logo via shared helper (same as AP/FG/LG)
    try:
        if org_name:
            tpl_ctx["company_logo"] = process_logo_image(tpl, org_name)
        else:
            tpl_ctx["company_logo"] = ""
    except (FileNotFoundError, Exception):
        tpl_ctx["company_logo"] = ""

    tpl.render(tpl_ctx, autoescape=True)

    # Save rendered template to temp, then reload as python-docx Document
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tpl.save(tmp.name)
        return Document(tmp.name)


# =============================================================================
# DOCX 4-Column Table Generator
# =============================================================================

def _set_header_cell(cell, text: str):
    """Style a table header cell with steel blue background and white text."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(255, 255, 255)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tc_pr = cell._element.get_or_add_tcPr()
    shading = tc_pr.makeelement(
        qn("w:shd"), {qn("w:fill"): "4472C4", qn("w:val"): "clear"},
    )
    tc_pr.append(shading)


def _add_colored_heading(doc, text: str, level: int = 2):
    """Add a heading with steel blue color."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = HEADING_COLOR


def _add_lu_header_row(table, text: str):
    """Add a Learning Unit header row that spans all columns with light blue background."""
    row = table.add_row()
    merged = row.cells[0].merge(row.cells[-1])
    merged.text = ""
    p = merged.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)  # Dark blue text
    tc_pr = merged._element.get_or_add_tcPr()
    shading = tc_pr.makeelement(
        qn("w:shd"), {qn("w:fill"): "D6E4F0", qn("w:val"): "clear"},
    )
    tc_pr.append(shading)


def _set_fixed_table_layout(table, col_widths):
    """Force a table to use fixed column layout at the XML level.

    This prevents Word from auto-resizing columns and ensures the table
    stays within page margins.
    """
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = tbl._new_tblPr()
        tbl.insert(0, tblPr)
    # Set layout to fixed
    tblLayout = tblPr.makeelement(qn("w:tblLayout"), {qn("w:type"): "fixed"})
    tblPr.append(tblLayout)
    # Set explicit total width in twips (dxa: 1 inch = 1440 twips)
    # Inches() returns EMU; 1 twip = 635 EMU
    total_twips = sum(int(w) // 635 for w in col_widths)
    # Remove existing tblW if any
    for existing in tblPr.findall(qn("w:tblW")):
        tblPr.remove(existing)
    tblW = tblPr.makeelement(
        qn("w:tblW"),
        {qn("w:w"): str(total_twips), qn("w:type"): "dxa"},
    )
    tblPr.append(tblW)


def generate_lesson_plan_docx(context: dict, schedule_data: dict, company: dict = None) -> str:
    """Generate a Lesson Plan DOCX with cover page, version control, and 4-column tables.

    Uses docxtpl template for cover page & version control (matching AP/FG/LG),
    then appends schedule tables programmatically.

    Args:
        context: Course context dict with Course_Title, etc.
        schedule_data: Output from build_lesson_plan_schedule().
        company: Company dict with name, uen, logo keys.

    Returns:
        Path to the generated DOCX file.
    """
    if company is None:
        company = {}

    # Render cover page & version control from template
    doc = _render_lp_template(context, company)

    # Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(f"Lesson Plan: {context.get('Course_Title', 'Course')}")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.name = "Calibri"

    # Metadata
    num_days = schedule_data.get("num_days", 1)
    instr_hrs = schedule_data.get("instructional_hours", 0)
    assess_hrs = schedule_data.get("assessment_hours", 0)

    methods = extract_unique_instructional_methods(context)
    methods_text = ", ".join(sorted(methods)) if methods else "N/A"

    org_name = company.get("name", "")
    metadata_lines = [
        f"Course Duration: {num_days} Day(s) (9:00 AM - 6:00 PM daily)",
        f"Total Training Hours: {instr_hrs} hrs",
        f"Total Assessment Hours: {assess_hrs} hrs",
        f"Instructional Methods: {methods_text}",
    ]
    if org_name:
        metadata_lines.insert(0, f"Organisation: {org_name}")

    for line in metadata_lines:
        p = doc.add_paragraph(line)
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(2)

    # Day-by-day 4-column tables
    days = schedule_data.get("days", {})
    # Columns sized to fit within 7.0" table budget
    col_widths = [Inches(1.3), Inches(0.7), Inches(3.0), Inches(2.0)]

    for day_num in sorted(days.keys()):
        _add_colored_heading(doc, f"Day {day_num}")

        slots = days[day_num]
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        table.autofit = False
        _set_fixed_table_layout(table, col_widths)

        for i, width in enumerate(col_widths):
            table.columns[i].width = width

        # Header row
        for i, width in enumerate(col_widths):
            table.rows[0].cells[i].width = width
        _set_header_cell(table.rows[0].cells[0], "Timing")
        _set_header_cell(table.rows[0].cells[1], "Duration")
        _set_header_cell(table.rows[0].cells[2], "Description")
        _set_header_cell(table.rows[0].cells[3], "Instructional Methods")

        # Data rows (with LU header rows)
        for slot in slots:
            slot_lu = slot.get("lu_num")
            if slot_lu:
                # Add LU header row
                is_contd = slot.get("is_contd", False)
                header = f"LU{slot_lu}: {slot.get('lu_title', '')}"
                if is_contd:
                    header += " (Cont'd)"
                _add_lu_header_row(table, header)

            row = table.add_row()
            # Timing, Duration, Methods - single-line cells
            for j, val in enumerate([
                slot.get("timing", ""),
                slot.get("duration", ""),
            ]):
                cell = row.cells[j]
                cell.width = col_widths[j]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(val)
                run.font.name = "Calibri"
                run.font.size = Pt(10)

            # Description cell - multi-line topics (one per line)
            desc_cell = row.cells[2]
            desc_cell.width = col_widths[2]
            desc_cell.text = ""
            desc_text = slot.get("description", "")
            lines = desc_text.split("\n") if "\n" in desc_text else [desc_text]
            for line_idx, line in enumerate(lines):
                if line_idx == 0:
                    p = desc_cell.paragraphs[0]
                else:
                    p = desc_cell.add_paragraph()
                p.paragraph_format.space_after = Pt(1)
                run = p.add_run(line)
                run.font.name = "Calibri"
                run.font.size = Pt(10)

            # Methods cell
            methods_cell = row.cells[3]
            methods_cell.width = col_widths[3]
            methods_cell.text = ""
            p = methods_cell.paragraphs[0]
            run = p.add_run(slot.get("methods", ""))
            run.font.name = "Calibri"
            run.font.size = Pt(10)

        doc.add_paragraph()  # Spacing between days

    # Margins are controlled by the template - no override needed

    # Save to temp file
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        return tmp.name
