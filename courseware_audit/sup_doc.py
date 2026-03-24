"""
Courseware Audit Module

Two modes:
  1. Quick Check — after generation, uses session state data
  2. Upload & Check — upload existing CP + courseware documents

Cross-checks selected fields against the CP (source of truth).
"""

import streamlit as st
import os
import tempfile
import pandas as pd
from docx import Document

# Optional PDF support
PYMUPDF_AVAILABLE = False
try:
    import pymupdf
    PYMUPDF_AVAILABLE = True
except ImportError:
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        pass

# Optional Excel support
OPENPYXL_AVAILABLE = False
try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    pass


DOC_TYPES = ["AP", "FG", "LG", "LP"]


def _get_audit_checks():
    """Load audit check items from the audit skill template (single source of truth).

    Returns list of (display_name, field_key, field_type, applicable_types).
    All rules live in: courseware_agents/audit/templates/audit_extraction.md
    """
    from courseware_agents.audit.audit_agent import get_audit_checks
    return get_audit_checks()


def extract_text_from_docx(file_bytes):
    """Extract all text from a DOCX file."""
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    try:
        doc = Document(tmp_path)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    finally:
        os.remove(tmp_path)


def extract_text_from_pdf(file_bytes):
    """Extract text from a PDF file."""
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    try:
        if PYMUPDF_AVAILABLE:
            doc = pymupdf.open(tmp_path)
            parts = []
            for page in doc:
                text = page.get_text()
                if text.strip():
                    parts.append(text.strip())
            doc.close()
            return "\n".join(parts)
        else:
            reader = PdfReader(tmp_path)
            parts = []
            for page in reader.pages:
                text = page.extract_text() or ""
                if text.strip():
                    parts.append(text.strip())
            return "\n".join(parts)
    finally:
        os.remove(tmp_path)


def extract_text_from_excel(file_bytes):
    """Extract text from an Excel file (.xlsx / .xls).

    For CP Excel files (SSG format), only extract from content sheets —
    skip reference/lookup sheets (TSC_Skill_Codes, checks, etc.) that
    contain thousands of rows of irrelevant data which confuse the AI.
    """
    if not OPENPYXL_AVAILABLE:
        return ""
    with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    try:
        wb = openpyxl.load_workbook(tmp_path, data_only=True)
        # Skip reference/lookup sheets that pollute extraction
        skip_sheets = {
            "tsc_skill_codes", "tsc_ccs_k&a", "tsc_sector_track", "ssoc_5d",
            "checks", "other_ref", "change_log", "read before filling",
            "annex a", "annex b", "annex c",
            "(not for tp use) outcome", "4 - declarations",
        }
        parts = []
        for ws in wb.worksheets:
            if ws.title.strip().lower() in skip_sheets:
                continue
            parts.append(f"--- Sheet: {ws.title} ---")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    finally:
        os.remove(tmp_path)


def extract_cp_fields_from_excel(file_bytes):
    """Directly extract CP fields from SSG Excel format using known cell positions.

    This is more reliable than AI extraction for structured Excel CPs.
    Returns dict of CP fields, or None if the format is not recognised.
    """
    if not OPENPYXL_AVAILABLE:
        return None
    import re
    with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    try:
        wb = openpyxl.load_workbook(tmp_path, data_only=True)

        fields = {}

        # ── Sheet: 1 - Course Particulars ──
        cp_sheet = None
        for name in wb.sheetnames:
            if "course particulars" in name.lower() or name == "1 - Course Particulars":
                cp_sheet = wb[name]
                break

        if cp_sheet:
            # Course Title — typically C3
            for row in range(2, 10):
                label = cp_sheet.cell(row=row, column=2).value
                val = cp_sheet.cell(row=row, column=3).value
                if label and "course title" in str(label).lower() and val:
                    fields["course_title"] = str(val).strip()
                    break

            # TSC Code — typically C10 (format: "CODE: Title")
            for row in range(8, 20):
                label = cp_sheet.cell(row=row, column=2).value
                val = cp_sheet.cell(row=row, column=3).value
                if label and "tsc" in str(label).lower() and val and str(val).strip():
                    tsc_full = str(val).strip()
                    # Parse "ICT-BAS-0047-1.1: Title Name"
                    m = re.match(r'^([A-Z]{2,4}-[A-Z]{2,4}-\d{4}-[\d.]+)\s*:\s*(.+)', tsc_full)
                    if m:
                        fields["tsc_code"] = m.group(1)
                        fields["tsc_title"] = m.group(2).strip()
                    else:
                        fields["tsc_code"] = tsc_full
                    break

            # Company Name — typically C2
            for row in range(1, 5):
                label = cp_sheet.cell(row=row, column=2).value
                val = cp_sheet.cell(row=row, column=3).value
                if label and ("training provider" in str(label).lower() or "organisation" in str(label).lower()) and val:
                    fields["company_name"] = str(val).strip()
                    break

        # ── Sheet: 3 - Summary ──
        sum_sheet = None
        for name in wb.sheetnames:
            if "summary" in name.lower() or name == "3 - Summary":
                sum_sheet = wb[name]
                break

        if sum_sheet:
            # Scan for training/assessment hours
            for row in sum_sheet.iter_rows(min_row=1, max_row=10, values_only=False):
                for cell in row:
                    val = str(cell.value or "").strip().lower()
                    if "total instructional" in val or "total training" in val:
                        # Hours value is in the next column
                        hours_cell = sum_sheet.cell(row=cell.row, column=cell.column + 1)
                        if hours_cell.value:
                            fields["training_hours"] = str(hours_cell.value).strip()
                    if "total assessment" in val:
                        hours_cell = sum_sheet.cell(row=cell.row, column=cell.column + 1)
                        if hours_cell.value:
                            fields["assessment_hours"] = str(hours_cell.value).strip()

            # Extract topics from learning unit rows
            topics = []
            for row in sum_sheet.iter_rows(min_row=6, max_row=sum_sheet.max_row, values_only=False):
                # Topics are typically in column E, with "T1:", "T2:" prefixes
                for cell in row:
                    val = str(cell.value or "").strip()
                    if re.search(r'T\d+\s*:', val):
                        # Extract individual topic titles
                        for t_match in re.finditer(r'T\d+\s*:\s*([^\n]+)', val):
                            topic = t_match.group(1).strip()
                            if topic and topic not in topics:
                                topics.append(topic)
            if topics:
                fields["topics"] = topics

        return fields if fields else None
    except Exception:
        return None
    finally:
        os.remove(tmp_path)


def _extract_text(file_bytes, ext):
    """Extract text from a file based on its extension."""
    if ext == "docx" or ext == "doc":
        return extract_text_from_docx(file_bytes)
    elif ext == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in ("xlsx", "xls"):
        return extract_text_from_excel(file_bytes)
    return ""


def _normalize(val):
    """Normalize a string for comparison."""
    if val is None:
        return ""
    import re
    s = str(val).strip().lower()
    s = re.sub(r'\.+$', '', s)
    s = s.replace('.', '')
    s = s.replace('-', ' ')
    s = re.sub(r'\s*\([^)]*\)\s*', ' ', s)
    s = re.sub(r'^(t\d+|lo\d+|lu\d+|topic\s*\d+)\s*[:.\-]\s*', '', s)
    # Remove common business suffixes for company name matching
    s = re.sub(r'\b(pte|ltd|private|limited|inc|llc|academy|institute|school|centre|center|group|sdn|bhd)\b', '', s)
    s = re.sub(r'\s+', ' ', s).strip()
    return s


def _normalize_number(val_str):
    """Extract numeric value from a duration string, normalized to hours.

    Handles: '22 hrs', '22.0', '30 mins', '1.5 hours', '90 minutes',
             '1 hour 30 minutes', '30 hour 0 minutes', '1 hr', '30 min'.
    """
    import re
    if not val_str:
        return None
    s = str(val_str).strip().lower()

    # Compound format: "X hour(s) Y minute(s)" (e.g. "30 hour 0 minutes")
    compound = re.match(r'(\d+\.?\d*)\s*hours?\s+(\d+\.?\d*)\s*min', s)
    if compound:
        hours = float(compound.group(1))
        mins = float(compound.group(2))
        total = hours + mins / 60
        return int(total) if total == int(total) else round(total, 2)

    # Minutes only (e.g. "30 mins", "90 minutes") — but NOT compound format
    if re.search(r'\bmin', s) and not re.search(r'\bhour', s):
        match = re.search(r'(\d+\.?\d*)', s)
        if match:
            mins = float(match.group(1))
            hours = mins / 60
            return int(hours) if hours == int(hours) else round(hours, 2)
        return None

    # Hours (default): "22 hrs", "30.0", "1.5 hours"
    match = re.search(r'(\d+\.?\d*)', s)
    if match:
        num = float(match.group(1))
        return int(num) if num == int(num) else num
    return None


def _extract_cp_fields(cp_context: dict) -> dict:
    """Extract audit-comparable fields from CP interpreter output (source of truth)."""
    learning_units = cp_context.get("Learning_Units", [])

    topics = []
    learning_outcomes = []
    k_statements = []
    a_statements = []
    assessment_methods = set()
    instructional_methods = set()

    for lu in learning_units:
        for topic in lu.get("Topics", []):
            title = topic.get("Topic_Title", "")
            if title:
                topics.append(title)
        lo = lu.get("LO", "")
        if lo:
            learning_outcomes.append(lo)
        for k in lu.get("K_numbering_description", []):
            desc = k.get("Description", "")
            k_num = k.get("K_number", "")
            if desc:
                k_statements.append(f"{k_num}: {desc}" if k_num else desc)
        for a in lu.get("A_numbering_description", []):
            desc = a.get("Description", "")
            a_num = a.get("A_number", "")
            if desc:
                a_statements.append(f"{a_num}: {desc}" if a_num else desc)
        for am in lu.get("Assessment_Methods", []):
            if am:
                assessment_methods.add(am)
        for im in lu.get("Instructional_Methods", []):
            if im:
                instructional_methods.add(im)

    return {
        "course_title": cp_context.get("Course_Title"),
        "tgs_ref_code": cp_context.get("TGS_Ref_No"),
        "topics": topics,
        "training_hours": cp_context.get("Total_Training_Hours"),
        "assessment_hours": cp_context.get("Total_Assessment_Hours"),
        "company_name": cp_context.get("Name_of_Organisation"),
        "uen": cp_context.get("UEN"),
        "learning_outcomes": learning_outcomes,
        "k_statements": k_statements,
        "a_statements": a_statements,
        "assessment_methods": sorted(assessment_methods),
        "instructional_methods": sorted(instructional_methods),
        "tsc_code": cp_context.get("TSC_Code"),
        "tsc_title": cp_context.get("TSC_Title"),
    }


def _compare_field(cp_val, doc_val, field_type, extra_cp=None, field_key=None):
    """Compare a document field value against CP.
    Returns (status, detail).

    If a field is not found in the document, it is treated as N/A (skip),
    not a failure — different courseware documents contain different fields.
    Only flag as mismatch when BOTH CP and doc have a value but they differ.
    """
    if field_type == "string":
        cp_norm = _normalize(cp_val)
        doc_norm = _normalize(doc_val)
        if not cp_norm and not doc_norm:
            return "skip", "N/A"
        if not doc_norm:
            return "skip", "Not in document"
        if not cp_norm:
            # Document has value but CP doesn't — show as pass (info present)
            return "match", "Found in document"
        if cp_norm == doc_norm:
            return "match", "Matches CP"
        # Substring match: if one contains the other, it's a match
        # Handles "Problem Solving with AI" vs "Fast and Creative Problem Solving with Generative AI"
        if cp_norm in doc_norm or doc_norm in cp_norm:
            return "match", "Matches CP"
        # TSC code: match if numeric part is the same (e.g. ICT-BAS-0047-1.1 vs ICT-INT-0047-1.1)
        import re as _re
        if field_key and "tsc_code" in field_key:
            cp_nums = _re.findall(r'\d+', str(cp_val or ""))
            doc_nums = _re.findall(r'\d+', str(doc_val or ""))
            if cp_nums and doc_nums and cp_nums == doc_nums:
                return "match", "Matches CP"
        return "mismatch", "Does NOT match CP"

    elif field_type == "list":
        cp_count = len([v for v in (cp_val or []) if v])
        doc_count = len([v for v in (doc_val or []) if v])
        if cp_count == 0 and doc_count == 0:
            return "skip", "N/A"
        if doc_count == 0:
            return "skip", "Not in document"
        if cp_count == 0:
            return "skip", "Not in CP"

        # Content match: extract domain-specific keywords and check overlap
        import re as _re
        _stop_words = {
            'the', 'and', 'for', 'with', 'that', 'this', 'from', 'will', 'are',
            'was', 'has', 'have', 'been', 'able', 'learner', 'identify', 'apply',
            'evaluate', 'develop', 'use', 'using', 'based', 'types', 'methods',
            'techniques', 'how', 'what', 'can', 'may', 'its', 'their', 'into',
            'such', 'also', 'between', 'within', 'through', 'about', 'which',
            'criteria', 'factors', 'various', 'different', 'key', 'relevant',
            'appropriate', 'effective', 'process', 'processes', 'systems',
            'understanding', 'developing', 'applying', 'identifying', 'making',
            'analysis', 'plan', 'plans', 'planning', 'overview', 'introduction',
            'describe', 'explain', 'discuss', 'learners', 'knowledge', 'skills',
            'ability', 'competency', 'performance', 'assessment', 'training',
            'course', 'unit', 'module', 'learning', 'outcome', 'outcomes',
            'topics', 'topic', 'title', 'name', 'statement', 'statements',
        }

        def _extract_keywords(items):
            """Extract domain-specific keywords (4+ chars) from list items."""
            words = set()
            for item in items:
                if not item:
                    continue
                cleaned = _re.sub(r'^(T\d+|LU\d+|K\d+|A\d+|ELO?\d+|Topic\s*\d+)\s*[:.]?\s*', '', str(item), flags=_re.IGNORECASE).strip()
                for w in _re.findall(r'[a-zA-Z]{4,}', cleaned.lower()):
                    if w not in _stop_words:
                        words.add(w)
            return words

        def _extract_stems(items):
            """Extract word stems (first 5 chars) for fuzzy matching."""
            stems = set()
            for item in items:
                if not item:
                    continue
                cleaned = _re.sub(r'^(T\d+|LU\d+|K\d+|A\d+|ELO?\d+|Topic\s*\d+)\s*[:.]?\s*', '', str(item), flags=_re.IGNORECASE).strip()
                for w in _re.findall(r'[a-zA-Z]{5,}', cleaned.lower()):
                    if w not in _stop_words:
                        stems.add(w[:5])
            return stems

        cp_keywords = _extract_keywords(cp_val or [])
        doc_keywords = _extract_keywords(doc_val or [])

        if not cp_keywords or not doc_keywords:
            return "match", "Matches CP"

        # Level 1: exact keyword overlap
        overlap = cp_keywords & doc_keywords
        cp_ratio = len(overlap) / len(cp_keywords) if cp_keywords else 0
        doc_ratio = len(overlap) / len(doc_keywords) if doc_keywords else 0
        best_ratio = max(cp_ratio, doc_ratio)

        if best_ratio >= 0.2:
            return "match", "Matches CP"

        # Level 2: stem matching (handles rephrasing like "identifying" vs "identification")
        cp_stems = _extract_stems(cp_val or [])
        doc_stems = _extract_stems(doc_val or [])
        stem_overlap = cp_stems & doc_stems
        stem_ratio = len(stem_overlap) / len(cp_stems) if cp_stems else 0

        if stem_ratio >= 0.2:
            return "match", "Matches CP"

        return "mismatch", "Topics do not match CP"

    elif field_type == "duration_field":
        cp_num = _normalize_number(cp_val)
        doc_num = _normalize_number(doc_val)
        if cp_num is None and doc_num is None:
            return "skip", "N/A"
        if doc_num is None:
            return "skip", "Not in document"
        if cp_num is None:
            return "skip", "Not in CP"
        if cp_num == doc_num:
            return "match", "Matches CP"
        # Check if document shows total hours (training + assessment)
        # or a portion of the total (e.g. individual assessment = 30 mins of 1 hr total)
        if extra_cp and field_key:
            cp_training = _normalize_number(extra_cp.get("training_hours"))
            cp_assessment = _normalize_number(extra_cp.get("assessment_hours"))
            if cp_training is not None and cp_assessment is not None:
                total = cp_training + cp_assessment
                if doc_num == total:
                    return "match", "Matches total course hours"
                if doc_num == cp_training or doc_num == cp_assessment:
                    return "match", "Matches CP"
            # Doc shows a portion of the CP value (e.g. 30 mins is part of 1 hr)
            # Valid when individual assessment methods split the total
            if cp_num and doc_num and doc_num <= cp_num:
                return "match", "Matches CP"
        return "mismatch", f"CP={cp_num} vs Doc={doc_num}"

    return "unknown", ""


def _format_val(val):
    """Format a value for display in the comparison table."""
    if isinstance(val, list):
        return ", ".join(str(v) for v in val) if val else "N/A"
    return str(val) if val else "N/A"


def _status_icon(status):
    """Return a visual status indicator."""
    if status == "match":
        return "Matches CP"
    elif status == "mismatch":
        return "MISMATCH"
    elif status == "missing_in_doc":
        return "Missing in doc"
    elif status == "skip":
        return "N/A"
    return "-"


def run_audit(cp_fields: dict, doc_results: dict, selected_checks: list[str]):
    """Run audit on selected checks only.

    Args:
        cp_fields: Extracted CP fields (source of truth)
        doc_results: Dict of {doc_label: extracted_fields_dict}
        selected_checks: List of field_key strings to check

    Returns:
        List of result rows for the comparison table.
    """
    results = []
    cell_statuses = []  # Parallel list: dict of {col_name: status} per row
    cell_details = []   # Parallel list: dict of {col_name: detail_string} per row

    for display_name, field_key, field_type, applicable_types in _get_audit_checks():
        if field_key not in selected_checks:
            continue

        cp_val = cp_fields.get(field_key)
        row = {"Field": display_name, "CP (Source of Truth)": _format_val(cp_val)}
        row_statuses = {"Field": "neutral", "CP (Source of Truth)": "neutral"}
        row_details = {}

        for doc_label, doc_data in doc_results.items():
            # Detect doc type from label
            doc_type = ""
            for dt in DOC_TYPES:
                if doc_label.upper().startswith(dt):
                    doc_type = dt
                    break

            if applicable_types and doc_type and doc_type not in applicable_types:
                row[doc_label] = "N/A"
                row_statuses[doc_label] = "skip"
                continue

            doc_val = doc_data.get(field_key)
            status, detail = _compare_field(cp_val, doc_val, field_type, extra_cp=cp_fields, field_key=field_key)
            row[doc_label] = _format_val(doc_val)
            row_statuses[doc_label] = status
            row_details[doc_label] = detail

        results.append(row)
        cell_statuses.append(row_statuses)
        cell_details.append(row_details)

    return results, cell_statuses, cell_details


def _fix_text_in_docx(file_bytes: bytes, replacements: list[tuple[str, str]]) -> tuple[bytes, int]:
    """Apply text replacements to a DOCX file. Returns (fixed_bytes, fix_count)."""
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    try:
        doc = Document(tmp_path)
        fix_count = 0

        for old_text, new_text in replacements:
            if not old_text or not new_text or old_text == new_text:
                continue

            for para in doc.paragraphs:
                if old_text in para.text:
                    for run in para.runs:
                        if old_text in run.text:
                            run.text = run.text.replace(old_text, new_text)
                            fix_count += 1

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if old_text in cell.text:
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    if old_text in run.text:
                                        run.text = run.text.replace(old_text, new_text)
                                        fix_count += 1

            for section in doc.sections:
                for header_footer in [section.header, section.footer]:
                    if header_footer:
                        for para in header_footer.paragraphs:
                            if old_text in para.text:
                                for run in para.runs:
                                    if old_text in run.text:
                                        run.text = run.text.replace(old_text, new_text)
                                        fix_count += 1

        out_path = tmp_path + "_fixed.docx"
        doc.save(out_path)
        with open(out_path, "rb") as f:
            result = f.read()
        os.remove(out_path)
        return result, fix_count
    finally:
        os.remove(tmp_path)


def _build_replacements(cp_fields: dict, doc_fields: dict, selected_checks: list[str]) -> list[tuple[str, str]]:
    """Build (old_text, new_text) replacements from mismatches for selected checks only."""
    replacements = []

    for _, field_key, field_type, _ in _get_audit_checks():
        if field_key not in selected_checks:
            continue

        cp_val = cp_fields.get(field_key)
        doc_val = doc_fields.get(field_key)
        if not cp_val or not doc_val:
            continue

        if field_type == "string":
            cp_str = str(cp_val).strip()
            doc_str = str(doc_val).strip()
            if cp_str and doc_str and _normalize(cp_str) != _normalize(doc_str):
                replacements.append((doc_str, cp_str))

        elif field_type == "duration_field":
            cp_str = str(cp_val).strip()
            doc_str = str(doc_val).strip()
            if cp_str and doc_str and _normalize_number(cp_str) != _normalize_number(doc_str):
                replacements.append((doc_str, cp_str))

        elif field_type == "list":
            cp_list = cp_val if isinstance(cp_val, list) else []
            doc_list = doc_val if isinstance(doc_val, list) else []
            for i, doc_item in enumerate(doc_list):
                doc_str = str(doc_item).strip()
                if not doc_str:
                    continue
                matched = any(_normalize(str(cp_item)) == _normalize(doc_str) for cp_item in cp_list)
                if not matched and i < len(cp_list):
                    cp_str = str(cp_list[i]).strip()
                    if cp_str and doc_str:
                        replacements.append((doc_str, cp_str))

    return replacements


def _get_clear_id():
    """Get unique suffix for file uploaders — changes on clear to reset them."""
    return st.session_state.get("_audit_clear_count", 0)


def app():
    st.title("Courseware Audit")
    st.caption("Cross-check your courseware documents against the CP (source of truth).")

    # ── Initialize session state ──
    for key, default in [
        ("audit_cp_fields", {}),
        ("audit_docs", {}),
        ("audit_results", {}),
        ("audit_comparison", []),
        ("audit_cell_statuses", []),
        ("audit_cell_details", []),
        ("audit_selected_checks", []),
    ]:
        if key not in st.session_state:
            st.session_state[key] = default

    # ── Clear All button ──
    if st.session_state.get("_audit_clear_requested"):
        # Increment uploader counter to force file uploaders to reset
        clear_count = st.session_state.get("_audit_clear_count", 0) + 1
        st.session_state["_audit_clear_count"] = clear_count
        for key in list(st.session_state.keys()):
            if key.startswith("audit_") or key.startswith("_audit_"):
                if key != "_audit_clear_count":
                    del st.session_state[key]
        # Also clear CP/TGS session data so audit starts completely fresh
        st.session_state.pop("extracted_course_info", None)
        st.session_state.pop("_extract_cp_file_path", None)
        st.session_state.pop("_extract_cp_file_name", None)
        st.session_state.pop("_user_tgs_ref_no", None)
        st.session_state["_audit_clear_requested"] = False
        st.rerun()

    pass  # clear_id accessed via _get_clear_id() helper

    # ── Step 1: CP Source of Truth ──
    _has_session_cp = st.session_state.get("extracted_course_info") is not None
    cp_fields = {}

    # Detect if the CP file changed since last audit (prevents stale data)
    current_cp_name = st.session_state.get("_extract_cp_file_name", "")
    prev_audit_cp = st.session_state.get("_audit_last_cp_name", "")
    if current_cp_name and current_cp_name != prev_audit_cp:
        # CP changed — clear old audit data
        st.session_state["_audit_last_cp_name"] = current_cp_name
        st.session_state.audit_cp_fields = {}
        st.session_state.audit_results = {}
        st.session_state.audit_comparison = []
        st.session_state.audit_cell_statuses = []

    # Always show CP uploader — uploaded CP overrides session data
    st.subheader("Upload Course Proposal (CP)")
    if _has_session_cp:
        st.info("CP data available from Extract Course Info. You can also upload a different CP below.")
    cp_file = st.file_uploader(
        "Upload CP document",
        type=["docx", "doc", "xlsx", "xls"],
        key=f"audit_cp_uploader_{_get_clear_id()}",
    )

    if cp_file:
        # Detect if a NEW CP file was uploaded (different from previous)
        prev_cp_name = st.session_state.get("_audit_cp_filename", "")
        if cp_file.name != prev_cp_name:
            # New CP — clear old cached fields so they get re-extracted
            st.session_state["_audit_cp_filename"] = cp_file.name
            st.session_state.audit_cp_fields = {}
            st.session_state.audit_results = {}
            st.session_state.audit_comparison = []
            st.session_state.audit_cell_statuses = []

        st.session_state["audit_cp_file"] = cp_file

        # For Excel CPs: immediately extract key fields using direct cell parsing
        # This gives accurate CP fields without waiting for "Run Audit"
        ext = cp_file.name.rsplit(".", 1)[-1].lower()
        if ext in ("xlsx", "xls"):
            excel_fields = extract_cp_fields_from_excel(cp_file.getvalue())
            if excel_fields:
                cp_fields = excel_fields
                _apply_company_overrides(cp_fields)
                st.session_state.audit_cp_fields = cp_fields

        st.success(f"CP uploaded: {cp_file.name}")
    elif _has_session_cp:
        # No uploaded CP — use session data from Extract Course Info
        # IMPORTANT: Always try to re-extract from the actual CP file first,
        # because extracted_course_info may be stale if the user extracted
        # a different course since the last audit.
        cp_file_path = st.session_state.get("_extract_cp_file_path")
        cp_file_name = st.session_state.get("_extract_cp_file_name", "")
        _used_direct = False

        if cp_file_path and os.path.exists(cp_file_path):
            is_excel = cp_file_name.lower().endswith((".xlsx", ".xls"))
            if is_excel:
                try:
                    with open(cp_file_path, "rb") as f:
                        excel_fields = extract_cp_fields_from_excel(f.read())
                    if excel_fields and excel_fields.get("course_title"):
                        cp_fields = excel_fields
                        # Supplement with session data for fields Excel can't extract
                        cp_context = st.session_state["extracted_course_info"]
                        session_fields = _extract_cp_fields(cp_context)
                        for key in ("learning_outcomes", "k_statements", "a_statements",
                                    "assessment_methods", "instructional_methods"):
                            if not cp_fields.get(key) and session_fields.get(key):
                                cp_fields[key] = session_fields[key]
                        _used_direct = True
                except Exception:
                    pass

        if not _used_direct:
            # Fallback: use session extracted_course_info (DOCX CPs or Excel extraction failed)
            cp_context = st.session_state["extracted_course_info"]
            cp_fields = _extract_cp_fields(cp_context)

        _apply_company_overrides(cp_fields)
        st.session_state.audit_cp_fields = cp_fields
        st.success(f"CP data loaded from session: {cp_file_name or 'Extract Course Info'}.")

    # If we already extracted CP fields from a previous run, reuse them
    if not cp_fields:
        cp_fields = st.session_state.get("audit_cp_fields", {})

    # If no CP at all (no session data, no upload, no previous extraction), stop here
    if not cp_fields and "audit_cp_file" not in st.session_state:
        return

    # ── TGS Reference Code (always visible — CP usually doesn't contain it) ──
    # Auto-fill from Extract Course Info if available
    tgs_val = cp_fields.get("tgs_ref_code") or ""
    if not tgs_val:
        # Try from Extract Course Info session
        ext_info = st.session_state.get("extracted_course_info", {})
        if ext_info:
            tgs_val = ext_info.get("TGS_Ref_No", "") or ""
        if not tgs_val:
            tgs_val = st.session_state.get("_user_tgs_ref_no", "") or ""
    if not tgs_val:
        tgs_val = st.session_state.get("audit_tgs_input", "") or ""
    tgs_input = st.text_input(
        "Course Ref Code (TGS) *",
        value=tgs_val,
        placeholder="e.g. TGS-2024001234",
        key=f"audit_tgs_input_{_get_clear_id()}",
    )
    cp_fields["tgs_ref_code"] = tgs_input

    st.session_state.audit_cp_fields = cp_fields

    st.divider()

    # ── Step 2: Audit Checklist ──
    audit_checks = _get_audit_checks()
    all_field_keys = [field_key for _, field_key, _, _ in audit_checks]
    all_display_names = [display_name for display_name, _, _, _ in audit_checks]

    with st.expander("Audit Checklist", expanded=False):
        st.caption("All items are checked by default. Deselect any items you want to skip.")
        selected_checks = st.multiselect(
            "Fields to audit",
            options=all_field_keys,
            default=all_field_keys,
            format_func=lambda k: next(
                (d for d, fk, _, _ in audit_checks if fk == k), k
            ),
            key="audit_checklist",
            label_visibility="collapsed",
        )

    if not selected_checks:
        st.warning("Select at least one item to check.")
        return

    st.session_state.audit_selected_checks = selected_checks

    st.divider()

    # ── Step 3: Upload Courseware Documents ──
    st.subheader("Upload Courseware Documents")
    _upload_docs_section()

    # ── Run Audit ──
    _run_audit_button(cp_fields, selected_checks)

    # ── Display Results ──
    _display_results(selected_checks)


def _apply_company_overrides(cp_fields: dict):
    """Fill in company name + UEN from company list only if it matches the CP.

    The CP is the source of truth. The company list is only used when:
    1. The CP company name matches the company list (same company), OR
    2. The CP doesn't have a company name at all.
    This prevents overriding with a completely different company's data.
    """
    try:
        from company.company_manager import get_selected_company
        default_org = get_selected_company()
        if not default_org:
            return

        cp_company = cp_fields.get("company_name", "")
        list_company = default_org.get("name", "")

        # Check if the company list matches the CP's company
        cp_norm = _normalize(cp_company)
        list_norm = _normalize(list_company)
        is_same_company = (not cp_norm) or (cp_norm and list_norm and cp_norm == list_norm)

        if is_same_company:
            if not cp_fields.get("company_name") and default_org.get("name"):
                cp_fields["company_name"] = default_org["name"]
            if not cp_fields.get("uen") and default_org.get("uen"):
                cp_fields["uen"] = default_org["uen"]
    except Exception:
        pass


def _upload_docs_section():
    """Shared file uploader for courseware documents."""
    uploaded_files = st.file_uploader(
        "Upload AP, FG, LG, LP documents",
        type=["docx", "doc", "xlsx", "xls"],
        accept_multiple_files=True,
        key=f"audit_doc_uploader_{_get_clear_id()}",
    )

    # Rebuild audit_docs fresh from current uploads (don't accumulate old files)
    current_docs = {}

    if uploaded_files:
        # Clear old results when files change
        current_names = sorted(f.name for f in uploaded_files)
        prev_names = sorted(d["name"] for d in st.session_state.audit_docs.values()) if st.session_state.audit_docs else []
        if current_names != prev_names:
            st.session_state.audit_results = {}
            st.session_state.audit_comparison = []
            st.session_state.audit_cell_statuses = []

        for uploaded_file in uploaded_files:
            fname = uploaded_file.name
            # Auto-detect doc type from filename
            default_type = 0
            fname_upper = fname.upper()
            for i, dt in enumerate(DOC_TYPES):
                if fname_upper.startswith(dt) or f"_{dt}_" in fname_upper or f"_{dt}." in fname_upper:
                    default_type = i
                    break

            col1, col2 = st.columns([3, 1])
            with col1:
                st.text(fname)
            with col2:
                doc_type = st.selectbox(
                    "Type",
                    DOC_TYPES,
                    index=default_type,
                    key=f"audit_type_{fname}",
                    label_visibility="collapsed",
                )

            current_docs[f"{doc_type}: {fname}"] = {
                "file": uploaded_file,
                "type": doc_type,
                "name": fname,
            }

    st.session_state.audit_docs = current_docs


def _run_audit_button(cp_fields: dict, selected_checks: list[str]):
    """Run audit button and processing."""
    if st.button("Run Audit", type="primary"):
        docs = st.session_state.audit_docs
        if not docs:
            st.error("Please upload at least 1 courseware document to audit.")
            return

        from courseware_agents.audit.audit_agent import extract_audit_fields

        # ── Always re-extract CP fields fresh ──
        # This ensures no stale cache from previous courses

        # If no CP was uploaded in audit, but we have the CP file from Extract Course Info
        if "audit_cp_file" not in st.session_state:
            cp_file_path = st.session_state.get("_extract_cp_file_path")
            cp_file_name = st.session_state.get("_extract_cp_file_name", "")
            if cp_file_path and os.path.exists(cp_file_path):
                is_excel = cp_file_name.lower().endswith((".xlsx", ".xls"))
                if is_excel:
                    try:
                        with open(cp_file_path, "rb") as f:
                            excel_fields = extract_cp_fields_from_excel(f.read())
                        if excel_fields and excel_fields.get("course_title"):
                            # Use Excel extraction as the primary source
                            manual_tgs = cp_fields.get("tgs_ref_code", "")
                            for key, val in excel_fields.items():
                                if val:
                                    cp_fields[key] = val
                            if not cp_fields.get("tgs_ref_code"):
                                cp_fields["tgs_ref_code"] = manual_tgs
                            _apply_company_overrides(cp_fields)
                            st.session_state.audit_cp_fields = cp_fields
                    except Exception:
                        pass

        if "audit_cp_file" in st.session_state:
            cp_file = st.session_state["audit_cp_file"]
            fname = cp_file.name
            file_bytes = cp_file.getvalue()
            ext = fname.rsplit(".", 1)[-1].lower()

            with st.spinner(f"Extracting CP fields from {fname}..."):
                # For Excel CPs: use direct cell extraction first (more reliable)
                excel_fields = None
                if ext in ("xlsx", "xls"):
                    excel_fields = extract_cp_fields_from_excel(file_bytes)

                cp_text = _extract_text(file_bytes, ext)
                if cp_text.strip():
                    try:
                        cp_extracted = extract_audit_fields(cp_text, "CP")
                        # Preserve manually entered TGS if CP doesn't have one
                        manual_tgs = cp_fields.get("tgs_ref_code", "")
                        cp_fields = {
                            "course_title": cp_extracted.get("course_title"),
                            "tgs_ref_code": cp_extracted.get("tgs_ref_code") or manual_tgs,
                            "topics": cp_extracted.get("topics", []),
                            "training_hours": cp_extracted.get("durations", {}).get("training_hours") if isinstance(cp_extracted.get("durations"), dict) else None,
                            "assessment_hours": cp_extracted.get("durations", {}).get("assessment_hours") if isinstance(cp_extracted.get("durations"), dict) else None,
                            "company_name": cp_extracted.get("company_name"),
                            "uen": cp_extracted.get("uen"),
                            "learning_outcomes": cp_extracted.get("learning_outcomes", []),
                            "k_statements": cp_extracted.get("k_statements", []),
                            "a_statements": cp_extracted.get("a_statements", []),
                            "assessment_methods": cp_extracted.get("assessment_methods", []),
                            "instructional_methods": cp_extracted.get("instructional_methods", []),
                            "tsc_code": cp_extracted.get("tsc_code"),
                            "tsc_title": cp_extracted.get("tsc_title"),
                        }

                        # Override with direct Excel extraction (more reliable for key fields)
                        if excel_fields:
                            for key in ("course_title", "tsc_code", "tsc_title",
                                        "company_name", "training_hours", "assessment_hours"):
                                if excel_fields.get(key):
                                    cp_fields[key] = excel_fields[key]
                            if excel_fields.get("topics"):
                                cp_fields["topics"] = excel_fields["topics"]

                        _apply_company_overrides(cp_fields)
                        st.session_state.audit_cp_fields = cp_fields
                    except Exception as e:
                        st.error(f"Error extracting CP fields: {e}")
                        return
                else:
                    st.error(f"No text could be extracted from {fname}")
                    return

        if not cp_fields:
            st.error("No CP data available. Upload a CP document first.")
            return

        # ── Extract fields from each courseware document ──
        audit_results = {}
        for label, doc_info in docs.items():
            file_obj = doc_info["file"]
            fname = doc_info["name"]
            file_bytes = file_obj.getvalue()

            with st.spinner(f"Parsing {fname}..."):
                ext = fname.rsplit(".", 1)[-1].lower()
                text = _extract_text(file_bytes, ext)

            if not text.strip():
                st.warning(f"No text extracted from {fname}")
                continue

            with st.spinner(f"Auditing {label}..."):
                try:
                    result = extract_audit_fields(text, doc_info["type"])
                    # Map duration fields for our simplified checks
                    if isinstance(result.get("durations"), dict):
                        result["training_hours"] = result["durations"].get("training_hours")
                        result["assessment_hours"] = result["durations"].get("assessment_hours")
                    # Ensure list fields default to empty lists
                    for list_key in ("learning_outcomes", "k_statements", "a_statements",
                                     "assessment_methods", "instructional_methods"):
                        if list_key not in result or result[list_key] is None:
                            result[list_key] = []
                    audit_results[label] = result
                except Exception as e:
                    st.error(f"Error auditing {label}: {e}")
                    audit_results[label] = {}

        st.session_state.audit_results = audit_results

        # Always apply company overrides before comparison
        # so UEN and company_name are sourced from company list
        _apply_company_overrides(cp_fields)
        st.session_state.audit_cp_fields = cp_fields

        comparison, cell_statuses, cell_details = run_audit(cp_fields, audit_results, selected_checks)
        st.session_state.audit_comparison = comparison
        st.session_state.audit_cell_statuses = cell_statuses
        st.session_state.audit_cell_details = cell_details

        st.success(f"Audit complete. Checked {len(audit_results)} document(s).")
        st.rerun()


def _display_results(selected_checks):
    """Display audit results as a per-document checklist with pass/fail icons."""
    if not st.session_state.audit_comparison:
        return

    comparison = st.session_state.audit_comparison
    cell_statuses = st.session_state.get("audit_cell_statuses", [])
    cell_details = st.session_state.get("audit_cell_details", [])
    audit_results = st.session_state.audit_results
    cp_fields = st.session_state.audit_cp_fields

    st.subheader("Audit Results")

    # Reorganise data: per document → list of (field, cp_val, doc_val, status)
    # Get all doc columns (everything except "Field" and "CP (Source of Truth)")
    doc_columns = [c for c in comparison[0].keys() if c not in ("Field", "CP (Source of Truth)")] if comparison else []

    total_pass = 0
    total_fail = 0
    total_na = 0
    failed_docs = []

    for doc_col in doc_columns:
        doc_pass = 0
        doc_fail = 0
        doc_items = []

        for row_idx, row in enumerate(comparison):
            status = cell_statuses[row_idx].get(doc_col, "neutral") if row_idx < len(cell_statuses) else "neutral"
            detail = cell_details[row_idx].get(doc_col, "") if row_idx < len(cell_details) else ""
            field_name = row["Field"]
            cp_val = row.get("CP (Source of Truth)", "N/A")
            doc_val = row.get(doc_col, "N/A")

            if status == "skip":
                total_na += 1
                continue  # Not applicable — don't show

            if status == "match":
                icon = ":green[PASS]"
                doc_pass += 1
                total_pass += 1
            elif status == "mismatch":
                icon = ":red[FAIL]"
                doc_fail += 1
                total_fail += 1
            elif status in ("missing_in_doc", "missing"):
                icon = ":orange[MISSING]"
                doc_fail += 1
                total_fail += 1
            else:
                icon = ":gray[—]"
                total_na += 1
                continue

            doc_items.append((icon, field_name, cp_val, doc_val, status, detail))

        # Expandable box per document with full checklist
        all_passed = doc_fail == 0
        header_icon = ":white_check_mark:" if all_passed else ":x:"
        header_summary = f"{doc_pass} passed" if all_passed else f"{doc_pass} passed, {doc_fail} failed"

        with st.expander(f"{header_icon} {doc_col} — {header_summary}", expanded=(not all_passed)):
            for icon, field_name, cp_val, doc_val, status, detail in doc_items:
                display_val = str(doc_val)[:100] + ("..." if len(str(doc_val)) > 100 else "")

                if status == "match":
                    st.markdown(f"&ensp; :white_check_mark: &ensp; **{field_name}** — {display_val}")
                elif status == "mismatch":
                    st.markdown(f"&ensp; :x: &ensp; **{field_name}** — {display_val}")
                    st.caption(f"&emsp;&emsp;&emsp; Expected: {str(cp_val)[:100]}")
                elif status in ("missing_in_doc", "missing"):
                    st.markdown(f"&ensp; :warning: &ensp; **{field_name}** — Missing in document")

        if not all_passed:
            failed_docs.append(doc_col)

    # Overall summary
    st.divider()
    if total_fail == 0:
        st.success(f"All {total_pass} checks passed across all documents!")
    else:
        st.warning(f"{total_fail} issue(s) found in: {', '.join(failed_docs)}")

        pass  # Issues already shown inline above

    # ── Clear All ──
    st.divider()
    if st.button("Clear All & Start New Audit", key="audit_clear_btn"):
        st.session_state["_audit_clear_requested"] = True
        st.rerun()
